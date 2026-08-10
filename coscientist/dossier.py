"""Compile validated research artifacts into a readable report, then export it.

The Markdown this module emits is the report itself, not a dump of the pipeline that
produced it: narrative synthesis lives in ``narrative``, and everything here is
concerned with how that synthesis is laid out on a page.
"""

from __future__ import annotations

import re
import textwrap
import unicodedata
from collections import Counter
from collections.abc import Iterable, Sequence
from datetime import date, datetime
from html import escape
from io import BytesIO
from pathlib import Path

from .advisories import (
    ADVISORY_CHAPTER,
    AUTO_APPROVAL_WARNING,
    Advisory,
    advisory_pointer,
    run_advisories,
)
from .debate import (
    readable_exchange,
    standalone_opening,
    strip_rationale_label,
    strip_turn_label,
    unemphasised,
)
from .evidence import GROUNDING_REDIRECT_MARKER, names_a_document
from .flowchart import flowchart_drawing, flowchart_steps
from .markdown_render import (
    FIGURE_INDEX_HEADING,
    Code,
    Details,
    Heading,
    ListBlock,
    Para,
    Quote,
    Rule,
    Table,
    cjk_markup,
    flatten_fragment_links,
    has_cjk,
    inline_markup,
    number_figures_and_tables,
    parse_blocks,
    parse_inline,
    plain_text,
    strip_table_of_contents,
    table_of_contents,
)
from .models import FACET_PHRASES, FORKED_STAGES, MERGE_PRODUCER, Session
from .narrative import (
    _AGENT_NAMES,
    _CONTRACT_FIELD_NAMES,
    CRITERION_SECTIONS,
    DEEP_DIVE_PREAMBLE,
    DISCOVERY_STOOD_IN,
    REVIEW_SECTIONS,
    AdjudicationNote,
    Citation,
    IdeaBrief,
    ProvenanceNote,
    ResearchOverview,
    ResearchRecord,
    _cited_reference_standing,
    _counted,
    _joined_titles,
    _judge_column,
    _judge_label,
    _labelled_bullets,
    _listed,
    _names,
    _number_word,
    _opening,
    _plural,
    _reviewer_lead_in,
    _sentence,
    _stage_words,
    build_idea_briefs,
    evidence_integrity_cases,
    evidence_integrity_ideas,
    evidence_integrity_lines,
    load_record,
    shared_coherence_notes,
    shared_grounding_reach,
    shared_review_questions,
    shared_support_notices,
    stage_name,
    synthesize_overview,
)
from .parity import DEFAULT_ELO, REVIEW_CRITERIA, UNMEASURED_MOVEMENT


def _bullets(items: Iterable[str]) -> list[str]:
    return [f"- {item}" for item in items]


def _front_matter(record: ResearchRecord, overview: ResearchOverview) -> list[str]:
    """The cover block: the goal restated in the reference report's four fields."""
    session, plan = record.session, record.plan
    requirements = [item.rstrip(".") for item in (plan.constraints if plan else [])]
    requirements += [
        f"{item.input_type} is {item.status.replace('_', ' ')} — {item.reason.rstrip('.')}"
        for item in session.input_requirements
    ]
    # Both of the first two labels used to read as descriptions of what this run did.
    # "Research mode: experimental" over a document in which no experiment was
    # performed, and "Analysis scope: full requested analysis" over a run that
    # produced proposals and nothing else, are the two lines a reader meets before any
    # caveat -- and section 1 then spends a paragraph undoing them. The qualifier is
    # cheaper on the cover than the retraction is inside.
    attributes = [
        f"Research mode proposed: {session.research_mode.replace('_', ' ')} — no "
        "experiment was performed in producing this report",
        f"Intended claim: {plan.intended_claim if plan else 'hypothesis'}",
        "Analysis scope: "
        + (
            "literature-only"
            if session.literature_only
            else "the full requested analysis, carried out as desk work"
        ),
        f"Approval profile: {session.approval_profile}",
    ]
    attributes += [
        f"Assumption: {item.rstrip('.')}" for item in (plan.assumptions if plan else [])
    ]
    # The review dimensions close the attribute list in the reference reports,
    # because they are what every idea is scored on rather than facts about the goal.
    # Each dimension named for the reviewer that scores it, in the words the rest of
    # the report gives that reviewer. Written out by hand this list invented two names
    # nothing else in the document uses -- "scored one to five by the methods and
    # feasibility review" and "by the safety and governance review" -- so a reader on
    # page one who went looking for either found Methods and statistics reviewer and
    # Ethics, safety and governance reviewer, and no way to tell whether those were
    # the same pass.
    attributes += [
        f"{CRITERION_SECTIONS[criterion]}: scored one to five by "
        f"the {_reviewer_lead_in(reviewer).rstrip(':').lower()}"
        for reviewer, (criterion, _label) in REVIEW_CRITERIA.items()
    ]
    # Two different things used to be poured into one bulleted list under this
    # heading: what would make the goal met, and what the ideas were ranked against.
    # They are set by different stages and answer different questions, and a reader
    # taking the eleventh bullet for a success criterion when it is a ranking axis has
    # been misled by the layout rather than by anything the text says.
    success = [item.rstrip(".") for item in (plan.success_criteria if plan else [])]
    comparison = [
        item.rstrip(".")
        for item in (record.population.comparison_criteria if record.population else [])
    ]
    criteria: list[str] = []
    if success or not comparison:
        criteria += [
            "**Success criteria — what would make the goal met:**",
            "",
            *_bullets(success or ["No success criterion was recorded for this goal."]),
            "",
        ]
    if comparison:
        criteria += [
            "**Comparison criteria — what every idea was scored against:**",
            "",
            # These arrive as "Label: what it means". Run through _bullets they print
            # as one unbroken line each with the label buried behind a colon, and the
            # label is the part the rest of the report cites them by.
            *_labelled_bullets(comparison).split("\n"),
            "",
        ]
    return [
        f"# {overview.goal_title}",
        "",
        # The cover notice, which the PDF and the DOCX both print under the title
        # and the Markdown printed nowhere. Markdown is the copy that gets pasted
        # into a message or opened in an editor, and in that copy the statement
        # that none of this is a finding sat in an appendix two thousand seven
        # hundred lines down. Three exports of one report, and only two of them
        # told the reader what they were holding.
        f"*{_DEFAULT_NOTICE}*",
        "",
        "# Research Goal Details",
        "",
        "## Goal",
        "",
        session.question,
        "",
        "## Requirements",
        "",
        # Numbered, because section 1 refers to them by number and a reader following
        # "constraint three" back to an unnumbered bullet list has to count.
        *(
            [f"{index}. {item}" for index, item in enumerate(requirements, start=1)]
            if requirements
            else ["No constraint was recorded for this goal."]
        ),
        "",
        "## Attributes",
        "",
        *_bullets(attributes),
        "",
        "## Criteria",
        "",
        *criteria,
    ]


def _attribution(session: Session) -> str:
    """The reference reports' single-line provenance stamp, with its exact spacing."""
    stamp = (session.updated_at or "")[:10] or date.today().isoformat()
    return f"Prepared by \U0001f9ec  AI co-scientist on {stamp}. For research purposes only."


def _overview_body(
    record: ResearchRecord, overview: ResearchOverview, pointer: str = ""
) -> list[str]:
    lines = [
        "# Research Overview",
        "",
        "## Top ideas",
        "",
        f"### {overview.report_title}",
        "",
        _attribution(record.session),
        "",
    ]
    if pointer:
        # One line where nine paragraphs of warning used to be. It is here rather
        # than on the cover because this is the first page of prose a reader meets,
        # and it is a sentence rather than a callout because a callout is the thing
        # the warnings had already become.
        lines.extend([pointer, ""])
    for section in overview.sections:
        lines.extend([f"#### {section.number}. {section.title}", ""])
        # A grid sits after the paragraph that introduces it, which is why it carries
        # an index rather than being appended: the paragraphs after it in this section
        # read off the standings above them.
        grids = {grid.after: grid for grid in section.grids}
        for index, paragraph in enumerate(section.paragraphs):
            lines.extend([paragraph, ""])
            grid = grids.get(index)
            if grid:
                lines.extend(_grid(grid.columns, grid.rows))
        for subsection in section.subsections:
            lines.extend([f"##### {subsection.number} {subsection.title}", ""])
            for paragraph in subsection.paragraphs:
                lines.extend([paragraph, ""])
            if subsection.table_rows:
                lines.extend(_idea_table(subsection.table_rows))
    lines.extend(
        [
            "## Research directions",
            "",
            *_bullets(overview.research_directions),
            "",
            "## Review summary",
            "",
            *_bullets(overview.review_summary),
            "",
        ]
    )
    lines.extend(_governance_block(record))
    return lines


def _shared_override_note(notes: Sequence[AdjudicationNote]) -> str:
    """What every override on this run left standing, said once over all of them."""
    return (
        f"{_opening(len(notes), 'of the decisions below is', 'of the decisions below are')} "
        "an override: the adjudicator accepted the flaw rather than fixing, "
        "withdrawing or mitigating it, and the decision on the record is to proceed "
        "while carrying it. "
        + _joined_titles([note.title for note in notes], fallback="No idea")
        + " therefore remain live in this report and are ranked alongside the other "
        "ideas. A rank says nothing about a flaw: no reviewer withdrew any of these "
        "findings and no mitigation was recorded against them. Read each of those "
        "ideas' sections with the decision below in mind."
    )


def _shared_withdrawal_note(notes: Sequence[AdjudicationNote]) -> str:
    """The same, for the ideas an adjudicator removed rather than accepted."""
    return (
        f"{_opening(len(notes), 'of the decisions below is', 'of the decisions below are')} "
        "a withdrawal: the adjudicator removed the hypothesis from the population in "
        "answer to the flaw, so it never entered the tournament and carries no rank "
        "and no Elo. "
        + _joined_titles([note.title for note in notes], fallback="No idea")
        + " are therefore absent from the per-idea sections of this report as well. "
        "They were removed by decision rather than defeated on merit, and the "
        "population they were removed from is retained as superseded history."
    )


def _governance_block(record: ResearchRecord) -> list[str]:
    """Replay every human answer to a fatal governance finding, verbatim.

    Nothing is emitted on a run where nobody was asked to adjudicate, which is the
    common case: an empty heading would train a reader to skip the one section that
    exists to be unskippable.
    """
    lines: list[str] = []
    if record.adjudications:
        answered = len(record.adjudications)
        open_blocks = len(record.open_governance_blocks)
        lines.extend(
            [
                "## Governance adjudications",
                "",
                # Two false statements stood here. The count was len(adjudications)
                # printed as a number of hypotheses carrying a fatal flaw, and those
                # are different quantities: two blocks answered against one hypothesis
                # counted as two hypotheses, and a flaw nobody answered counted as
                # none. And "each block was answered by hand" was asserted
                # unconditionally, so a run holding an unanswered blocker told the
                # reader every block had been decided, two headings above the list of
                # the ones that had not.
                # "A person's answer" is more than the record supports. What the run
                # holds is a name typed at the command line, and the paragraph three
                # below says so; a live report opened this section by calling
                # "Automated verification run (Claude Code)" a person.
                f"{_opening(answered, 'governance adjudication')} "
                + (
                    "is recorded for this run: an answer entered"
                    if answered == 1
                    else "are recorded for this run, each of them an answer entered"
                )
                + " against a fatal flaw the safety and governance review raised. "
                + ("It is" if answered == 1 else "Each is")
                + " set out below with the flaw it responds to and the reason given, "
                "both reprinted word for word. The wording is theirs; nothing here is "
                "a summary, because a reader has to be able to judge the decision "
                "rather than learn that one was taken.",
                "",
                (
                    f"{_opening(open_blocks, 'further fatal finding')} "
                    + ("has" if open_blocks == 1 else "have")
                    + " been recorded and not answered by anyone. "
                    + ("It is" if open_blocks == 1 else "They are")
                    + " listed under Unanswered governance blocks below and "
                    + ("is" if open_blocks == 1 else "are")
                    + " not covered by any decision on this page."
                    if open_blocks
                    # Scoped to this page's subject. Written of fatal findings in
                    # general it contradicted the governance paragraph in the body,
                    # which on the same run reported fatal flaws against a further
                    # three ideas that nobody adjudicated: those come from the other
                    # reviews and are not governance blocks, and the two sentences
                    # were describing different sets in the same words.
                    else "No further safety and governance block was left "
                    "unanswered in this run. A fatal flaw recorded by one of the "
                    "other reviews is not a governance block and is not covered "
                    "here; those are printed under the ideas that carry them."
                ),
                "",
                # The name is a free-text argument to the adjudication command. The
                # run cannot tell whether it belongs to the person who typed it, and
                # a report that prints it as an attribution invites the reader to
                # treat it as one -- which turns an unverified string into the record
                # of who is accountable for overriding a safety finding.
                "The names below are as entered by whoever ran the adjudication. "
                "This system does not authenticate them, so a name here is a record "
                "of what was claimed at the time and carries only as much weight as "
                "the process that produced it.",
                "",
            ]
        )
        # What an override means, and what it leaves standing, is the same wherever it
        # is said. Three overrides on a live run printed the identical two sentences
        # under the resolution line and the identical four-sentence paragraph after
        # the quotations -- a hundred and twenty words repeated twice over, between
        # which the only thing that changed was the title.
        overridden = [note for note in record.adjudications if not note.withdrawn]
        pulled = [note for note in record.adjudications if note.withdrawn]
        hoisted: set[str] = set()
        if len(overridden) > 1:
            hoisted.add("override")
            lines.extend([_shared_override_note(overridden), ""])
        if len(pulled) > 1:
            hoisted.add("withdrawal")
            lines.extend([_shared_withdrawal_note(pulled), ""])
        # One justification answering several flaws. Reprinted under each of them it
        # reads as several considered answers, and on a live run the same forty words
        # stood under three different flaws and named containment controls that bore
        # on only one of them.
        shared_reasons = {
            text: count
            for text, count in Counter(
                " ".join(note.justification.split())
                for note in record.adjudications
                if note.justification.strip()
            ).items()
            if count > 1
        }
        for text, count in shared_reasons.items():
            given_by = _listed(
                sorted(
                    {
                        note.adjudicator
                        for note in record.adjudications
                        if " ".join(note.justification.split()) == text
                    }
                )
            )
            lines.extend(
                [
                    f"One justification below is given word for word against "
                    f"{_plural(count, 'flaw')}. It was written once by {given_by} and "
                    "applied to "
                    + ("both" if count == 2 else "all of them")
                    + ", so it is quoted here once and read as a single decision "
                    "covering them rather than as a separate answer to each:",
                    "",
                    f"> {text}",
                    "",
                ]
            )
    for index, note in enumerate(record.adjudications, start=1):
        outcome = "withdrawal" if note.withdrawn else "override"
        reason = " ".join(note.justification.split())
        lines.extend(
            [
                f"### {index}. {note.heading}",
                "",
                # No "Adjudicated by" line: the resolution sentence opens with the
                # adjudicator's name and the justification below is attributed to it
                # again, so the label was the third copy of one string inside a block
                # four lines deep.
                f"Resolution: {outcome}, adjudicated by {note.adjudicator}."
                if outcome in hoisted
                else f"Resolution: {outcome}. {note.resolution_sentence}",
                "",
                "Fatal flaw recorded by the safety and governance review, verbatim:",
                "",
                f"> {note.flaw_text}",
                "",
            ]
        )
        if reason in shared_reasons:
            lines.extend(
                [
                    f"Justification: the wording quoted above, which {note.adjudicator}"
                    " wrote once and applied to this flaw along with "
                    f"{_plural(shared_reasons[reason] - 1, 'other')}.",
                    "",
                ]
            )
        else:
            lines.extend(
                [
                    f"Justification given by {note.adjudicator}, verbatim:",
                    "",
                    f"> {reason}",
                    "",
                ]
            )
        if outcome in hoisted:
            continue
        if note.withdrawn:
            lines.extend(
                [
                    # That it never entered the tournament and carries no rank is in
                    # the resolution sentence a few lines above. What is left to say
                    # is the part that does not follow from it: the removal was a
                    # decision rather than a defeat, and the population survives it.
                    f"{note.title} is therefore absent from the per-idea sections of "
                    "this report as well. It was removed by decision rather than "
                    "defeated on merit, and the population it was removed from is "
                    "retained as superseded history.",
                    "",
                ]
            )
        else:
            lines.extend(
                [
                    f"{note.title} remains live in this report and still carries this "
                    "flaw. It is ranked alongside the other ideas, and its rank says "
                    "nothing about the flaw: no reviewer withdrew the finding and no "
                    "mitigation was recorded against it. Read its section with this "
                    "decision in mind.",
                    "",
                ]
            )
    if record.open_governance_blocks:
        # A block nobody answered is not an adjudication, so it gets its own heading
        # rather than a slot under one; filing it there would imply a decision exists.
        lines.extend(
            [
                ("###" if record.adjudications else "##")
                + " Unanswered governance blocks",
                "",
                "The following fatal findings have not been adjudicated by anyone. The "
                "work they cover is blocked rather than approved, and no part of this "
                "report constitutes permission to begin it.",
                "",
                *_bullets(
                    f"{item.title}: {item.flaw_text}"
                    for item in record.open_governance_blocks
                ),
                "",
            ]
        )
    return lines


_ORDINAL_WORDS = ("", "first", "second", "third", "fourth", "fifth", "sixth")


def _reference_lines(references: Sequence[Citation]) -> list[str]:
    """The reference list, with entries no title tells apart marked as distinct.

    Where the search captured nothing but a publisher, two different documents on
    the same publisher render as the same line: a live list carried "Untitled
    source on nih.gov" at 5 and again at 7, so a reader meeting [5] and [7] in the
    text had no way to tell whether one source had been numbered twice.
    """
    # An entry that carries its own link is already told apart by the link, so the
    # clause is only worth its length where nothing else on the line differs.
    # A standing true of every entry is the lead-in's to state, not each entry's:
    # twenty entries each ending "Not retrieved" is the sentence above the list
    # printed twenty times over.
    uniform = len({_entry_standing(citation) for citation in references}) == 1
    unlinked = [
        citation for citation in references if not names_a_document(citation.url)
    ]
    repeated = Counter(citation.title.strip().lower() for citation in unlinked)
    position: Counter[str] = Counter()
    lines = []
    for citation in references:
        key = citation.title.strip().lower()
        distinguisher = ""
        if not names_a_document(citation.url) and repeated[key] > 1:
            position[key] += 1
            distinguisher = (
                f", the {_ORDINAL_WORDS[position[key]]} of "
                f"{_number_word(repeated[key]).lower()} separate records the search "
                "returned "
                # The clause was written for entries the search left untitled and
                # was then printed under entries that had one, telling the reader
                # there was no title directly after setting the title down.
                + (
                    "under that title"
                    if _titled(citation)
                    else "under that publisher without a title"
                )
            )
        lines.append(
            _reference_line(
                citation, distinguisher=distinguisher, mark_standing=not uniform
            )
        )
    return lines


_LINK_TEXT_CEILING = 96


def _link_text(url: str) -> str:
    """What a link shows, which is not always the whole of where it goes.

    One ResearchGate locator carried the paper's title inside its path and ran to
    three hundred characters; printed as its own link text it took four lines of
    the reference list and wrapped mid-word in the PDF. The href stays whole -- it
    has to still resolve -- and the ellipsis marks that the text is not all of it.
    """
    if len(url) <= _LINK_TEXT_CEILING:
        return url
    # On a separator the URL itself supplies. The bare slice cut inside the words a
    # locator carries in its path -- ".../Effect_of_Al2O3_coating_on_the_electroch…"
    # -- which reads as a broken link rather than an abbreviated one.
    head = url[: _LINK_TEXT_CEILING - 1]
    cut = max(head.rfind(character) for character in "/_-.?&=")
    if cut > _LINK_TEXT_CEILING // 2:
        head = head[:cut]
    return head.rstrip("/_-.?&=") + "…"


def _reference_line(
    citation: Citation, *, distinguisher: str = "", mark_standing: bool = True
) -> str:
    """One reference, numbered to match its marker and linked where it can be.

    This list was flat, unnumbered and unlinked because the only locators
    discovery produced were grounding redirects, which name no document. Now
    that those are followed to the paper they open, a marker like [2] can be
    matched to an entry and the entry to a DOI. A lead that still has nothing
    better than a hostname says so, rather than offering a link to a homepage.
    """
    # A title that is itself a question already ends in punctuation, and
    # "... cathodes?." reads as a rendering fault.
    title = citation.title.rstrip() + distinguisher
    stop = "" if title.endswith((".", "?", "!")) else "."
    standing = _entry_standing(citation) if mark_standing else ""
    # Where the search returned no title the name is read off the locator's own path,
    # which is a good deal better than "Untitled source on researchgate.net" and is
    # still not what the document calls itself: the path may have been shortened, and
    # its hyphens do not say which of them joined a compound. The entry says so, once,
    # and only on the entries it is true of.
    named_by_address = (
        " Named from its address, the search having returned no title for it."
        if citation.named_by_address
        else ""
    )
    if names_a_document(citation.url):
        return (
            f"{citation.number}. {title}{stop} [{_link_text(citation.url)}]"
            f"({citation.url}){named_by_address}{standing}"
        )
    # "No resolvable locator was recorded" is the renderer describing its own state,
    # and it is not even accurate: what discovery returned for these is usually a
    # publisher's front page or a search redirect, which is a locator that does not
    # reach the document. Saying which of those it was tells a reader how to find the
    # paper themselves; saying "no locator" tells them not to try.
    site = _site_of(citation.url)
    if site:
        return (
            # "discovery recorded no link" named the stage rather than the act; the
            # reader has no stage called discovery, they have a literature search.
            f"{citation.number}. {title}{stop} Retrieved from {site}; the literature "
            f"search recorded no link to the document itself.{standing}"
        )
    if _redirect_only(citation):
        # "It has to be found by title" over an entry that has no title is advice the
        # entry itself refutes. These are the search's own redirects, which name no
        # document and no host: the publisher in the title is all the run captured.
        # Why that is so is stated once above the list, because five entries carrying
        # the same twenty-word explanation is the explanation printed four times too
        # often.
        return f"{citation.number}. {title}{stop}{standing}"
    if GROUNDING_REDIRECT_MARKER in citation.url:
        # A link was recorded. It is the search's own redirector, which expires and
        # names no publisher, so it cannot be printed and cannot be followed -- but
        # "no link was recorded" said the run had captured nothing when what it had
        # captured was a locator that stopped working.
        return (
            f"{citation.number}. {title}{stop} The literature search recorded only "
            "its own redirect link for this source, which no longer resolves, so it "
            f"has to be found by title.{standing}"
        )
    return (
        f"{citation.number}. {title}{stop} No link to this source was recorded, so it "
        f"has to be found by title.{standing}"
    )


# What the run established about the document behind an entry, said on the entry.
# Two lead-in sentences promised "which is which is recorded against each entry in
# the evidence appendix"; no entry recorded it, and the appendix that name points at
# lists ideas whose grounding is in doubt and no entry of this list at all.
_ENTRY_STANDING = {
    "verified": "",
    "corrected": " Retrieved and checked; the document carries a correction.",
    "metadata_verified": (
        " Not checked against the document: only its catalogue record was reached."
    ),
    "retracted": " Retrieved and found retracted. Nothing here is grounded by it.",
    "inaccessible": (
        " Could not be retrieved when this run went back to it. Nothing here is "
        "grounded by it."
    ),
}
_UNCHECKED_STANDING = (
    " Not retrieved: this entry records where a statement came from, not that the "
    "document says it."
)


def _entry_standing(citation: Citation) -> str:
    """The retrieval verdict this run recorded for the document, or nothing.

    A verified entry says nothing, because a mark printed against every entry is not
    a mark. What the reader needs is which entries fall short of the sentence over
    the list, and those are the ones that speak.
    """
    return _ENTRY_STANDING.get(citation.verification_status, _UNCHECKED_STANDING)


def _titled(citation: Citation) -> bool:
    """Whether the entry prints a title, as opposed to standing in for one."""
    return not citation.title.lstrip().startswith("Untitled source")


def _redirect_only(citation: Citation) -> bool:
    """A reference the search left as a publisher name and no way to reach it."""
    return (
        not names_a_document(citation.url)
        and not _site_of(citation.url)
        and not _titled(citation)
    )


def _site_of(url: str) -> str:
    """The host a non-citable locator points at, if it points at one."""
    if not url.startswith(("http://", "https://")) or GROUNDING_REDIRECT_MARKER in url:
        return ""
    _, _, remainder = url.partition("://")
    host, _, _ = remainder.partition("/")
    return host.removeprefix("www.")


# A qualifier that holds of every source is dropped from the citations themselves,
# because a tag on all of them distinguishes none of them. It is stated here instead,
# once, above the list it is true of -- which is what the suppression assumed and no
# code did, so a report citing four unverified sources said so nowhere near them.
_REFERENCE_QUALIFIERS = {
    "unsupported": (
        # "They are what the literature search returned" stood above four entries on a
        # run whose appendix records fifty-three source leads. What the list is, is the
        # sources this report cites; how many the search turned up is the appendix's to
        # state, and it does.
        "None of the sources below was checked against the document it names. They "
        "are the sources this report cites; what they are cited for is what the "
        "search said they hold, not what reading them established."
    ),
    "leaning accurate": (
        "Every source below was checked against the document it names, and every "
        "check came back short of full confidence. Each is corroborating rather than "
        "settled."
    ),
    "inaccurate": (
        # "or it could not be retrieved" used to be here, and being unable to open a
        # page is not a finding about the paper behind it. What the qualifier now
        # marks is retraction and only retraction.
        "Every source below names a document that has since been retracted. Nothing "
        "in this report is grounded by them."
    ),
    "disputed": (
        "Every source below was found to be disputed by the check against the "
        "document it names."
    ),
}


def _uniform_reference_standing(record: ResearchRecord) -> str:
    """The blanket sentence, where the run's own retrieval record bears it out.

    The qualifier is read off what support the *claims* carry, and every sentence
    above says something about whether the *documents* were retrieved and read --
    two records that can disagree. On a live run they did: twenty-four entries of
    which three had been retrieved and checked (1, 12 and 20, which print no
    standing of their own for exactly that reason) stood under "None of the sources
    below was checked against the document it names." A source with no claim
    annotation against it reads as unsupported, and verification had checked it.

    Where they disagree the count wins, because ``_cited_reference_standing`` is
    written from the retrieval record and says how many fall into each case rather
    than asserting one case of them all.
    """
    qualifier = record.citations.universal_qualifier
    sentence = _REFERENCE_QUALIFIERS.get(qualifier, "")
    if not sentence:
        return ""
    checked, total = record.citations.cited_standing
    # "None was checked" wants nothing checked; the other three open "Every source
    # below was checked" and want nothing unchecked.
    return (
        sentence
        if (checked == 0 if qualifier == "unsupported" else checked == total)
        else ""
    )


def _knowledge_base(record: ResearchRecord, overview: ResearchOverview) -> list[str]:
    references = record.citations.references()
    lines = [
        "# Knowledge Base",
        "",
        "## Knowledge Summary",
        "",
        overview.knowledge_summary,
        "",
        "## Open Questions",
        "",
        *(
            [overview.open_questions_lead_in, ""]
            if overview.open_questions_lead_in
            else []
        ),
        *_bullets(overview.open_questions),
        "",
        "## Unexpected Connections",
        "",
        *([overview.connections_lead_in, ""] if overview.connections_lead_in else []),
        *_bullets(overview.unexpected_connections),
        "",
        "## References",
        "",
    ]
    if references:
        # A qualifier that holds of the whole list gets the sentence written for it.
        # Where the list is mixed there is no universal qualifier, and this stood
        # empty -- so the reference list of a run that had checked some of its
        # sources and not others said nothing at all about which were which, while
        # the findings section two chapters above stated a figure.
        lead_in = _uniform_reference_standing(record) or _cited_reference_standing(
            record
        )
        if lead_in:
            lines.extend([lead_in, ""])
        if any(_redirect_only(citation) for citation in references):
            lines.extend(
                [
                    "Where an entry names a publisher and nothing else, the literature "
                    "search returned a redirect rather than the document, so it "
                    "captured neither the document's title nor a link that reaches it.",
                    "",
                ]
            )
        lines.extend(_reference_lines(references))
    else:
        lines.append(
            "No source was cited in this report, because discovery resolved no "
            "external literature for this goal."
        )
    lines.append("")
    return lines


def _grid(columns: Sequence[str], rows: Iterable[Sequence[str]]) -> list[str]:
    """A section's grid: any number of columns, each row already stringified."""
    return [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
        *["| " + " | ".join(row) + " |" for row in rows],
        "",
    ]


def _idea_table(rows: Iterable[tuple[str, str]]) -> list[str]:
    """The per-idea comparison grid: two columns, one bolded label set, five rows.

    The values are the specialist's own prose, and a pipe or a newline inside one
    ends the cell it is in. Narrative rewrites an embedded Markdown table as prose
    before it reaches here; this is the backstop for a stray pipe in a sentence,
    which would otherwise split the row into columns the header cannot name.
    """
    return [
        "| Category | Description |",
        "| --- | --- |",
        *[
            f"| **{label}** | {' '.join(value.split()).replace('|', '—')} |"
            for label, value in rows
        ],
        "",
    ]


def _numbered_subsections(titles: Iterable[str], bodies: Iterable[str]) -> list[str]:
    lines: list[str] = []
    for index, (title, body) in enumerate(zip(titles, bodies, strict=True), start=1):
        lines.extend([f"##### {index}. {title}", "", body, ""])
    return lines


def _revised_form_block(brief: IdeaBrief) -> list[str]:
    """The rewrite the recommendation is actually for, set out where the idea is.

    The report used to recommend the evolved form of four ideas and print the evolved
    text of none of them: the closing section named the changes in the abstract and a
    reader wanting the hypothesis they were being asked to fund had nowhere to go.
    """
    if not brief.revised_form:
        return []
    # The heading asserts a recommendation, so it only carries one where the
    # meta-review made it. Evolution rewrites the whole shortlist.
    heading = (
        "### Revised Form Recommended"
        if brief.revised_is_recommended
        else "### Revised Form"
    )
    lines = [heading, "", brief.revised_lead_in, ""]
    for label, text in brief.revised_form:
        lines.extend([f"**{label}:** {text}", ""])
    return lines


def _review_tally(review) -> str:
    """What this review raised and answered, in place of a second copy of either.

    Every objection used to be printed twice inside one idea: once as a bullet list
    under the review that raised it, and again under Deep Verification a page later.
    On a live run that was forty-one objections and eighty-two printings, and a reader
    who has already read a list does not read it again -- they skip the section it is
    in, which is the one section of the idea written to be checked item by item. Deep
    Verification is the copy that survives, because it carries what this one could
    not: the raising review, whether that review answered anything, and a number to
    refer to the objection by. Where that copy is printed is said once in the preamble
    above the ideas rather than under each of the forty reviews that has objections.

    The responses were the same defect one section earlier and outlasted the fix:
    "Rebuttals offered:" under the review, and the identical sentence again under
    Addressed Objections in the same idea's summary. That copy is the one that
    survives, for the same reason as before -- it attributes each response to the
    review that wrote it, which is what recovers the subject of one written as "could
    be useful for verifying the exact thickness dependence, but fundamentally lacks
    novelty". What is left here is the count, which is the one thing the summary
    cannot show: that this review answered at all rather than stood on its objection.
    """
    parts = []
    if review.objections:
        parts.append(
            f"raised {_number_word(len(review.objections)).lower()} "
            + ("objection" if len(review.objections) == 1 else "objections")
        )
    if review.rebuttals:
        parts.append(
            f"recorded {_number_word(len(review.rebuttals)).lower()} "
            + ("response" if len(review.rebuttals) == 1 else "responses")
        )
    return f"This review {' and '.join(parts)}."


def shared_review_tally(briefs: Sequence[IdeaBrief]) -> tuple[list[str], bool]:
    """The count above, where every review of every idea recorded the same one.

    The count says of a review that it answered at all rather than stood on its
    objection, which is worth a sentence where the reviews differ. On a live run of
    eight ideas they did not: "This review raised one objection and recorded one
    response." stood between the findings and the score of all five reviews of all
    eight, forty times, saying nothing about the review under which it was printed
    that was not equally true of the other thirty-nine. Where it does not vary it is
    a property of the run, so it is stated once with the other properties of the run
    and each review closes on its findings and its score.
    """
    scored = [review for brief in briefs for review in brief.reviews]
    tallies = {_review_tally(review) for review in scored}
    # Anything that varies stays where it is: a review that recorded no response is
    # the one this sentence exists to distinguish, and a hoisted count would say of
    # it what is true only of the others.
    if len(scored) < 2 or len(tallies) != 1:
        return [], False
    # A placeholder raised nothing and answered nothing, and is printed as the
    # placeholder it is; counting it among the reviews below would put words in the
    # mouth of a reviewer that never wrote one.
    if any(
        review.stood_in or not (review.objections or review.rebuttals)
        for review in scored
    ):
        return [], False
    said = next(iter(tallies)).removeprefix("This review ").rstrip(".")
    return (
        [
            f"Each of the {_number_word(len(scored)).lower()} reviews below "
            f"{said}, and that is the same under every idea, so it is given here "
            "rather than under each of them. What a review raised is printed under "
            "Deep Verification in the idea's own section and what it answered under "
            "Addressed Objections; what is said here is only that none of them "
            "stood on its objection without answering.",
            "",
        ],
        True,
    )


def _review_block(
    brief: IdeaBrief,
    hoisted: frozenset[tuple[str, str, str]] = frozenset(),
    *,
    tally_hoisted: bool = False,
) -> list[str]:
    """The five-to-six review sections, four of which close on a matched score."""
    lines: list[str] = ["### Reviews", "", "#### Summary", ""]
    lines.extend(_numbered_subsections(brief.summary.keys(), brief.summary.values()))
    for section in REVIEW_SECTIONS:
        reviews = brief.reviews_in(section)
        lines.extend([f"#### {section}", ""])
        if not reviews:
            lines.extend(
                [
                    f"No {section.lower()} review was recorded against this idea, so "
                    "this dimension is unassessed rather than passed.",
                    "",
                ]
            )
            continue
        for review in reviews:
            # A placeholder's findings and objection say the same nothing under every
            # idea a reviewer skipped, and printed as prose they read as this idea's
            # review. What is printed instead is that nobody wrote one. The verdict and
            # score stay, because they went into the averages and the ranking and the
            # reader has to be able to reconcile the arithmetic above with them.
            if review.stood_in:
                lines.extend(
                    [
                        f"No {section.lower()} review of this idea was written. This "
                        "is the fixed placeholder the run puts in a reviewer's place "
                        "where it answered for some ideas and not others: the verdict "
                        "and score below are the placeholder's, and they entered the "
                        "averages, the spread and the ranking as though a reviewer had "
                        "set them down.",
                        "",
                    ]
                )
                lines.extend(
                    [f"Answer: {review.answer}", "", f"Score: {review.score}", ""]
                )
                continue
            # Who the reviewer is and what it asked are the same under every idea, so
            # where more than one idea carries the pair it is stated once above them
            # all and this section opens straight onto what the review found.
            if (section, review.lead_in.rstrip(":"), review.question) not in hoisted:
                lines.extend([review.lead_in, "", review.question, ""])
            for finding in review.findings:
                lines.extend([finding, ""])
            lines.extend(_review_finding_tables(review))
            if (review.objections or review.rebuttals) and not tally_hoisted:
                lines.extend([_review_tally(review), ""])
            # The reference reports close every scored review with a matched pair, so
            # a reader can check the verdict against the number without scrolling.
            lines.extend([f"Answer: {review.answer}", "", f"Score: {review.score}", ""])
    lines.extend(["#### Coherence", ""])
    for paragraph in brief.coherence:
        lines.extend([paragraph, ""])
    lines.extend(["#### Deep Verification", ""])
    if brief.deep_verification_lead_in:
        lines.extend([brief.deep_verification_lead_in, ""])
    for index, (title, body) in enumerate(brief.deep_verification, start=1):
        lines.extend([f"##### {index}. {title}", "", body, ""])
    return lines


def _verdict_line(match, *, transcript_above: bool = True) -> str:
    """The result of one match, without restating an argument already printed.

    The judge states its rationale in the closing turn, so printing it again
    underneath repeated a full paragraph verbatim on every debated match.

    Only where that turn is on the page. A pair that met twice is transcribed
    under one of the two ideas and cross-referenced under the other, and the
    cross-referenced chapter suppressed the reason against a transcript the
    reader has to go elsewhere to find: "This is the same exchange as the one
    under MLD Alucone Hybrid Coating ... The verdict below is how it went for
    this idea." stood above "The judge ruled this a loss with confidence 0.70."
    and nothing else.
    """
    verdict = (
        f"The judge ruled this a {match.outcome} with confidence "
        f"{match.confidence:.2f}."
    )
    # What the reader was shown, which is what a repeat would be repeating.
    # ``readable_turn`` is not that: the turns are printed contribution by
    # contribution through ``readable_exchange``, which gives the closing rationale
    # the sentence capital that its "Rationale:" label had kept off it. So the test
    # read "Rationale: this idea provides a more feasible ..." against a rationale
    # capitalised on its own, missed by the one letter, and printed the paragraph
    # twice -- four times over on one live report.
    tail = (
        " ".join(said for _prefix, said in readable_exchange(match.debate_turns[-1]))
        if match.debate_turns and transcript_above
        else ""
    )
    # A rematch note is recorded on the rationale so the pairing explains
    # itself; it belongs with the result, not inside the judge's reasoning.
    note, _, remainder = match.rationale.strip().partition("]")
    if note.startswith("["):
        note = note[1:].strip()
    else:
        note, remainder = "", match.rationale
    rationale = " ".join(
        unemphasised(strip_rationale_label(strip_turn_label(remainder))).split()
    )
    # Containment, not suffix: a judge that repeats its rationale mid-turn and
    # then closes on a further sentence still printed the paragraph twice.
    if rationale and rationale in " ".join(tail.split()):
        # The judge argues its verdict in the closing turn above, so the verdict
        # says nothing further and the preamble above the ideas says why. Pointing
        # at the paragraph directly overhead was twelve copies of one sentence on
        # a live run. Silence is unambiguous only because the case with nothing to
        # point at now says so.
        pass
    elif rationale:
        # Opened after the containment test above, which reads the judge's own
        # words: an opening dropped before the test stops matching the turn that
        # carries it, and the paragraph goes back to being printed twice.
        verdict = f"{verdict} Rationale: {standalone_opening(rationale)}"
    else:
        verdict = f"{verdict} No rationale was recorded for it."
    # The note used to sit between the verdict and the sentence that opens on "Its",
    # so the pronoun reached back across a bracketed aside about a different match.
    # It is an aside about the pairing, so it goes after what it is an aside from.
    return f"{verdict} {note}" if note else verdict


def _match_rationale(match) -> str:
    """One judge's stated reason, as prose rather than as the markup it arrived in.

    Read at render time as well as at parse time, because the sessions this has to
    print were recorded before the parser cleaned any of it: a live dossier carried
    "**Conclusion:** Hypothesis 2 ..." inside a bullet whose own label is bold, and
    the stray asterisks closed the report's emphasis instead of the judge's.
    """
    remainder = match.rationale
    note, _, tail = match.rationale.strip().partition("]")
    if note.startswith("["):
        remainder = tail
    return standalone_opening(
        " ".join(
            unemphasised(strip_rationale_label(strip_turn_label(remainder))).split()
        )
    )


def _undebated(brief: IdeaBrief) -> list:
    """This idea's matches that were decided without an exchange to read."""
    return [
        match
        for match in brief.matches
        if not match.debate_turns and not match.unreadable_turns
    ]


def shared_match_notes(briefs: Sequence[IdeaBrief]) -> tuple[list[str], frozenset[str]]:
    """What is true of every undecided-by-debate match, said once above the ideas.

    Who judged those matches is a property of the tournament, not of the idea whose
    table they appear under, and a judge that reuses one line of reasoning across the
    whole tournament has said something about the tournament rather than about any
    match in it. Both were printed under every idea: eight copies of the first and,
    on the adjudication run, seven of the second.
    """
    groups = [matches for matches in map(_undebated, briefs) if matches]
    if len(groups) < 2:
        return [], frozenset()
    judged = [
        frozenset(_judge_label(match.judge) for match in group) for group in groups
    ]
    # Two ideas whose matches were judged differently cannot share one sentence about
    # who judged them; joining the labels would attribute both judges to both ideas.
    if len(set(judged)) > 1:
        return [], frozenset()
    matches = [match for group in groups for match in group]
    reasons = {
        said.rstrip(".") for match in matches if (said := _match_rationale(match))
    }
    # Whether every one of them carries a reason decides both what this note may
    # claim and whether the ideas below may stop repeating it. Written flat, the note
    # asserted a stated reason for matches whose judge recorded none.
    complete = all(_match_rationale(match) for match in matches)
    hoisted = {"judges"} | ({"tail"} if complete else set())
    notes = [
        "Not every match was argued. Where an idea's table below shows a match with "
        f"no debate under it, that match was decided by {' and '.join(sorted(judged[0]))} "
        "rather than by argued rounds: there is a verdict"
        + (
            " and the reason the judge recorded"
            if complete
            else ", a stated reason where the judge recorded one"
        )
        + ", and no exchange behind either to read."
    ]
    if len(reasons) == 1 and complete:
        hoisted.add("reason")
        notes.append(
            "One reason is recorded for every one of those matches, word for word, so "
            "it accounts for the tournament rather than for any single result: "
            f"{next(iter(reasons))}."
        )
    return [line for note in notes for line in (note, "")], frozenset(hoisted)


def _match_summary(
    brief: IdeaBrief,
    hoisted: frozenset[str] = frozenset(),
    transcribed: set[tuple[int, frozenset[str]]] | None = None,
) -> list[str]:
    """The tournament block, with the debate transcripts the ranking was decided on.

    ``transcribed`` carries the matches already reproduced. A match has two sides and
    the report gives each idea a chapter, so every exchange was printed twice, once
    under each -- about six thousand words of a live report, and a reader who
    recognised the second copy had no way to be sure it was the same one. Passing it
    is what makes the second chapter point at the first instead; omitting it prints
    every transcript, which is what a single idea's block should do.
    """
    lines = [
        "### Tournament",
        "",
        "#### Match summary",
        "",
        f"- Total matches: {len(brief.matches)}",
        f"- Matches won: {brief.wins}",
        f"- Matches lost: {brief.losses}",
        f"- Matches tied: {brief.ties}",
        f"- Win rate: {brief.win_rate}%",
        "",
    ]
    if not brief.matches:
        lines.extend(
            [
                "This idea played no ranked match, so its position rests on its "
                "reviews alone.",
                "",
            ]
        )
        return lines
    lines.extend(
        [
            "| Round | Opponent | Result | Elo before | Elo after | Judge |",
            "| ---: | --- | --- | ---: | ---: | --- |",
        ]
    )
    for match in brief.matches:
        lines.append(
            f"| {match.round_number} | {match.opponent_title} | {match.outcome} | "
            f"{match.shown_before} | {match.shown_after} | "
            f"{_judge_column(match.judge)} |"
        )
    lines.append("")
    for match in brief.matches:
        if not match.debate_turns and not match.unreadable_turns:
            continue
        lines.extend([f"#### Debate against {match.opponent_title}", ""])
        # Keyed on the round as well as the pair, because two ideas can meet twice:
        # once in the Swiss rounds and again in the top round robin.
        seen = (match.round_number, frozenset({brief.title, match.opponent_title}))
        if transcribed is not None and seen in transcribed:
            lines.extend(
                [
                    "This is the same exchange as the one under "
                    f"{match.opponent_title}, read from the other side, and it is "
                    "reproduced there rather than in both chapters. The verdict "
                    "below is how it went for this idea.",
                    "",
                    _verdict_line(match, transcript_above=False),
                    "",
                ]
            )
            continue
        if transcribed is not None:
            transcribed.add(seen)
        if match.debate_turns:
            # One bullet per contribution rather than per recorded turn: a turn is
            # often a whole exchange, and printing it whole put three experts and a
            # closing rationale into one unbroken paragraph.
            lines.extend(
                [
                    *[
                        f"- **{prefix}:** {said}" if prefix else f"- {said}"
                        for turn in match.debate_turns
                        for prefix, said in readable_exchange(turn)
                    ],
                    "",
                ]
            )
        if match.unreadable_turns:
            lines.extend(
                [
                    f"{_number_word(match.unreadable_turns)} "
                    + ("turn" if match.unreadable_turns == 1 else "turns")
                    + " of this match's debate "
                    + ("was" if match.unreadable_turns == 1 else "were")
                    # "a serialised payload" is the shape the data arrived in, which
                    # is the renderer's problem and not the reader's; what the reader
                    # needs is that there is no readable argument here.
                    + " recorded as raw data rather than as prose and "
                    "held back. The argument behind this result cannot be read in "
                    "full, so the outcome should be treated as undebated.",
                    "",
                ]
            )
        lines.extend([_verdict_line(match), ""])
    undebated = _undebated(brief)
    if undebated:
        # Two things were wrong with the sentence this replaces. "3 of these
        # matches" printed under "Total matches: 3" invites the reader to look
        # for the ones that do carry a transcript, and "no reasoning underlies
        # those results" contradicted the Judge column beside it, which named an
        # LLM comparison. A single-pass judgement is not an arithmetic one; what
        # it lacks is a transcript, and that is what the reader cannot audit.
        count = (
            "None of these matches carries a debate transcript"
            if len(undebated) == len(brief.matches)
            else f"{_number_word(len(undebated))} of these matches "
            + ("carries" if len(undebated) == 1 else "carry")
            + " no debate transcript"
        )
        judges = sorted({_judge_label(match.judge) for match in undebated})
        # "the reasoning behind it is not [recorded]" was false: every one of these
        # judgements carries a one-line rationale in the session, and the report was
        # telling the reader that reasoning it was withholding did not exist. What
        # these matches actually lack is a transcript -- the argument that produced
        # the line -- so the rationale is printed and the shortfall stated exactly.
        reasoned = [match for match in undebated if _match_rationale(match)]
        # A judge that writes one line and reuses it is not giving a reason per match.
        # Printed as a bullet per round it read as four findings; over seven ideas it
        # was twenty-four copies of one sentence. Identical text is folded to one
        # statement, and that it is identical is itself the thing worth reporting.
        distinct = {_match_rationale(match).rstrip(".") for match in reasoned}
        folded = len(distinct) == 1 and len(reasoned) > 1
        # What a missing transcript costs the reader is set out once, under Comparison
        # of Candidate Ideas. Eight ideas each carrying their own copy of that
        # explanation is eight paragraphs of the same two sentences; what belongs here
        # is only what is true of this idea's matches.
        tail = (
            ""
            # "tail" says every undebated match in the run recorded a reason and none
            # of them has an exchange, which is what this sentence said under each of
            # the eight ideas; the note above the ideas now says it once.
            if "reason" in hoisted or "tail" in hoisted
            else "there is nothing to reproduce below: no reason was recorded for "
            "any of them either."
            if not reasoned
            else "the reason the judge recorded is below and the exchange behind it "
            "is not."
            if folded
            else "the judge's stated reason for each is below and the exchange behind "
            "it is not."
            if len(reasoned) == len(undebated)
            # Saying "for each" over a list that is short of the count above leaves
            # the reader hunting for bullets that were never written.
            else "the one that recorded a reason is below, without the exchange "
            "behind it, and the rest recorded no reason either."
            if len(reasoned) == 1
            else f"the {_number_word(len(reasoned))} of them that recorded a reason "
            "are below, without the exchange behind them, and the rest recorded no "
            "reason either."
        )
        # Who judged them is the same sentence under every idea when one judge decided
        # the whole tournament, so it is hoisted and this opens on the count alone.
        decided = (
            ""
            if "judges" in hoisted
            else f" They were decided by {' and '.join(judges)} rather than by "
            "argued rounds,"
        )
        stated = f"{count}.{decided}"
        if tail:
            stated += f" so {tail}" if decided else f" {tail[0].upper()}{tail[1:]}"
        elif decided:
            stated = stated.rstrip(",") + "."
        lines.extend([stated, ""])
        if "reason" in hoisted:
            pass
        elif folded:
            lines.extend(
                [
                    "The same reason is recorded for "
                    + (
                        "every one of them"
                        if len(reasoned) == len(undebated)
                        else f"{_number_word(len(reasoned))} of them"
                    )
                    + ", word for word, so it accounts for the set rather than for any "
                    f"one result: {distinct.pop()}.",
                    "",
                ]
            )
        else:
            for match in reasoned:
                lines.append(
                    f"- **Round {match.round_number} against {match.opponent_title}"
                    f" ({match.outcome}):** {_match_rationale(match).rstrip('.')}."
                )
            if reasoned:
                lines.append("")
    return lines


def _idea_deep_dive(
    record: ResearchRecord,
    brief: IdeaBrief,
    *,
    grounding_hoisted: bool = False,
    authors_own_hoisted: bool = False,
    hoisted_questions: frozenset[tuple[str, str, str]] = frozenset(),
    hoisted_matches: frozenset[str] = frozenset(),
    tally_hoisted: bool = False,
    transcribed: set[tuple[int, frozenset[str]]] | None = None,
) -> list[str]:
    lines = [
        f"## {brief.title}",
        "",
        brief.rank_line,
        "",
        f"Elo: {brief.elo}",
        "",
        f"Category: {brief.category}",
        "",
    ]
    # An idea that is still live because a person accepted its fatal flaw cannot be
    # read past this point without meeting the flaw, so it precedes the grounding note.
    if brief.chapter_governance_notice:
        lines.extend([brief.chapter_governance_notice, ""])
    # The verdict is a field of this idea and stays here. What it means is the same
    # sentence for every idea that shares it, and where they all do it is stated above
    # the ideas instead of eight times below.
    lines.extend(
        [
            # Hoisting left "Evidence support: unverified." standing alone, a thousand
            # lines below the paragraph that says what the word means -- and this is
            # the page a reader is on when they decide whether to act on the idea.
            # The Executive Candidate Summary says where its evidence column is
            # explained; the idea's own heading said nothing.
            f"Evidence support: {brief.support_label} — the verdict explained under "
            "Candidate Ideas above."
            if grounding_hoisted
            else brief.support_notice,
            "",
        ]
    )
    lines.extend(["### Idea Proposal", "", brief.proposal, "", "### Description", ""])
    for paragraph in brief.description:
        lines.extend([paragraph, ""])
    # The protocol and the diagram of it are one thing to a reader who intends to run
    # the idea, so they are set together and ahead of the specialist arguing for it.
    lines.extend(_validation_protocol(brief))
    lines.extend(_workflow_diagram(brief))
    lines.extend(_authors_own_sections(brief, hoisted=authors_own_hoisted))
    lines.extend(_self_rating(brief))
    lines.extend(_evidence_assessment(brief))
    lines.extend(_revised_form_block(brief))
    lines.extend(_review_block(brief, hoisted_questions, tally_hoisted=tally_hoisted))
    lines.extend(_match_summary(brief, hoisted_matches, transcribed))
    return lines


def _validation_protocol(brief: IdeaBrief) -> list[str]:
    """The experiment the specialist designed, step by step as it numbered them.

    ``validation_protocol`` is required by the contract, asked for by every generation
    prompt and checked by normalisation, and no exporter read it: the sample size and
    its power rationale, the calibration, the blinding, the abort limits and the
    go/no-go threshold were in the saved session for all eight live ideas and in none
    of the three exports. Without it the report proposes hypotheses and says nothing
    about how to test them.
    """
    if not brief.validation_protocol:
        return []
    lines = ["### Validation Protocol", ""]
    if len(brief.validation_protocol) == 1:
        lines.extend([brief.validation_protocol[0], ""])
        return lines
    lines.extend(
        f"{index}. {step}"
        for index, step in enumerate(brief.validation_protocol, start=1)
    )
    lines.append("")
    return lines


def shared_authors_own_note(briefs: Sequence[IdeaBrief]) -> list[str]:
    """What The Specialist's Own Sections is, said once above the ideas that have one.

    The note under the heading describes the generation contract -- one prose field
    for a mechanism the prompt asks four things of -- rather than the idea it stands
    under, and it stood in the same words under four of eight ideas.
    """
    if sum(1 for brief in briefs if brief.authors_own_sections) < 2:
        return []
    return [
        "Some ideas below carry a section headed The Specialist's Own Sections. The "
        "contract gives a specialist one prose field for the mechanism, and where a "
        "specialist headed parts of its answer inside that field, those parts are "
        "reproduced there under its own labels rather than read out as the mechanism. "
        "That section is the proposer arguing for its own idea, not a finding of the "
        "run.",
        "",
    ]


def _authors_own_sections(brief: IdeaBrief, *, hoisted: bool = False) -> list[str]:
    """The sections the proposing specialist headed inside its own mechanism field.

    The generation prompt asks for four parts and the schema gives them one prose
    field to arrive in, so the report read the lot out as the mechanism: one live
    Mechanism cell opened "Motivation and Supporting Evidence:", ran 1,475 characters
    through a "Critical Scientific Judgment:" and never said what the mechanism was.
    Kept whole and printed under the labels the specialist gave them, in its voice.
    """
    if not brief.authors_own_sections:
        return []
    lines = ["### The Specialist's Own Sections", ""]
    if not hoisted:
        lines.extend(
            [
                "The specialist that proposed this idea headed "
                + (
                    "a section of its own"
                    if len(brief.authors_own_sections) == 1
                    else f"{_number_word(len(brief.authors_own_sections)).lower()} "
                    "sections of its own"
                )
                + " inside the mechanism above. "
                + (
                    "It is reproduced"
                    if len(brief.authors_own_sections) == 1
                    else "They are reproduced"
                )
                + " here under the specialist's own label: this is the proposer "
                "arguing for its own idea, not a finding of the run.",
                "",
            ]
        )
    for label, body in brief.authors_own_sections:
        lines.extend([f"**{label}.** {_sentence(body)}", ""])
    return lines


def _self_rating(brief: IdeaBrief) -> list[str]:
    """The table the proposing specialist appended to one of its own prose fields.

    It arrived inside a prose field, so it was read out as clauses of the mechanism
    and printed into the Mechanism cell of the comparison grid, where a judgement
    the specialist awarded itself stood among the fields the run filled in. Kept,
    because the specialist wrote it; moved, and said to be its own.

    Which field it was appended to is the specialist's choice, so the sentence names
    it rather than assuming the mechanism: two of eight live ideas appended it to the
    protocol, where it ran on after the last bench step.

    The heading is set beside the table and not at the end of the sentence, where the
    nearest noun took it: "appended a table of its own to the validation protocol,
    headed Evaluation of Idea Table" gives the heading to the protocol, which is a
    field of the idea and has no heading of its own to give.
    """
    if not brief.self_rating:
        return []
    header, *rows = brief.self_rating
    width = len(header)
    appended = brief.self_rating_source or "mechanism"
    return [
        "### The Specialist's Own Rating",
        "",
        "The specialist that proposed this idea appended a table of its own"
        + (
            f" — headed {brief.self_rating_title} — "
            if brief.self_rating_title
            else " "
        )
        + f"to the {appended}. "
        + "The ratings in it are its own assessment of what it had just written, "
        "not a result of the reviews or the tournament below.",
        "",
        *_grid(
            [_cell(name) for name in header],
            [
                [
                    _cell(row[index] if index < len(row) else "")
                    for index in range(width)
                ]
                for row in rows
            ],
        ),
    ]


def _review_finding_tables(review) -> list[str]:
    """The tables a reviewer wrote inside its findings, kept as tables.

    The discipline critics answer the findings field with a Markdown table under
    "**Structured Evaluation Table:**", and the flattener read each row out as a run
    of clauses -- "Aggregation Control (Description: ALD on pre-fabricated electrodes
    prevents agglomeration; Judgment: High)". Fourteen findings on a live run arrived
    that way, five rows apiece, in the one section of the report a reader consults to
    find out what was wrong with an idea.
    """
    lines: list[str] = []
    for title, table in review.finding_tables:
        header, *rows = table
        width = len(header)
        if title:
            lines.extend([f"**{title}.**", ""])
        lines.extend(
            _grid(
                [_cell(name) for name in header],
                [
                    [
                        _cell(row[index] if index < len(row) else "")
                        for index in range(width)
                    ]
                    for row in rows
                ],
            )
        )
    return lines


def _cell(text: str) -> str:
    """One table cell, with anything that would end the cell early taken out."""
    return " ".join(str(text).split()).replace("|", "—")


def _workflow_diagram(brief: IdeaBrief) -> list[str]:
    """The specialist's own diagram of the idea, where it drew one.

    ``workflow_diagram_mermaid`` was on the contract, asked for in the prompt and
    returned by the generators, and no exporter read it: every diagram the run
    produced was carried in the saved session and printed nowhere.
    """
    if not brief.mermaid:
        return []
    return [
        "### Proposed Workflow",
        "",
        "```mermaid",
        brief.mermaid,
        "```",
        "",
    ]


def _evidence_assessment(brief: IdeaBrief) -> list[str]:
    """What the proposing specialist thought the literature did to its own idea."""
    if not brief.evidence_notes:
        return []
    lines = ["### Evidence Assessment", ""]
    current = ""
    for heading, badge, statement in brief.evidence_notes:
        if heading != current:
            if current:
                lines.append("")
            current = heading
            lines.extend([f"**{heading}:**", ""])
        lines.append(f"- {f'**{badge}** ' if badge else ''}{statement}")
    lines.append("")
    return lines


# The producer field holds whatever ran the stage, and not all of them are models. The
# offline provider's key is the model id it reports, "deterministic-offline": the key
# here used to be "deterministic", which nothing writes, so the appendix printed the
# raw sentinel in the column headed by a model name and a reader auditing the report
# had no way to tell a template apart from a model they had not heard of.
#
# Each qualifier is bracketed rather than set off by a comma or a colon, because these
# labels are also listed against each other in the run facts. "a fixed template, not a
# model, deep-research-preview-04-2026, and gemini-3.1-pro-preview" is a list of three
# producers written as though it were four, and the join cannot repair a comma that is
# inside the item.
_PRODUCER_LABELS = {
    "google_search_grounding": "Google Search grounding (model not recorded)",
    "unavailable": "nothing (the stage produced no output)",
    "deterministic-offline": "a fixed template (not a model)",
    MERGE_PRODUCER: "a merge of the specialists' answers (no model call)",
}
# What a stage handed back, named for what it is rather than for the class it
# validates against. "DossierManifest" and "EvolutionCycle" are this codebase's words
# for them, and the appendix is read by someone auditing the report, not the source.
_RECORD_TYPES = {
    "ResearchPlan": "research plan",
    "DiscoveryManifest": "literature discovery manifest",
    "EvidencePacket": "evidence packet",
    "CandidatePopulation": "hypothesis population",
    "ReviewSet": "set of reviews",
    "TournamentState": "tournament record",
    "EvolutionCycle": "revision cycle",
    "ResearchLandscape": "idea landscape",
    "DossierManifest": "meta-review of the round",
}


def _agent_name(agent: str) -> str:
    """A specialist as the rest of the report names it, not as the run files it.

    The appendix rendered agent ids by taking the underscores out, which is the one
    place in this codebase that did: everywhere else the reader meets "evidence and
    correctness review" and "clustering by mechanism", the table said "reflection" and
    "proximity". A reader matching a warning to the row it is about had to know that
    those are the same stage.
    """
    return _AGENT_NAMES.get(agent, agent.replace("_", " "))


def _specialist_label(record: ResearchRecord, note) -> str:
    """The specialist a row credits, minus the credit where nothing was produced.

    The evidence row named the Deep Research discovery specialist on a run whose
    Literature discovery section, two paragraphs above, states that the Deep Research
    agent never ran and that a single grounded search pass stood in for it. The row is
    still where the manifest is filed, so the specialist is still the right label for
    it -- but printed alone it credits work nobody did, and the "Produced by" cell
    beside it ("Google Search grounding") is not a contradiction a reader should have
    to notice unaided.
    """
    name = _agent_name(note.agent)
    if note.agent.startswith("deep_research") and record.deep_research_stood_in:
        return f"{name} — did not run; see Literature discovery above"
    return name


# Words that carry no meaning of their own in a two- or three-word name, so that
# "scoping the goal" and "goal scoping" can be recognised as the one name they are.
_STAGE_FILLER = frozenset({"a", "and", "by", "of", "the"})


def _one_name(text: str) -> frozenset[str]:
    """A name reduced to the words that distinguish it, in no particular order."""
    return frozenset(word for word in text.lower().split() if word not in _STAGE_FILLER)


def _specialist_cell(stage: str, label: str, alone: bool) -> str:
    """The specialist, or what its name would say twice if the row printed it.

    Five of the nine stages run one specialist that carries the stage's own name, so
    naming the stage in words rather than by its id puts the same phrase in two
    adjacent cells: "| tournament ranking | tournament ranking |", and "| scoping the
    goal | goal scoping |" for the same repetition with the words swapped round. What
    the second cell is there to say on those rows is that nothing fanned out.
    """
    if alone and _one_name(stage) == _one_name(label):
        return "the stage's only specialist"
    return label


def _elapsed(start: str, end: str) -> str:
    """How long the run took, which is what the two timestamps are printed for."""
    try:
        span = datetime.fromisoformat(end) - datetime.fromisoformat(start)
    except ValueError:
        return ""
    minutes = round(span.total_seconds() / 60)
    if span.total_seconds() < 60:
        return f", {round(span.total_seconds())} seconds after it started"
    return f", {_plural(minutes, 'minute')} after it started"


def _producer_label(model: str) -> str:
    return _PRODUCER_LABELS.get(model, model)


def _produced_by(note: ProvenanceNote) -> str:
    """What wrote one stage's payload, distinguishing a merge from a substitution.

    The generation stage runs four specialists and then folds their answers into one
    population, and that fold calls no model. Labelled like any other template it
    reached the table as "a fixed template (not a model)" -- the phrase this report
    uses for a stage whose specialist failed -- directly above a sentence saying no
    stage fell back to a template. One of the two had to be wrong, and it was the
    label: nothing failed, and no answer in that row is a template's.

    Runs from here on record ``MERGE_PRODUCER`` on that artifact and need none of
    this. Sessions saved before then are still re-rendered on every download, so the
    one agent that wrote the misleading default is corrected by name -- and only when
    the note is not a fallback, since a wholly offline run's aggregate really is a
    template like everything else in it.
    """
    if not note.model:
        return "not recorded"
    if (
        note.agent == "generation_aggregator"
        and note.model == "deterministic-offline"
        and note.source != "deterministic_fallback"
    ):
        return _PRODUCER_LABELS[MERGE_PRODUCER]
    return _producer_label(note.model)


def _record_type(schema_name: str) -> str:
    return _RECORD_TYPES.get(schema_name, schema_name)


# What the "Written by" column asks for, against the enum the payload is filed
# under. Printed raw, four rows of a live table answered "Written by" with
# "repaired", which names neither an author nor anything the page has defined by
# the time the reader meets it.
_WRITTEN_BY = {
    "specialist": "the specialist",
    "repaired": "the specialist, then repaired",
    "deterministic_fallback": "a fixed template",
}


def _run_facts(record: ResearchRecord) -> list[str]:
    """What produced this document, in the terms someone would need to reproduce it.

    The appendix opened straight into evidence integrity, which says whether the run
    was sound but nothing about which run it was. A reader holding two dossiers from
    the same goal had no field to tell them apart, and nobody auditing one could say
    which session, which models or which tournament settings it came out of.
    """
    session = record.session
    # The last thing the run recorded, rather than the session field, which until
    # recently nothing wrote after the session was created: every dossier printed a
    # "Last updated" equal to its start time to the microsecond, and sessions saved
    # by older builds still carry it.
    finished = max(
        [event.created_at for event in session.events]
        + [artifact.created_at for artifact in session.artifacts]
        + [session.updated_at],
        default=session.updated_at,
    )
    facts = [
        # The bare id told a reader nothing about what to do with it. It is the handle
        # every other record of this run is filed under, which is the reason to print
        # it at all.
        f"Session: {session.id} (the id this run's saved session, ledger rows and "
        "stored artifacts are filed under)",
        f"Started: {session.created_at[:19].replace('T', ' ')} UTC",
        f"Last recorded activity: {finished[:19].replace('T', ' ')} UTC"
        + _elapsed(session.created_at, finished),
        f"Workflow version: {session.workflow_version}",
        f"Stages completed: {session.current_stage} of 8",
        _approval_fact(session),
    ]
    # Provenance is the section a reader goes to for what this run did, and on a
    # fork two of the lines around this one are about work it did not do: the
    # stage count includes an evidence stage it started past, and the models it
    # names include the Deep Research model it never called.
    if session.seeded_evidence_from:
        facts.append(
            f"Evidence forked from: {session.seeded_evidence_from} — this run did "
            "not search the literature. Its scope and evidence base were carried "
            "over from an earlier run of the same question, and the stage count "
            # Both of them, because both are carried: the bullet named only the
            # evidence stage, so "8 of 8" less the one stage disclaimed read as
            # seven stages of this run's own work where it was six.
            "above includes the scope and evidence stages it started past."
        )
    # "Models: gemini-3.1-pro-preview, google_search_grounding" named a tool as a
    # model, in a field a reader uses to reproduce the run.
    #
    # Sorted case-insensitively, because a case-sensitive sort is not alphabetical to
    # a reader: it put "Google Search grounding" ahead of "a fixed template" and
    # "deep-research-preview" for no reason the page could show.
    produced: dict[bool, set[str]] = {True: set(), False: set()}
    for note in record.provenance:
        if note.model:
            produced[note.stage in FORKED_STAGES].add(_produced_by(note))
    producers = sorted(produced[True] | produced[False], key=str.lower)
    fact = f"Produced by: {_listed(producers, fallback='not recorded')}"
    # The forked scope and evidence carry their own provenance, so the models that
    # wrote them land in this list beside the models this run called. A live fork
    # named the Deep Research model here, unqualified, one bullet under the line
    # saying the run did not search the literature -- in the field an auditor uses
    # to reproduce it.
    #
    # The set is the models that ran in a forked stage and nowhere else, and the
    # qualifier has to say that rather than claim the forked work. Written as
    # "produced the forked scope and evidence", it credited the whole of both to
    # whatever survived the difference: a live fork read "of which
    # deep-research-preview-04-2026 produced the forked scope and evidence", while
    # Table 27 four bullets below gave the scope to gemini-3.1-pro-preview -- dropped
    # from the qualifier only because it also ran stages this run executed.
    inherited = sorted(produced[True] - produced[False], key=str.lower)
    if session.seeded_evidence_from and inherited:
        fact += (
            f" — of which {_listed(inherited)} ran only in the scope and evidence "
            "this run forked, and in nothing it executed itself"
        )
    facts.append(fact)
    prompts = sorted(
        {note.prompt_version for note in record.provenance if note.prompt_version}
    )
    if prompts:
        facts.append(f"Prompt version: {_listed(prompts)}")
    if record.tournament:
        facts.extend(_tournament_facts(record.tournament))
        judges = sorted(
            {_judge_label(item.judge) for item in record.tournament.comparisons}
        )
        if judges:
            # ", ".join gave "a multi-turn model debate, a single-pass model
            # comparison", which is two noun phrases with no conjunction between them
            # -- a list the reader has to reread to see is a list. Everywhere else in
            # this module a series of judges is written out as a series.
            facts.append(f"Judged by: {_listed(judges)}")
    return _bullets(facts)


def _approval_fact(session: Session) -> str:
    """What "approved" meant on this run, restated where the audit trail is read.

    The approval regime is declared in the front matter and warned about in section 1,
    and the appendix then prints a table of what every stage produced under a sentence
    about stages being accepted -- with nothing on the page to say who or what did the
    accepting. A reader who opens the report at the appendix, which is what auditing it
    looks like, could not tell whether a person had ever looked at any of this. The
    regime is restated here and nothing more; the consequences of it are stated once,
    under Research Goal, and pointed at from here.
    """
    profile = f"the {session.approval_profile} approval profile"
    # A gate is an acceptance. Every decision was counted here, and a run whose
    # scope draft a researcher edited before accepting it therefore reported "nine
    # gate decisions" two lines under "Stages completed: 8 of 8" -- the ninth being
    # that edit, which opened a gate rather than closing one. The count is off the
    # acceptances, which is one per stage, and the revisions are their own sentence
    # because sending a draft back is a fact about the run worth having.
    accepted = [
        decision for decision in session.decisions if decision.action == "accept"
    ]
    automatic = [decision for decision in accepted if decision.automatic]
    revised = [
        decision for decision in session.decisions if decision.action == "revise"
    ]
    sent_back = ""
    if revised:
        stages = _listed(_stage_words({decision.stage for decision in revised}))
        sent_back = (
            f", and {_plural(len(revised), 'draft')} "
            f"{'was' if len(revised) == 1 else 'were'} sent back for revision before "
            f"being accepted ({stages})"
        )
    if not accepted:
        return (
            f"Approvals: {profile} was in force and this run recorded no stage "
            "acceptance at all" + sent_back
        )
    if len(automatic) == len(accepted):
        return (
            f"Approvals: every stage gate in this run was accepted automatically "
            f"under {profile}, so acceptance here records a well-formed payload "
            f"rather than a person's agreement — the warning headed {AUTO_APPROVAL_WARNING} "
            "says what it does not amount to" + sent_back
        )
    if not automatic:
        return (
            "Approvals: every stage gate in this run was accepted by a person under "
            f"{profile}" + sent_back
        )
    # "3 of this run's nine stage gates" set one count as a digit and the other as a
    # word. Both are spelled, and the second count is stated rather than left as "the
    # rest", so a reader can check the arithmetic against the stage count above.
    return (
        f"Approvals: {_number_word(len(automatic)).lower()} of this run's "
        f"{_plural(len(accepted), 'stage gate')} were accepted automatically under "
        f"{profile} and the other "
        f"{_number_word(len(accepted) - len(automatic)).lower()} by a person — the "
        f"warning headed {AUTO_APPROVAL_WARNING} above names the automatic ones"
        + sent_back
    )


def _tournament_facts(tournament) -> list[str]:
    """The protocol as configured, and the tournament the comparisons actually record.

    One line used to print the configured protocol as though it were the schedule that
    was played: "3 Swiss rounds then a top-4 round robin, 12 matches in all" appeared
    over a record whose last round holds three matches between four ideas, which is
    half of the round robin it claims. ``swiss_rounds`` and ``top_round_robin_size``
    are settings; what happened is in ``comparisons``, and both are worth printing so
    long as neither is passed off as the other.

    The convergence clause is split off for a reading fault of its own: "stopped
    without converging with a final round that moved a rating by 46 points" attaches
    "with a final round" to "converging", so the sentence appeared to name the thing
    the tournament had failed to converge with.
    """
    rounds: dict[int, int] = {}
    for item in tournament.comparisons:
        rounds[item.round_number] = rounds.get(item.round_number, 0) + 1
    played = [rounds[number] for number in sorted(rounds)]
    facts = [
        f"Tournament protocol configured: {_counted(tournament.swiss_rounds, 'Swiss round')} "
        f"then a top-{tournament.top_round_robin_size} round robin"
    ]
    if not played:
        facts.append("Tournament as played: no match was recorded")
        return facts
    facts.append(
        f"Tournament as played: {_counted(len(played), 'round')} of "
        + _listed([str(count) for count in played])
        + f" matches, {len(tournament.comparisons)} in all"
        + _round_robin_shortfall(tournament, sorted(rounds))
    )
    facts.append(
        "Tournament outcome: "
        + ("converged" if tournament.converged else "stopped without converging")
        # score_movement is a fraction of the 1200 starting rating, not a rating.
        # "a final score movement of 0.04" reads as four hundredths of a point
        # against a round that actually moved a rating by forty-six of them. Its
        # sentinel is not a fraction at all, and multiplied out it printed "moved a
        # rating by 1200 points" in a run whose largest single move was sixteen.
        + (
            "; its final round moved a rating by "
            + _counted(round(tournament.score_movement * DEFAULT_ELO), "point")
            if tournament.score_movement < UNMEASURED_MOVEMENT
            else "; no rating change was recorded for its final round"
        )
    )
    return facts


def _round_robin_shortfall(tournament, rounds: list[int]) -> str:
    """Whether the round robin the settings configured is in the record, in full.

    A round robin over four ideas is six matches and nothing else, so a final round
    of three is not the configured protocol played out -- it is a fact about this run
    that the settings line would otherwise conceal.
    """
    final = tournament.swiss_rounds + 1
    if final not in rounds:
        return (
            "; no round beyond the Swiss rounds was recorded, so the configured "
            "round robin was never played"
        )
    matches = [item for item in tournament.comparisons if item.round_number == final]
    ideas = {item.candidate_a_id for item in matches} | {
        item.candidate_b_id for item in matches
    }
    complete = len(ideas) * (len(ideas) - 1) // 2
    if len(matches) >= complete:
        return ""
    return (
        f"; its last round played {_counted(len(matches), 'match')} between "
        f"{_counted(len(ideas), 'idea')}, where a round robin over them is "
        f"{_counted(complete, 'match')}"
    )


# What stopped the literature search, in the words of the thing that stopped it. The
# manifest records a machine token; the appendix is read by a person deciding how much
# of the field this report saw.
_DISCOVERY_STOPS = {
    "coverage_sufficient": (
        "the coverage of the question was judged sufficient, so no further pass ran"
    ),
    "maximum_passes_reached": (
        "the pass limit was reached with coverage still short of sufficient"
    ),
    "cost_limit_reached": (
        "the cost limit was reached with coverage still short of sufficient"
    ),
    "coverage_improvement_below_threshold": (
        "a further pass was judged unlikely to add enough to be worth running"
    ),
    "no_material_incremental_value": (
        "the last pass added no authoritative source and closed no material gap"
    ),
    "interaction_in_progress": (
        "the search was still running when the report was compiled, so what is here "
        "is a partial result"
    ),
    "deep_research_timed_out": "the search exceeded its deadline and was abandoned",
    "empty_deep_research_report": "the search returned without a report",
    # Three reasons the orchestrator records and this table did not hold, so the
    # fallback below printed them: a live appendix read "It stopped because the run
    # recorded gap_directed_search.", which names an identifier out of the source
    # and tells a reader deciding how much of the field this report saw nothing.
    "gap_directed_search": (
        "the last passes were aimed at the gaps the fan-out left open, and the "
        "search ended with them rather than searching the question again"
    ),
    "deep_research_start_failed": (
        "no pass of the search could be started, so nothing here came back from one"
    ),
}


def _stop_reason(reason: str) -> str:
    """What stopped the search, in words, whatever token the manifest recorded."""
    if reason in _DISCOVERY_STOPS:
        return _DISCOVERY_STOPS[reason]
    if reason.startswith("fan_out_truncated_by_budget:"):
        # The token carries its own payload: the facets the pass budget never
        # reached. Naming them is the point of the reason, and the fallback below
        # printed them as a comma-joined list of enum members.
        dropped = [
            FACET_PHRASES.get(facet, facet.replace("_", " "))
            for facet in reason.partition(":")[2].split(",")
            if facet
        ]
        return (
            "the pass budget was reached before the fan-out was complete, so "
            + _joined_titles(dropped, fallback="no facet")
            + (" was" if len(dropped) == 1 else " were")
            + " never searched"
        )
    return f"the run recorded {reason or 'no reason'}"


def _discovery_provenance(record: ResearchRecord) -> list[str]:
    """How the literature under this report was found, and by what.

    The appendix listed the evidence stage as a row in a table of schema names and
    said nothing else about it. On a live run the Deep Research agent -- the stage's
    designed path, an iterated search that stops when coverage converges -- never ran,
    and a single grounded search pass stood in for it. The manifest recorded the
    substitution in ``convergence_reason`` and the failure in ``runs[*].status``,
    neither of which anything read, so the report presented four leads from one query
    set as though a converging literature search had produced them.
    """
    discovery = record.discovery
    if discovery is None:
        return [
            "## Literature discovery",
            "",
            "No discovery manifest was recorded, so nothing states where the "
            "literature in this report came from or how much of the field it "
            "represents.",
            "",
        ]
    reason = discovery.convergence_reason
    passes = [run for run in discovery.runs if run.status == "completed"]
    attempted = len(discovery.runs)
    leads = len(discovery.source_leads)
    lines = ["## Literature discovery", ""]
    # The Knowledge Summary opens by saying a forked run did not search, but that
    # sentence guards one heading and this appendix is thirty pages further down.
    # Standing alone it read "Deep Research ran seven passes, at an estimated cost
    # of $21.00" as a plain fact about a run that spent nothing and ran none of
    # them -- the section of the report a reader goes to for exactly that number.
    if record.session.seeded_evidence_from:
        lines.extend(
            [
                "The search below is not this run's. Its scope and evidence base "
                "were carried over from "
                f"{record.session.seeded_evidence_from}, an earlier run of the same "
                "question, and the passes, the leads and the cost recorded here are "
                "what that run spent to build the corpus this one reasoned over. "
                "This run ran no pass of its own.",
                "",
            ]
        )
    if reason in DISCOVERY_STOOD_IN:
        failed = next((run for run in discovery.runs if run.error), None)
        lines.append(
            "The Deep Research agent did not run"
            + (f": {failed.error.rstrip('.')}. " if failed and failed.error else ". ")
            + (
                # "A single search-grounded pass, from one set of queries" was the
                # honest description of the stage when it ran one query. It now runs
                # one per facet and one per success criterion, and repeating the old
                # sentence would understate the search by an order of magnitude --
                # the opposite failure to the one the sentence was written to fix,
                # and just as misleading.
                "The literature in this report is what grounded search returned "
                f"instead — {_plural(leads, 'source lead')} from "
                f"{_plural(len(discovery.discovery_angles), 'parallel search')}, "
                "one for each facet of the evidence and each of the plan's success "
                "criteria. What no amount of parallel search does is iterate: "
                "Deep Research is the designed path for this stage because it runs "
                "again against the gaps the previous pass left open, and nothing "
                "here did that."
                if leads and discovery.discovery_angles
                else f"The literature in this report is what a single "
                f"search-grounded pass returned instead — "
                f"{_plural(leads, 'source lead')}, from one set of queries, with "
                "no second pass against the gaps the first left open. Deep "
                "Research is the designed path for this stage because it iterates "
                "until coverage stops improving; the breadth of what is cited "
                "below is that of one query set, not of a converged search."
                if leads
                else "Nothing stood in for it and no source lead was returned, so "
                "this report cites no literature at all."
            )
        )
    else:
        # This sentence was joined out of two counts written in two notations: the
        # attempted passes came from _plural, which prints a digit, and the completed
        # ones from _number_word, which prints a word. "Deep Research ran 1 pass, of
        # which one completed" reads as two different quantities, and on the common
        # run where every pass completes it is also the same quantity twice. Both
        # counts are words here, and the second clause is dropped when there is
        # nothing to contrast it with.
        lines.append(
            f"Deep Research ran {_opening(attempted, 'pass').lower()}"
            + (
                ""
                if len(passes) == attempted
                else f", of which {_number_word(len(passes)).lower()} completed"
            )
            + f", at an estimated cost of ${discovery.estimated_cost_usd:.2f}, and "
            f"returned {_plural(leads, 'source lead')}. It stopped because "
            + _stop_reason(reason)
            + "."
        )
    incomplete = [
        run for run in discovery.runs if run.status not in {"completed", "queued"}
    ]
    if incomplete and reason not in DISCOVERY_STOOD_IN:
        lines.extend(
            [
                "",
                *_bullets(
                    [
                        f"Pass {run.pass_number} ended {run.status.replace('_', ' ')}"
                        + (f": {run.error.rstrip('.')}." if run.error else ".")
                        for run in incomplete
                    ]
                ),
            ]
        )
    lines.extend(_carried_from_the_evidence_stage(record, leads))
    lines.extend(_sources_per_pass(record))
    lines.append("")
    return lines


def _carried_from_the_evidence_stage(record: ResearchRecord, leads: int) -> list[str]:
    """Why the corpus can be bigger than what the search returned.

    The evidence packet carries sources of its own, and any the search did not
    already list are added to the corpus so that a claim resting on one can be
    numbered and followed. Nothing said so. This paragraph read "returned 86 source
    leads" and the Knowledge Summary, twenty-two hundred lines above it, read "the
    literature search returned eighty-eight leads" -- the same search, two counts, and
    the two sources that account for the difference recorded only in the code.

    Compared before the folding, not after. The registry keys on the document, so its
    own total is what the leads resolve to; subtracting the search's raw count from
    that would have made the two sources carried in look like one.
    """
    _, documents = record.citations.verification_standing
    corpus = documents + record.citations.folded_duplicates
    carried = corpus - leads
    if carried <= 0:
        return []
    it = "it" if carried == 1 else "them"
    return [
        "",
        f"The corpus is larger than that: {_plural(carried, 'further source')} "
        f"{'was' if carried == 1 else 'were'} carried in from the evidence stage, "
        f"which listed {it} against a claim without this search having returned {it}. "
        # Through _plural, like every other count in this paragraph. Writing it
        # straight would have put "returned eight source leads" and "the corpus is
        # 9 leads" in adjacent sentences, the same quantity in two notations.
        f"Counted with {'that' if carried == 1 else 'those'}, the corpus is "
        f"{_plural(corpus, 'lead')}, which is the count the Knowledge Summary and "
        "the reference list work from.",
    ]


def _sources_per_pass(record: ResearchRecord) -> list[str]:
    """Which sources each pass found, which the knowledge summary promises is here.

    "Which sources a pass found is recorded per pass in the discovery appendix" stood
    at the head of the reproduced pass reports, and the appendix gave one total for
    the whole search and no way to reach a pass from it. Every lead in the manifest
    carries ``originating_passes``, so the breakdown was always on the record.
    """
    discovery = record.discovery
    leads = discovery.source_leads if discovery else []
    by_pass: dict[int, list[str]] = {}
    for lead in leads:
        for number in dict.fromkeys(lead.originating_passes):
            by_pass.setdefault(number, []).append(lead.canonical_url)
    if len(by_pass) < 2:
        return []
    rows = []
    for number in sorted(by_pass):
        found = by_pass[number]
        cited = sorted(
            {
                marked
                for url in found
                if (marked := record.citations.numbered(url)) is not None
            }
        )
        rows.append(
            f"Pass {number} returned {_plural(len(found), 'source lead')}"
            + (
                ", none of them cited in this report."
                if not cited
                else ", cited in this report as "
                + _names([f"[{marked}]" for marked in cited])
                + "."
            )
        )
    # No silent caps: a lead two passes returned is counted under both, so the rows
    # total more than the search returned, and a lead the manifest recorded no pass
    # for appears in no row at all.
    shared = sum(1 for lead in leads if len(set(lead.originating_passes)) > 1)
    unattributed = sum(1 for lead in leads if not lead.originating_passes)
    caveats = ""
    if shared:
        caveats += (
            f" {_opening(shared, 'lead')} came back from more than one pass and "
            + ("is" if shared == 1 else "are")
            + " counted under each of them, so the rows total more than the search "
            "returned."
        )
    if unattributed:
        caveats += (
            f" {_opening(unattributed, 'lead')} "
            + ("records" if unattributed == 1 else "record")
            + " no pass and "
            + ("is" if unattributed == 1 else "are")
            + " in no row at all."
        )
    return [
        "",
        "### Sources per pass",
        "",
        "The pass reports under Knowledge Summary point here for which sources a pass "
        "found. The numbers are the entries under References; a lead nothing in this "
        "report cites has no number to give." + caveats,
        "",
        *_bullets(rows),
    ]


def _advisory_appendix(advisories: Sequence[Advisory]) -> list[str]:
    """The warnings chapter: every run-level caveat, worst first, stated once.

    It opens the appendix half rather than closing it. Provenance is a record for
    someone auditing the report; this is the last thing a reader who is deciding
    whether to act on it needs to have read, so it comes first of the two.
    """
    if not advisories:
        return []
    lines = [f"# {ADVISORY_CHAPTER}", ""]
    blocking = [item for item in advisories if item.blocking]
    lines.extend(
        [
            (
                "Everything below qualifies the report as a whole. "
                + (
                    "None of it blocks the work proposed above."
                    if not blocking
                    else (
                        "The first entry says"
                        if len(blocking) == 1
                        else f"The first {_number_word(len(blocking)).lower()} "
                        "entries say"
                    )
                    + " the work should not proceed on the material "
                    + ("it names" if len(blocking) == 1 else "they name")
                    + "."
                )
                + " Caveats that apply to one hypothesis rather than to the run are "
                "under that hypothesis, not here."
            ),
            "",
        ]
    )
    for advisory in advisories:
        lines.extend([f"## {advisory.title}", ""])
        for paragraph in advisory.body.split("\n\n"):
            lines.extend([paragraph, ""])
    return lines


_REPAIR_REASON = re.compile(r"^(?P<detail>.*\S) \((?P<reason>[^()]+)\)$")

# How the repair pass names what it repaired: the path to the field in the model it
# was enforcing. "Candidate.score_novelty 6 -> 3" reached a live Provenance chapter
# twenty times in one sentence -- a Pydantic attribute path, printed to a reader who
# has never seen the model, for the one record of what the run changed.
_REPAIR_PATH = re.compile(r"^(?P<model>[A-Z]\w+)(?:\.(?P<field>\w+))?(?=[ :])")

# A rescaled number, so the field it rescaled can be named once for all of them.
_REPAIR_RESCALE = re.compile(r"^(?P<field>.*?) (?P<pair>-?[\d.]+ → -?[\d.]+)$")


def _spelled_model(name: str) -> str:
    """A model's class name as words: ``EvidenceItem`` is an evidence item."""
    return re.sub(r"(?<!^)(?=[A-Z])", " ", name).lower()


def _named_field(detail: str, *, qualified: bool) -> str:
    """One repair detail with its field named as the report names that field.

    The model is dropped unless the same stage repaired fields of more than one of
    them, because the sentence around this already says which answer was repaired
    and a lone qualifier only re-states it.
    """
    detail = detail.replace(" -> ", " → ")
    path = _REPAIR_PATH.match(detail)
    if not path:
        return detail
    field = path["field"] or ""
    spelled = _CONTRACT_FIELD_NAMES.get(field, field.replace("_", " "))
    lead = f"{_spelled_model(path['model'])} {spelled}" if qualified else spelled
    return f"{lead.strip()}{detail[path.end() :]}"


def _fielded(details: Sequence[str]) -> list[str]:
    """One field name over all the values rescaled in it, in the recorded order.

    The repair pass writes a line per value, so a stage that rescaled five scores
    for each of four ideas wrote the same five field names four times over. Printed
    as recorded that is twenty entries in which "novelty score" appears four times
    carrying four different numbers, and nothing in the list says the four are four
    ideas rather than four readings of one.
    """
    ordered: list[tuple[str, list[str]]] = []
    seen: dict[str, list[str]] = {}
    for detail in details:
        rescale = _REPAIR_RESCALE.match(detail)
        if not rescale:
            ordered.append((detail, []))
            continue
        field, pair = rescale["field"], rescale["pair"]
        if field in seen:
            seen[field].append(pair)
            continue
        seen[field] = [pair]
        ordered.append((field, seen[field]))
    return [f"{field} {_names(pairs)}" if pairs else field for field, pairs in ordered]


def _grouped_repairs(repairs: Sequence[str]) -> tuple[str, bool]:
    """The repair list as sentences: the shared reason once, each field named once.

    Every rescaled score carries the same parenthetical, and the renderer joined the
    raw strings. A live report printed "(answered on a 1-10 scale)" twenty times in
    one sentence -- once after each of the twenty fields the run rescaled -- which
    reads as twenty findings where there is one, and buries the field names it is
    the only record of. Grouped in first-seen order, so the sentence still walks the
    fields in the order the repair pass met them.
    """
    grouped: dict[str, list[str]] = {}
    for repair in repairs:
        match = _REPAIR_REASON.match(repair)
        if match:
            grouped.setdefault(match["reason"], []).append(match["detail"])
        else:
            # Its own group, keyed on itself so it is never folded into another's
            # reason, and printed back exactly as the repair pass wrote it.
            grouped.setdefault(f"\x00{repair}", []).append(repair)
    models = {
        path["model"]
        for details in grouped.values()
        for detail in details
        if (path := _REPAIR_PATH.match(detail))
    }
    repeated = False
    stated = []
    for reason, details in grouped.items():
        named = [_named_field(detail, qualified=len(models) > 1) for detail in details]
        folded = _fielded(named)
        repeated = repeated or len(folded) < len(named)
        if reason.startswith("\x00"):
            stated.extend(folded)
        else:
            stated.append("; ".join(folded) + f" ({reason})")
    # Sentence per reason, because a field's own rescalings are already a
    # semicolon-separated list and a second level of them would read as one list.
    # The first stays lower case: it follows the colon in the sentence that says
    # which answer was repaired.
    sentences = [stated[0], *(part[0].upper() + part[1:] for part in stated[1:])]
    # Whether any field was folded goes back to the caller rather than out on the
    # end of this sentence: four stages were repaired the same way, so the caveat
    # closed four consecutive paragraphs in the same thirty-one words -- the shape
    # this function exists to take out of the sentence, put back around it.
    return ". ".join(sentences) + ".", repeated


def _provenance_appendix(record: ResearchRecord) -> list[str]:
    """Where each stage's payload came from, so a template can never pass as reasoning."""
    lines = ["# Provenance", "", "## Run", "", *_run_facts(record), ""]
    lines.extend(_discovery_provenance(record))
    lines.extend(["## Evidence integrity", ""])
    integrity = evidence_integrity_lines(record)
    if integrity:
        # "The following ideas" stood over a list that on both live runs was every idea
        # in the run, and the list is grouped by case rather than by idea, so no line
        # under it counted them either. A reader auditing the report could work out
        # that nothing here has verified grounding only by tallying the titles against
        # the population themselves.
        every = all(item.qualified for item in record.evidence_support.values())
        # The lead-in used to agree with the line count, and two of the four cases
        # name every idea they cover on one line. So "The grounding of the following
        # ideas carries a qualification" stood over a single line naming a single
        # idea on the run that finished today.
        covered = evidence_integrity_ideas(record)
        # "Each line states one of the four and names the ideas it applies to" stood
        # over two lines, and the four cases it listed included two this run never
        # recorded. Only the cases that produced a line are named.
        cases = evidence_integrity_cases(record)
        # Comma before the final "or" at every count, not only past two. One of the
        # four cases is itself an alternation -- "its evidence was retracted or could
        # not be retrieved" -- so the bare join printed "retracted or could not be
        # retrieved or its evidence was never checked against its source", three
        # branches with nothing saying which two are the pair.
        stated = (
            cases[0] if len(cases) == 1 else ", ".join(cases[:-1]) + f", or {cases[-1]}"
        )
        lines.extend(
            [
                # This used to say the listed ideas "do not rest on the evidence they
                # cite", which is true of a missing or a retracted citation and flatly
                # untrue of the third case the same list prints -- an idea that does
                # rest on its evidence, which simply has not been checked against its
                # source. On a live run that third case was most of the list, so the
                # sentence contradicted the bullets directly beneath it.
                (
                    (
                        "The one idea in this run carries a qualification on its "
                        "grounding"
                        if covered == 1
                        else "Every idea in this run carries a qualification on its "
                        "grounding"
                    )
                    if every
                    # The cases are written in the singular, about one idea's own
                    # grounding, so the subject has to be singular too. "The grounding
                    # of the following ideas carries a qualification: its evidence was
                    # retracted" put a plural subject over four lines of "its".
                    else (
                        "The following idea carries a qualification on its grounding"
                        if covered == 1
                        else "Each of the following ideas carries a qualification on "
                        "its grounding"
                    )
                )
                + f": {stated}."
                + (
                    " The line below names the "
                    + ("idea" if covered == 1 else "ideas")
                    + " it applies to."
                    if len(integrity) == 1
                    else " Each line below states one of the "
                    f"{_number_word(len(cases)).lower()} and names the ideas it "
                    "applies to."
                    if len(cases) > 1
                    else " Each line below names the ideas it applies to."
                ),
                "",
                *_bullets(integrity),
                "",
            ]
        )
    elif record.evidence_support:
        lines.extend(
            [
                "Every evidence id cited by a candidate resolves to a record in this "
                "session, and no cited record was retracted or unretrievable.",
                "",
            ]
        )
    else:
        lines.extend(
            [
                "No candidate population was recorded, so no citation could be "
                "resolved and nothing here confirms that any grounding exists.",
                "",
            ]
        )
    # "Payload provenance" over a table of "Contract" and "Payload source" columns is
    # the run describing itself in its own words. The appendix is written for someone
    # auditing the report rather than reading it, but auditing it does not require
    # knowing what the code calls the thing a stage handed back.
    lines.extend(["## What each stage produced", ""])
    if not record.provenance:
        lines.extend(["No stage of this run recorded what it produced.", ""])
        return lines
    # A fork's carried-over stages are listed here exactly like the stages it ran, so
    # the heading claims work for the run that a reader can only discount by carrying
    # a caveat from two sections above. A live fork's first two rows were the scope
    # and the Deep Research discovery it inherited, printed under "What each stage
    # produced" between "This run searched no literature" and "This run ran no pass
    # of its own".
    ran = {note.stage for note in record.provenance}
    carried = [stage_name(stage) for stage in FORKED_STAGES if stage in ran]
    if record.session.seeded_evidence_from and carried:
        named = _listed(carried)
        executed = any(note.stage not in FORKED_STAGES for note in record.provenance)
        rest = " Every row below them is work this run did itself." if executed else ""
        lines.append(
            named[0].upper()
            + named[1:]
            + (" is" if len(carried) == 1 else " are")
            + " the forked run's work rather than this one's, carried over with the "
            "evidence base and listed here because the report rests on "
            + ("it" if len(carried) == 1 else "them")
            + " either way."
            + rest
        )
        lines.append("")
    # The source column held "specialist" on every row of a healthy run, and the
    # paragraph under the table said so again in words. A column that cannot vary is
    # not evidence; it is only worth a column on the runs where it does vary.
    varied = len({note.source for note in record.provenance}) > 1
    header = ["Stage", "Specialist", "What it produced", "Produced by"]
    # And the values under it as an answer to the question the header asks. The
    # column printed the enum the payload is filed under, so four rows answered
    # "Written by" with "repaired", which is not an author -- and the word is not
    # defined until four paragraphs below the table.
    if varied:
        header.append("Written by")
    lines.extend(
        [f"| {' | '.join(header)} |", f"| {' | '.join(['---'] * len(header))} |"]
    )
    # A run whose evidence stage verifies six times records six notes, and every
    # column of those notes holds the same value: the live table printed the row
    # "| evidence | source verification | evidence packet | gemini-3.1-pro-preview |"
    # six times over, which a reader can only read as a repeat until they count them.
    # The repeats are real work and are not dropped -- the count says how many.
    rows: list[tuple[list[str], int]] = []
    # Which stages ran one specialist and which fanned out, which is what decides
    # whether a specialist cell has anything to add to the stage cell beside it.
    worked = Counter(note.stage for note in record.provenance)
    for note in record.provenance:
        # The column printed the pipeline's ids -- "scope", "evidence", "generate",
        # "reflect" -- in a report that calls those same passes scoping the goal,
        # literature discovery, idea generation and independent review in every
        # sentence it writes about them, including the sentence directly under this
        # table. The id is the run's word for the stage, not the reader's.
        stage = stage_name(note.stage)
        row = [
            stage,
            _specialist_cell(
                stage,
                _specialist_label(record, note),
                worked[note.stage] == 1,
            ),
            _record_type(note.schema_name),
            _produced_by(note),
        ]
        if varied:
            row.append(_WRITTEN_BY.get(note.source, note.source.replace("_", " ")))
        if rows and rows[-1][0] == row:
            rows[-1] = (row, rows[-1][1] + 1)
            continue
        rows.append((row, 1))
    for row, count in rows:
        if count > 1:
            row = [*row]
            row[2] = f"{row[2]}, {_number_word(count).lower()} of them"
        lines.append(f"| {' | '.join(row)} |")
    lines.append("")
    if record.superseded_populations:
        lines.extend(
            [
                # "rewritten 1 time" is a count where prose has a word for it, and the
                # withdrawal that produced it is a single event on every run that has
                # one so far.
                "The candidate population was rewritten "
                + (
                    "once"
                    if record.superseded_populations == 1
                    else f"{_number_word(record.superseded_populations).lower()} times"
                )
                + " after a governance adjudication withdrew a hypothesis. Only the "
                "current version is listed above; the earlier version is retained in "
                "the session as superseded rather than deleted, so the population each "
                "stage actually ran against stays readable.",
                "",
            ]
        )
    if record.fallback_stages:
        lines.extend(
            [
                f"{_number_word(len(record.fallback_stages))} "
                + ("stage" if len(record.fallback_stages) == 1 else "stages")
                + " fell back to a fixed template because the specialist's answer "
                "came back incomplete or malformed. What that costs the report is "
                f"stated under {ADVISORY_CHAPTER}; on a healthy run no stage falls "
                "back at all.",
                "",
            ]
        )
    folded = False
    for note in record.repaired_stages:
        stated, repeated = _grouped_repairs(
            note.repairs or ["no repair detail was recorded"]
        )
        folded = folded or repeated
        lines.extend(
            [
                f"The {_agent_name(note.agent)} answer was repaired before it "
                "could be accepted: " + stated,
                "",
            ]
        )
    if folded:
        # Once, under all of them. Which idea's copy of a field a value belongs to
        # is not in the record, so the folded lists cannot say, and a reader
        # counting four novelty scores against eight ideas should not have to guess
        # whether they can.
        lines.extend(
            [
                "The repair pass records the field it repaired and not which idea's "
                "copy of it, so the values above are in the order it met them.",
                "",
            ]
        )
    # A stage whose answer was mostly the specialist's own is recorded as the
    # specialist's own, whole. Where the reviewer skipped an idea and the run filled the
    # gap from a template, the substitution is real, is printed under an idea, and is
    # nowhere in the table above -- so the count is stated here rather than left to the
    # reader to notice one review at a time.
    stood_in = record.stood_in_reviews
    if stood_in:
        one = len(stood_in) == 1
        lines.extend(
            [
                "The review stage's answer was accepted with "
                + (
                    "one review missing"
                    if one
                    else f"{_number_word(len(stood_in)).lower()} reviews missing"
                )
                + ": a reviewer answered for some of the ideas and not others, and "
                + ("the gap was" if one else "the gaps were")
                + " filled from the same fixed template a whole stage falls back to. "
                + ("That review carries" if one else "Those reviews carry")
                + " a verdict and a score that entered the averages, the spreads and "
                "the ranking, and "
                + ("it is named" if one else "each is named")
                + " where "
                + ("it is" if one else "they are")
                + " printed. The stage is listed above as the specialist's own "
                "because the rest of it was.",
                "",
            ]
        )
    if not record.fallback_stages and not record.repaired_stages and not stood_in:
        # "nothing substituted" was printed unconditionally, directly under a table
        # whose evidence row was produced by the stand-in for a specialist that never
        # ran. A template substitution is not the only kind there is, so the sentence
        # is scoped to what it actually checked and defers to the section that knows.
        substituted = record.deep_research_stood_in
        lines.extend(
            [
                "Every stage was accepted on the specialist's own answer, with "
                "nothing repaired and no stage falling back to a fixed template."
                + (
                    " Discovery is the exception, and what stood in for it is set "
                    "out under Literature discovery above."
                    if substituted
                    else ""
                ),
                "",
            ]
        )
    return lines


def compile_dossier(session: Session) -> str:
    """Compile the report as a research document rather than an artifact dump.

    The nine-part structure — front matter, overview, narrative, directions, review
    summary, knowledge base, references, per-idea deep dives — mirrors the reference
    reports this project is modelled on. Everything it prints is derived from a
    validated payload; nothing is restated as raw JSON, because a reader cannot audit
    what they cannot read.
    """
    record = load_record(session)
    overview = synthesize_overview(record)
    briefs = build_idea_briefs(record)
    # A grounding verdict more than one idea carries is explained once above them all
    # rather than in full under each of the eight. What it means was already explained
    # where the ideas were listed, so here it is the counts and a pointer to that.
    grounding, hoisted = shared_support_notices(
        [brief.support for brief in briefs], detail=False
    )
    # Likewise the reviewer and the question under each review heading, which are
    # properties of the review rather than of the idea it was written about.
    questions, hoisted_questions = shared_review_questions(briefs)
    # And who decided the matches nobody argued, which is the tournament's fact rather
    # than any one idea's.
    match_notes, hoisted_matches = shared_match_notes(briefs)
    # And the count of what each review raised and answered, where it did not vary.
    tally, tally_hoisted = shared_review_tally(briefs)
    # Collected before anything is laid out, because the overview carries the count
    # and the count cannot be taken from a chapter that has not been built yet.
    advisories = run_advisories(record, overview=overview, briefs=tuple(briefs))
    authors_own: list[str] = []
    lines = _front_matter(record, overview)
    lines += _overview_body(record, overview, advisory_pointer(advisories))
    lines += _knowledge_base(record, overview)
    if briefs:
        # The reference reports reopen the deep-dive half with a second "Top ideas",
        # which is what separates the per-idea sections from the overview above. It
        # is the parent of the idea sections, so it is a level above them. Qualified
        # because the overview's own "Top ideas" heading is still on the page: two
        # entries reading "Top ideas" in one table of contents point at different
        # halves of the report and give a reader no way to tell which is which.
        lines += ["# Top ideas in detail", ""]
        # Before the fourteen paragraphs that explain how to read the sections,
        # because a reader who wants one idea should not have to read the manual
        # to find out which one. Everything in the row is in the idea's own
        # section below; the row is the handle for reaching it.
        # The judge's reading only. The computed fallback is this same table in
        # sentences, and printing it directly under the table it restates would
        # cost a paragraph to say nothing new.
        lines += _candidate_summary_table(
            briefs,
            record.tournament.briefing
            if record.tournament and record.tournament.briefing_author == "judge"
            else "",
        )
        # Under a heading of its own, at the level of the sections around it. The
        # manual used to run on unheaded from the bottom of the summary table to
        # the first idea -- fourteen paragraphs a reader had no way to recognise
        # as a manual, skip past, or come back to.
        lines += [READING_GUIDE_HEADING, ""]
        for paragraph in DEEP_DIVE_PREAMBLE:
            lines += [paragraph, ""]
        if grounding:
            lines += [grounding, ""]
        # After the verdict on each idea's own grounding and before the reviews: a
        # reader who has just been told three ideas are grounded is owed the fact
        # that it is one finding grounding them.
        lines += shared_grounding_reach(record, briefs)
        lines += questions
        lines += tally
        lines += shared_coherence_notes(briefs)
        lines += match_notes
        # And what the specialist's own sections are, which is a fact about the
        # generation contract rather than about any of the ideas that carry one.
        authors_own += shared_authors_own_note(briefs)
        lines += authors_own
    # Shared across the chapters, so the second idea of a pair points at the first
    # rather than reprinting the exchange they both played in.
    transcribed: set[tuple[int, frozenset[str]]] = set()
    for brief in briefs:
        lines += _idea_deep_dive(
            record,
            brief,
            grounding_hoisted=brief.support in hoisted,
            authors_own_hoisted=bool(authors_own),
            hoisted_questions=frozenset(hoisted_questions),
            hoisted_matches=hoisted_matches,
            tally_hoisted=tally_hoisted,
            transcribed=transcribed,
        )
    lines += _advisory_appendix(advisories)
    lines += _provenance_appendix(record)
    report = _em_dashed(
        _densely_numbered(_without_math_markup("\n".join(lines).rstrip() + "\n"))
    )
    # Numbering runs before the contents list, so the index of exhibits it adds
    # is itself an entry in the contents rather than a section nothing points at.
    return table_of_contents(number_figures_and_tables(report))


SUMMARY_TABLE_HEADING = "## Executive Candidate Summary"
READING_GUIDE_HEADING = "## How to read each idea"
CHAPTER_SECTIONS = frozenset(
    heading.removeprefix("## ")
    for heading in (SUMMARY_TABLE_HEADING, READING_GUIDE_HEADING)
)
"""The sections of the deep-dive chapter that are not an idea."""

_SUMMARY_CELL_CEILING = 140

# Words that cannot be the last word of an abbreviated statement: each one governs
# something the cut has taken away, so a cell ending on one reads as a cell that was
# cut off rather than one that was shortened.
_GOVERNING_WORDS = frozenset(
    """a an the and or nor but so if while when where which who whose that than then
    at by for from in into of on onto over to under with within without via versus
    about after before during per across between among against toward towards upon
    leading resulting causing compared relative due such is are was were be been
    its their his her our this these those""".split()
)


def _cell(text: str, ceiling: int | None = _SUMMARY_CELL_CEILING) -> str:
    """One statement, short enough that the row it is in stays one screen wide.

    A pipe inside a cell ends the cell, so any that survived into a claim would
    silently split the row into columns the header has no names for.

    ``ceiling=None`` prints the cell whole. Abbreviating what a row says is fine;
    abbreviating what a row *is* is not, and this cut two candidate titles -- "ALD
    Alumina Conformal Barrier for Transition Metal Dissolution…" -- so the reader
    could not match the row to the section heading it summarises, in a report with
    two ideas whose full names both open "Sacrificial H".
    """
    flat = " ".join(plain_text(text).split()).replace("|", "—")
    if ceiling is None or len(flat) <= ceiling:
        return flat
    # On a word boundary. The bare slice cut inside words -- "improving capacity
    # reten…", "by the chemical red…" -- in five of the eight rows of one live
    # summary table, which reads as a rendering fault rather than as an
    # abbreviation. An unclosed bracket goes with the word that opened it.
    head = flat[: ceiling - 1]
    # And on a clause boundary where there is one to cut on. Cut inside a clause, the
    # cell states half a condition: every Falsifier Summary cell of a live summary
    # table ended on "or if the capacity retention after…" or "compared to…", which
    # is a cell a reader cannot hold the idea to. Ending the cell where the writer
    # ended a clause leaves a statement that is short rather than unfinished.
    clause = max(head.rfind(", "), head.rfind("; "), head.rfind(" — "))
    space = head.rfind(" ")
    for cut in (clause, space):
        if cut > ceiling // 2:
            head = head[:cut]
            break
    # A clause boundary is not always there to cut on, and the word boundary that
    # stands in for it can fall after a word that governs the words it cut. Five
    # Primary Claim and Falsifier Summary cells of a live table ended "at the same…",
    # "than the…", "while…", "into…" and "leading to…" -- each one an abbreviation
    # that reads as a truncation, because the last word left standing is the one
    # asking for what follows it.
    words = head.split()
    while len(words) > 1 and words[-1].strip("(),;:—-").lower() in _GOVERNING_WORDS:
        words.pop()
    head = " ".join(words)
    if head.count("(") > head.count(")"):
        head = head[: head.rfind("(")]
    return head.rstrip(" ,;.:—-([") + "…"


def _candidate_summary_table(briefs: Sequence, briefing: str = "") -> list[str]:
    """Every idea on one page, in the order the tournament ranked them.

    The deep-dive half opened on eight sections of roughly a hundred and fifty
    lines each, with no way to see what the eight were without reading them. A
    reader comparing ideas -- which is what a ranked population is for -- had to
    hold eight claims in their head across nine pages.
    """
    if not briefs:
        return []
    lines = [
        # Level two, with the eight idea sections it summarises. At level three it
        # was a child of them in the contents while standing above all of them on
        # the page, and its parent heading is level one, so the document skipped
        # a level here and nowhere else.
        SUMMARY_TABLE_HEADING,
        "",
        "| Rank | Candidate Title | Strategy | Primary Claim | Falsifier Summary | Elo | Evidence |",
        "| ---: | --- | --- | --- | --- | ---: | --- |",
    ]
    for brief in briefs:
        shortlist = " ★" if brief.shortlisted else ""
        lines.append(
            f"| {brief.rank} | **{_cell(brief.title, None)}**{shortlist} "
            f"| {_cell(brief.strategy, 24)} "
            f"| {_cell(brief.facts.get('Core idea', ''))} "
            f"| {_cell(brief.facts.get('Falsifier', ''))} "
            # The verdict as the rest of the report words it. The raw field went
            # into this column, so one row read "partially_grounded" against the
            # paragraph below that calls the same verdict "partially grounded".
            f"| {brief.elo} | {brief.support_label or brief.support} |"
        )
    lines.extend(
        [
            "",
            "A star marks an idea the tournament shortlisted. Elo is the rating it "
            "finished the tournament on, and the evidence column is the grounding "
            "verdict explained under Candidate Ideas above.",
            "",
        ]
    )
    if briefing:
        # A column of ratings says which idea finished ahead; it does not say what
        # the matches decided it on, or which of these gaps is too narrow to read
        # anything into. The judge that played the tournament says both here.
        lines.extend(["**What the tournament found.** " + briefing.strip(), ""])
        corrected = _places_the_run_never_produced(briefing, briefs)
        if corrected:
            lines.extend([corrected, ""])
    return lines


# Place words a judge writes, and the position each one claims. Only as far as the
# report ever ranks; past that a judge writes the numeral.
_PLACE_WORDS = {
    "first": 1,
    "second": 2,
    "third": 3,
    "fourth": 4,
    "fifth": 5,
    "sixth": 6,
    "seventh": 7,
    "eighth": 8,
    "ninth": 9,
    "tenth": 10,
}
_PLACE_WORD = re.compile(rf"\b(?:{'|'.join(_PLACE_WORDS)})\b", re.IGNORECASE)
# A judge writes the places as a series and puts the noun after the last of them --
# "the fourth, fifth, and sixth place hypotheses" -- so a word is a place only where
# the series it sits in ends in one. "the first move" and "a second coating" are not
# positions, and reading them as positions would answer a sentence nobody wrote.
_PLACE_SERIES = re.compile(
    rf"{_PLACE_WORD.pattern}(?:(?:,\s*|\s+and\s+|,\s*and\s+){_PLACE_WORD.pattern})*"
    r"[\s-]+(?:place|placed|ranked|position)\b",
    re.IGNORECASE,
)


def _places_the_run_never_produced(briefing: str, briefs: Sequence) -> str:
    """Where the judge counted places through a tie the standings do not count through.

    The briefing is the judge's prose over a table of computed positions, and on a
    live run the two disagreed about how many places the run had: three ideas finished
    on an Elo of 1184 and the table gave all three position 4, while the paragraph
    above it asked readers "not to draw distinctions between the fourth, fifth, and
    sixth place hypotheses". No idea in that run finished fifth or sixth. Rewriting
    the judge's sentence would be the report putting words in the judge's mouth, so
    the standings answer it instead, and only where a place it names is one nothing
    in the run holds.
    """
    ranks = Counter(brief.rank for brief in briefs)
    spoken = {
        word.lower()
        for series in _PLACE_SERIES.finditer(briefing)
        for word in _PLACE_WORD.findall(series.group(0))
    }
    named = sorted(
        _PLACE_WORDS[word] for word in spoken if _PLACE_WORDS[word] not in ranks
    )
    ties = [
        (rank, [brief for brief in briefs if brief.rank == rank])
        for rank in sorted(ranks)
        if ranks[rank] > 1
    ]
    if not named or not ties:
        return ""
    counted = _listed([_ORDINALS[place] for place in named])
    stated = _listed(
        [
            f"{_number_word(len(group)).lower()} ideas finished level on an Elo of "
            f"{group[0].elo} and share position {rank}"
            for rank, group in ties
        ]
    )
    return (
        f"The paragraph above is the judge's, and it counts places this run did not "
        f"produce: nothing here finished {counted}. What the standings hold is a tie "
        f"— {stated} — and the positions in the table are the ones the tournament "
        "returned."
    )


# The place words again, by position, for saying back what the judge named.
_ORDINALS = {place: word for word, place in _PLACE_WORDS.items()}


# A specialist writing chemistry and units reaches for TeX, and nothing downstream
# of here renders it: a live report carried "aluminum oxide ($Al_2O_3$)" in the body
# and "a precisely controlled $\mathbf{1\text{--}5 \text{ nm}}$ cathode coating"
# across a page of the Knowledge Base, in the Markdown, the PDF and the DOCX alike.
# What a reader wants out of those is Al2O3 and a 1-5 nm coating.
_MATH_SPAN = re.compile(r"(?<!\$)\$(?!\$)([^\$\n]{1,240})\$(?!\$)")
_TEX_WRAPPER = re.compile(r"\\(?:mathbf|mathrm|mathit|textbf|textit|text|bm)\s*\{")
_TEX_SYMBOLS = {
    r"\geq": "\u2265",
    r"\ge": "\u2265",
    r"\leq": "\u2264",
    r"\le": "\u2264",
    r"\approx": "\u2248",
    r"\times": "\u00d7",
    r"\cdot": "\u00b7",
    r"\pm": "\u00b1",
    r"\sim": "~",
    r"\circ": "\u00b0",
    r"\degree": "\u00b0",
    r"\mu": "\u00b5",
    r"\alpha": "\u03b1",
    r"\beta": "\u03b2",
    r"\Delta": "\u0394",
    r"\rightarrow": "\u2192",
    r"\to": "\u2192",
    r"\%": "%",
    r"\&": "&",
    r"\$": "$",
}
_TEX_SUPERSCRIPTS = str.maketrans(
    "0123456789+-=()n",
    "\u2070\u00b9\u00b2\u00b3\u2074\u2075\u2076\u2077\u2078"
    "\u2079\u207a\u207b\u207c\u207d\u207e\u207f",
)
_TEX_SUPERSCRIPT = re.compile(r"\^\{?([^{}\s$]{1,6})\}?")
_TEX_SUBSCRIPT = re.compile(r"_\{([^{}]{1,12})\}|_([A-Za-z0-9])")
_TEX_COMMAND = re.compile(r"\\[A-Za-z]+|\\[,;:!]|\\ ")


def _is_math(span: str) -> bool:
    """Whether a dollar-delimited span is notation rather than two sums of money.

    "$21.00" is the estimated cost of a discovery run and appears in the provenance
    appendix; two of those on one line must not read as a formula between them.
    """
    return bool(set(span) & set("\\_^{")) or span[:1].isalpha()


def _plain_superscript(match: re.Match[str]) -> str:
    body = match.group(1)
    lifted = body.translate(_TEX_SUPERSCRIPTS)
    # Only where every character has a raised form. "x^{max}" lifted letter by letter
    # would be unreadable, and a caret a reader can parse beats one they cannot.
    return lifted if not set(lifted) & set(body) else f"^{body}"


def _without_math_markup(report: str) -> str:
    """Every inline TeX span in the report, set as the plain text it stands for."""

    def spelled(match: re.Match[str]) -> str:
        span = match.group(1)
        if not _is_math(span):
            return match.group(0)
        # Wrappers first, so "\mathbf{n \geq 5}" is unwrapped before its contents are
        # read; the closing brace goes with the rest of the braces at the end.
        text = _TEX_WRAPPER.sub("", span)
        for command, symbol in sorted(_TEX_SYMBOLS.items(), key=len, reverse=True):
            text = text.replace(command, symbol)
        text = _TEX_SUBSCRIPT.sub(lambda hit: hit.group(1) or hit.group(2), text)
        text = _TEX_SUPERSCRIPT.sub(_plain_superscript, text)
        text = _TEX_COMMAND.sub(" ", text).replace("{", "").replace("}", "")
        # A double hyphen inside notation is a range -- "1\text{--}5 \text{ nm}" -- and
        # the em-dash pass below only touches the spaced form used in prose.
        return " ".join(text.replace("--", "\u2013").split()) or match.group(0)

    return _MATH_SPAN.sub(spelled, report)


def _em_dashed(report: str) -> str:
    """Set every parenthetical dash in the prose as an em dash.

    The report is written with em dashes and had a handful of ASCII double hyphens left
    in the emitted strings. In a typeset PDF that is the first thing a reader sees: the
    same page carries both, and the document reads as an un-proofed draft. Fixing the
    strings fixes today's; doing it here as well means a ``--`` typed into a future
    string cannot reach the page. Fenced code is left alone, where a double hyphen is
    a command-line argument rather than punctuation.
    """
    out: list[str] = []
    fenced = False
    for line in report.split("\n"):
        if line.lstrip().startswith("```"):
            fenced = not fenced
        out.append(line if fenced else line.replace(" -- ", " — "))
    return "\n".join(out)


_MARKER = re.compile(r"\[(\d+(?:, \d+)*)\]")
_REFERENCE_ENTRY = re.compile(r"^(\d+)\. (.+)$")


def _densely_numbered(report: str) -> str:
    """Drop any reference the finished report never cites, and close the gap.

    A number is assigned to a source the first time the narrative builder reaches for
    it, but a paragraph can be dropped afterwards to hold the section inside its word
    band. The source keeps its number and the report then lists a reference nothing
    points at -- a live run ended on "[4] Stable High-Voltage LiCoO2 Cathode Enabled
    by LiF Nanoshells" under a sentence claiming all four were cited. Numbering can
    only be settled once the text is final, so it is settled here.
    """
    head, separator, tail = report.partition("\n## References\n")
    if not separator:
        return report
    entries, rest = _split_reference_list(tail)
    prose = head + rest
    used = sorted(
        {
            int(number)
            for group in _MARKER.findall(prose)
            for number in group.split(", ")
        }
    )
    if used == [index for index, _ in enumerate(entries, start=1)]:
        return report
    renumbered = {old: new for new, old in enumerate(used, start=1)}

    def close_gaps(text: str) -> str:
        return _MARKER.sub(
            lambda match: (
                "["
                + ", ".join(
                    str(renumbered[int(number)])
                    for number in match.group(1).split(", ")
                )
                + "]"
            ),
            text,
        )

    kept = [f"{renumbered[old]}. {text}" for old, text in entries if old in renumbered]
    kept_block = "\n".join(kept) if kept else _NO_CITED_SOURCE
    return f"{close_gaps(head)}\n## References\n\n{kept_block}\n\n{close_gaps(rest)}"


def _split_reference_list(tail: str) -> tuple[list[tuple[int, str]], str]:
    """The numbered entries under References, and everything that follows them."""
    entries: list[tuple[int, str]] = []
    lines = tail.splitlines(keepends=True)
    for index, line in enumerate(lines):
        match = _REFERENCE_ENTRY.match(line.strip())
        if match:
            entries.append((int(match.group(1)), match.group(2)))
        elif entries and line.strip():
            return entries, "".join(lines[index:])
    return entries, "".join(lines[len(entries) :]) if entries else tail


_NO_CITED_SOURCE = (
    "No source was cited in this report, because discovery resolved no external "
    "literature for this goal."
)


_TITLE = "Co-Scientist Research Dossier"
# "review artifacts" named a part of the system on the one page every reader sees.
# The notice is the first thing on the cover and the last thing quoted back, so it
# says what the document holds in the words the document itself uses.
_DEFAULT_NOTICE = (
    "Draft for internal review. This report holds proposed hypotheses and the "
    "reviews they received, not verified findings; every claim must be independently "
    "verified against primary sources before it is acted upon."
)


def _without_cover_notice(content: str) -> str:
    """The compiled Markdown without the notice these two exporters set themselves.

    The notice is cover matter, and each format puts it where that format's cover
    is: under the title in Markdown, on the title page in the PDF and the DOCX.
    Left in the body for those two it would be set twice, a page apart.
    """
    return content.replace(f"*{_DEFAULT_NOTICE}*\n\n", "", 1)


_SERIF = "Times-Roman"
_SERIF_BOLD = "Times-Bold"
_SERIF_ITALIC = "Times-Italic"
_MONO = "Courier"
_CJK = "STSong-Light"
_MARGIN = 54
# SimpleDocTemplate pads its frame by 6pt on each side, and ``document.width`` is the
# frame's outer width, so anything sized to it is 12pt too wide. Paragraphs are laid
# out inside the frame and simply reflow; a Table is given an explicit column total and
# does not, so every table, code box and quote in the report hung 6pt past the right
# margin -- visibly further out than the text above and below it.
_FRAME_PADDING = 6

# A print hierarchy against the 10.5pt body below, not the browser's.
#
# These used to be {32, 24, 18.7, 16, 13.3, 10.7}, which is the CSS default cascade
# (2em, 1.5em, 1.17em, 1em, 0.83em, 0.67em) on a 16px base, copied across as points.
# Two things go wrong when it lands on a 10.5pt body. The top is enormous: a 32pt H1
# is 3.05x body where the reference sets 2x, and on A4 it takes three lines for a
# heading that is a question. The bottom inverts: H5 and H6 come out at or below body
# size, and this report leans on both -- every idea carries eight H5 subsections -- so
# the deepest headings were smaller than the text they introduced. H6 in particular
# sat 0.2pt above body in a non-bold serif, which is not a heading, it is a paragraph.
_HEADING_SIZES = {1: 21.0, 2: 16.5, 3: 14.0, 4: 12.5, 5: 11.5, 6: 10.5}
# Word renders on-screen at 11pt body, so the same hierarchy is compressed.
_DOCX_HEADING_SIZES = {1: 24.0, 2: 18.0, 3: 15.0, 4: 13.0, 5: 12.0, 6: 11.0}
# A4, shared by both exports. The DOCX used to inherit python-docx's US Letter
# template, so the two renderings of one report paginated differently and the column
# arithmetic below was handed a text measure the DOCX did not actually have.
_PAGE_WIDTH_MM = 210.0
_PAGE_HEIGHT_MM = 297.0
# Table cells are set smaller than body text so the wider grids fit the measure. The
# DOCX used to size its columns with this number and then set the cells at the 11pt
# Normal size, which is the mid-word header break _column_widths exists to prevent.
_CELL_FONT_SIZE = 8.5
_TOC_LEVELS = {"DossierH1": 0, "DossierH2": 1, "DossierH3": 2}
_BULLETS = ("\u2022", "\u2013", "\u00b7")
# Every list marker in the report is drawn from this face rather than from the body
# serif. reportlab's winansi codec encodes U+2022 as byte 0x7f, which conforming
# viewers draw as a bullet -- the PDF spec makes unused WinAnsi codes render as one --
# but which carries no ToUnicode mapping, so every bullet in the document came out of
# copy-paste, of a screen reader, and of any text-extraction pipeline as the U+007F
# control character. Bitstream Vera ships inside reportlab, so pinning the markers to
# a TrueType face adds no dependency and gives them a real ToUnicode entry.
_BULLET_FACE = "DossierMarker"


def _register_pdf_fonts() -> None:
    """Register the CID and marker faces: CJK goals and list bullets both need one."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    if _CJK not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(_CJK))
    pdfmetrics.registerFontFamily(
        _CJK, normal=_CJK, bold=_CJK, italic=_CJK, boldItalic=_CJK
    )
    if _BULLET_FACE not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(_BULLET_FACE, "Vera.ttf"))


def _pdf_styles() -> dict:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.styles import ParagraphStyle

    body = ParagraphStyle(
        "DossierBody",
        fontName=_SERIF,
        fontSize=10.5,
        leading=15.5,
        alignment=TA_LEFT,
        spaceAfter=7,
        allowWidows=0,
        allowOrphans=0,
    )
    styles = {"body": body}
    for level, size in _HEADING_SIZES.items():
        styles[f"h{level}"] = ParagraphStyle(
            f"DossierH{level}",
            parent=body,
            fontName=_SERIF_BOLD,
            fontSize=size,
            leading=size * 1.22,
            spaceBefore=size * 0.7,
            spaceAfter=size * 0.32,
            keepWithNext=1,
        )
    # H6 sets at body size, so bold is the only thing marking it as a heading. It used
    # to be set in the roman face, which left it indistinguishable from a paragraph.
    styles["title"] = ParagraphStyle(
        "DossierTitle",
        parent=styles["h1"],
        fontSize=26,
        leading=31,
        alignment=TA_CENTER,
        spaceBefore=0,
    )
    styles["subtitle"] = ParagraphStyle(
        "DossierSubtitle",
        parent=body,
        fontSize=14,
        leading=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#202124"),
    )
    styles["title_meta"] = ParagraphStyle(
        "DossierTitleMeta",
        parent=body,
        fontSize=10,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#5f6368"),
    )
    styles["notice"] = ParagraphStyle(
        "DossierNotice",
        parent=body,
        fontName=_SERIF_ITALIC,
        fontSize=9.5,
        leading=14,
        alignment=TA_CENTER,
        leftIndent=36,
        rightIndent=36,
        textColor=colors.HexColor("#5f6368"),
    )
    styles["toc_title"] = ParagraphStyle(
        "DossierTocTitle", parent=styles["h2"], spaceBefore=0
    )
    styles["quote"] = ParagraphStyle(
        "DossierQuote",
        parent=body,
        fontName=_SERIF_ITALIC,
        leftIndent=18,
        rightIndent=12,
        textColor=colors.HexColor("#3c4043"),
    )
    styles["code"] = ParagraphStyle(
        "DossierCode",
        parent=body,
        fontName=_MONO,
        fontSize=8,
        leading=10,
        # Space must clear borderPadding: the tinted box is drawn outside the
        # flowable's measured height and would otherwise collide with neighbours.
        spaceBefore=13,
        spaceAfter=15,
        leftIndent=6,
        rightIndent=6,
        borderPadding=6,
        backColor=colors.HexColor("#f5f6f7"),
        borderColor=colors.HexColor("#dadce0"),
        borderWidth=0.5,
    )
    styles["appendix_code"] = ParagraphStyle(
        "DossierAppendixCode", parent=styles["code"], fontSize=7, leading=8.6
    )
    styles["summary"] = ParagraphStyle(
        "DossierSummary",
        parent=body,
        fontName=_SERIF_BOLD,
        fontSize=10,
        leading=13,
        spaceBefore=6,
        spaceAfter=2,
    )
    styles["table_cell"] = ParagraphStyle(
        "DossierTableCell",
        parent=body,
        fontSize=_CELL_FONT_SIZE,
        leading=11,
        spaceAfter=0,
    )
    styles["table_header"] = ParagraphStyle(
        "DossierTableHeader", parent=styles["table_cell"], fontName=_SERIF_BOLD
    )
    styles["toc"] = [
        ParagraphStyle(
            "DossierToc0",
            parent=body,
            fontName=_SERIF_BOLD,
            fontSize=11.5,
            leading=17,
            spaceBefore=7,
            firstLineIndent=-14,
            leftIndent=14,
        ),
        ParagraphStyle(
            "DossierToc1",
            parent=body,
            fontSize=10,
            leading=14,
            firstLineIndent=-14,
            leftIndent=32,
        ),
        ParagraphStyle(
            "DossierToc2",
            parent=body,
            fontSize=9,
            leading=12.5,
            firstLineIndent=-14,
            leftIndent=50,
            textColor=colors.HexColor("#3c4043"),
        ),
    ]
    return styles


_PICTOGRAPHS = re.compile("[\U0001f000-\U0001faff☀-➿️]")

# What the body face can draw at all. Times-Roman is encoded in WinAnsi here, and
# reportlab draws a filled box for every character outside it -- so anything the
# Markdown carries that WinAnsi does not hold has to be turned into something that
# says the same thing before it reaches the page.
_PDF_ENCODING = "cp1252"

# The folds, each one a character a live report carried into the PDF. "significant"
# came off page 15 as "signi<box>cant", from the ligature a model copied out of a
# source PDF, and "1-5 nm" as "1<box>5 nm" from a typographic hyphen. The rest are
# their near neighbours, folded here so the next run does not find them first.
_PDF_FOLDS = {
    "\ufb00": "ff",
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
    "\u2010": "-",
    "\u2011": "-",
    "\u2012": "-",
    "\u2212": "-",
    # ACS sets the dash in a paper's name as a box-drawing rule, and a title arrives
    # here with it in. Named rather than folded, it printed mid-title in a live
    # reference list: "Degradation Effects in Li4Ti5O12-Based
    # Cells(box drawings light horizontal)Learning from Electrode Potential
    # Profiles". It is punctuation wherever this report meets it -- nothing here
    # draws a table out of characters -- so it is set as the dash it is being used as.
    "\u2500": "\u2014",
    "\u2501": "\u2014",
    "\u2015": "\u2014",
    "\u2e3a": "\u2014",
    "\u2e3b": "\u2014",
    "\u2043": "-",
    # The shortlist mark. The star is inside the pictograph range, so it was dropped
    # from all ninety-two pages of a live PDF -- under a table whose own caption says
    # "a star marks an idea the tournament shortlisted", over a column of nothing.
    "\u2605": "*",
    "\u2606": "*",
    "\u2032": "'",
    "\u2033": '"',
    "\u2044": "/",
    "\u2215": "/",
    "\u2192": "->",
    "\u2190": "<-",
    "\u2194": "<->",
    "\u21d2": "=>",
    "\u2264": "<=",
    "\u2265": ">=",
    "\u2248": "~",
    "\u2260": "!=",
    # The micro sign is in WinAnsi and the Greek letter is not, and a run writing
    # micrometres reaches for whichever one its source used.
    "\u03bc": "\u00b5",
    # Spaces that are not the space bar, and the widths that are no width at all.
    "\u2007": " ",
    "\u2009": " ",
    "\u200a": " ",
    "\u202f": " ",
    "\u200b": "",
    "\u2060": "",
    "\ufeff": "",
}
_PDF_FOLDED = re.compile(f"[{re.escape(''.join(_PDF_FOLDS))}]")


def _spelled(char: str) -> str:
    """A character with no glyph, written out rather than quietly dropped.

    Dropping is right for decoration and wrong for everything else: "$\\Delta$G" with
    the delta taken out is a different quantity, and nothing on the page would say so.
    Greek is spelled the way plain text has always spelled it, and anything rarer is
    named, because a name a reader can look up beats both a filled box and a silence.
    """
    try:
        name = unicodedata.name(char)
    except ValueError:
        return ""
    if name.startswith("GREEK ") and " LETTER " in name:
        letter = name.split(" LETTER ", maxsplit=1)[1].split(" ")[0].lower()
        return letter.capitalize() if "CAPITAL" in name else letter
    return f"({name.lower()})"


def _renderable(text: str) -> str:
    """The text with every character the PDF's Type 1 faces cannot set folded away.

    Three populations reach here. Emoji are decoration -- the attribution stamp
    carries one, and the stamp says who prepared the report in words -- so they go.
    Characters with a plain-text equivalent are folded to it. Everything else is
    spelled out, because a report about battery chemistry cannot afford to lose a
    character out of a formula without saying that it did.
    """
    folded = _PDF_FOLDED.sub(lambda match: _PDF_FOLDS[match.group(0)], text)
    stripped = _PICTOGRAPHS.sub("", folded)
    if any(not _drawable(char) for char in stripped):
        stripped = "".join(
            char if _drawable(char) else _spelled(char) for char in stripped
        )
    if stripped == text:
        return text
    # Collapsing the gap a dropped emoji left keeps the stamp from reading as a typo;
    # it is done only on text that lost a glyph, so code indentation is untouched.
    return re.sub(r" {2,}", " ", stripped) if len(stripped) < len(text) else stripped


def _drawable(char: str) -> bool:
    """Whether the body face can set this character as it stands."""
    if has_cjk(char) or char in _SUBSCRIPTS or char in _SUPERSCRIPTS:
        # The CID face draws the first, and _shift_scripts turns the other two into
        # markup further down the pipeline. Neither is the body face's problem.
        return True
    try:
        char.encode(_PDF_ENCODING)
    except UnicodeEncodeError:
        return False
    return True


_SUBSCRIPTS = {chr(0x2080 + digit): str(digit) for digit in range(10)} | {
    "₊": "+",
    "₋": "-",
    "₌": "=",
    "₍": "(",
    "₎": ")",
    "ₙ": "n",
    "ₓ": "x",
}
_SUPERSCRIPTS = {
    "⁰": "0",
    "¹": "1",
    "²": "2",
    "³": "3",
    **{chr(0x2074 + digit): str(digit + 4) for digit in range(6)},
    "⁺": "+",
    "⁻": "-",
    "ⁿ": "n",
}
_SUBSCRIPT_RUN = re.compile(f"[{''.join(_SUBSCRIPTS)}]+")
_SUPERSCRIPT_RUN = re.compile(f"[{''.join(_SUPERSCRIPTS)}]+")


def _shift_scripts(markup: str) -> str:
    """Set Unicode sub- and superscripts as real ones instead of as filled boxes.

    Times has no glyph for U+2082, so ``Li2CO3`` written the way a chemist writes it
    came off the page as ``Li`` box ``CO`` box -- a formula rendered wrong, in a report
    about battery chemistry, where the reader has no way to recover the missing digit.
    Dropping the characters would silently change the formula, so they are translated
    into the markup reportlab sets properly instead.

    This runs after escaping, where the angle brackets it introduces are the only ones
    in the string, and it deliberately leaves U+00B2 and U+00B3 alone in ``_markup``'s
    unescaped input -- they are in WinAnsi and Times draws them.
    """
    for pattern, table, tag in (
        (_SUBSCRIPT_RUN, _SUBSCRIPTS, "sub"),
        (_SUPERSCRIPT_RUN, _SUPERSCRIPTS, "super"),
    ):
        markup = pattern.sub(
            lambda match, table=table, tag=tag: "<{0}>{1}</{0}>".format(
                tag, "".join(table[char] for char in match.group(0))
            ),
            markup,
        )
    return markup


def _markup(text: str) -> str:
    """Inline Markdown to reportlab markup, with CJK runs pinned to the CID face."""
    return _shift_scripts(inline_markup(_renderable(text), cjk_font=_CJK))


def _literal(text: str) -> str:
    """Escape verbatim text (code, titles) and keep CJK runs on the CID face."""
    return _shift_scripts(cjk_markup(escape(_renderable(text), quote=False), _CJK))


def _para(markup: str, style, klass=None):
    from reportlab.platypus import Paragraph

    return (klass or Paragraph)(markup, style)


def _dossier_meta(blocks: list) -> tuple[str, str, str]:
    """Pull title-page facts out of the compiled Markdown rather than re-deriving them.

    The title page restates the report's own front matter — the goal title, the goal
    itself and the attribution stamp — so the two can never disagree.
    """
    title = ""
    question = ""
    stamp = ""
    in_goal = False
    for block in blocks:
        if isinstance(block, Heading):
            if block.level == 1 and not title:
                title = plain_text(block.text)
            in_goal = block.level == 2 and plain_text(block.text).strip() == "Goal"
            continue
        if isinstance(block, Para):
            if in_goal and not question:
                question = plain_text(block.text).replace("\n", " ").strip()
                in_goal = False
            if not stamp and block.text.startswith("Prepared by"):
                stamp = plain_text(block.text).replace("\n", " ").strip()
    return title or _TITLE, question, stamp


_CELL_PADDING = 10.0


def _all_tables(blocks: list) -> list[Table]:
    """Every table in the block tree, including the ones nested inside ``Details``."""
    found: list[Table] = []

    def walk(items: list) -> None:
        for item in items:
            if isinstance(item, Table):
                found.append(item)
            elif isinstance(item, Details):
                walk(item.blocks)

    walk(blocks)
    return found


def _shared_column_widths(blocks: list, available: float) -> dict[tuple, list[float]]:
    """One set of column widths per table shape, measured over every table of it.

    ``_column_widths`` weights a table by its own contents, and the report sets the
    same table eight times over -- one tournament grid per idea, one review grid per
    idea. Measured separately they came out at visibly different widths: "Judge" was
    set anywhere from 56 to 85pt and "Opponent" varied by 17% across tables that are
    the same table, purely because one idea's rationales ran longer than another's. A
    reader scanning eight of these in sequence reads a moved column rule as a
    difference in the data.

    Shape is the header row, so tables that are not the same table are still measured
    on their own.
    """
    grouped: dict[tuple, list[Table]] = {}
    for table in _all_tables(blocks):
        grouped.setdefault(tuple(plain_text(cell) for cell in table.header), []).append(
            table
        )
    widths: dict[tuple, list[float]] = {}
    for shape, tables in grouped.items():
        if len(tables) < 2:
            continue
        merged = Table(
            header=list(tables[0].header),
            rows=[row for table in tables for row in table.rows],
            aligns=list(tables[0].aligns),
        )
        widths[shape] = _column_widths(merged, available)
    return widths


def _column_widths(
    block: Table, available: float, shared: dict[tuple, list[float]] | None = None
) -> list[float]:
    """Weight columns by their widest cell so identifiers do not force a squeeze.

    Weights are measured in points rather than in characters. Character counts treat
    "Discriminating" and "IIIIIIIIIIIIII" as the same width, and the report's tables
    are headed by long words in bold, so the column that most needed room was the one
    the count under-served: "Discriminating predictions" was allotted less width than
    its longest single word and Word broke the header mid-word, twice.

    Each column is then floored at its own longest unbreakable word, which is the
    narrowest it can be set without hyphenating something. Where those floors do not
    all fit, they are scaled together -- a table that cannot be set without breaking a
    word should break it in proportion rather than starve one column.

    The header row is floored whole wherever the whole row fits on one line. Without
    that the floors depend only on the body cells, so one tournament table set its
    "Judge" and "Rationale" headers on one line and the next table set the same two
    headers on two, purely because the rationales under them ran longer. A reader
    scanning eight of these tables reads a change in the header as a change in the
    table.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    settled = (shared or {}).get(tuple(plain_text(cell) for cell in block.header))
    if settled is not None and len(settled) == len(block.header):
        return list(settled)
    weights: list[float] = []
    floors: list[float] = []
    headers: list[float] = []
    for index, header in enumerate(block.header):
        cells = [plain_text(header)] + [plain_text(row[index]) for row in block.rows]
        weights.append(
            max(stringWidth(cell, _SERIF_BOLD, _CELL_FONT_SIZE) for cell in cells)
            + _CELL_PADDING
        )
        headers.append(
            stringWidth(plain_text(header), _SERIF_BOLD, _CELL_FONT_SIZE)
            + _CELL_PADDING
        )
        floors.append(
            max(
                (
                    stringWidth(word, _SERIF_BOLD, _CELL_FONT_SIZE)
                    for cell in cells
                    for word in cell.split()
                ),
                default=0.0,
            )
            + _CELL_PADDING
        )
    if sum(headers) <= available:
        floors = [max(floor, head) for floor, head in zip(floors, headers, strict=True)]
    total = sum(weights) or 1.0
    scaled = [available * weight / total for weight in weights]
    if sum(floors) > available:
        crowding = available / sum(floors)
        return [floor * crowding for floor in floors]
    # Give every starved column its floor and take the difference from the slack in
    # the columns that have room, in proportion to how much slack each of them has.
    debt = sum(
        max(floor - width, 0.0) for floor, width in zip(floors, scaled, strict=True)
    )
    slack = (
        sum(
            max(width - floor, 0.0) for floor, width in zip(floors, scaled, strict=True)
        )
        or 1.0
    )
    return [
        floor if width < floor else width - (width - floor) * debt / slack
        for floor, width in zip(floors, scaled, strict=True)
    ]


def _table_flowable(
    block: Table,
    styles: dict,
    available: float,
    shared: dict[tuple, list[float]] | None = None,
):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import Table as PlatypusTable
    from reportlab.platypus import TableStyle

    alignments = {"left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT}
    cell_styles = {
        align: styles["table_cell"].clone(f"DossierTableCell{align}", alignment=code)
        for align, code in alignments.items()
    }
    header_styles = {
        align: styles["table_header"].clone(f"DossierTableHead{align}", alignment=code)
        for align, code in alignments.items()
    }
    data = [
        [
            _para(_markup(cell), header_styles[block.aligns[index]])
            for index, cell in enumerate(block.header)
        ]
    ]
    for row in block.rows:
        data.append(
            [
                _para(_markup(cell), cell_styles[block.aligns[index]])
                for index, cell in enumerate(row)
            ]
        )
    table = PlatypusTable(
        data,
        colWidths=_column_widths(block, available, shared),
        repeatRows=1,
        splitByRow=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eaed")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9aa0a6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#fafafa")],
                ),
            ]
        )
    )
    return table


def _list_flowables(block: ListBlock, styles: dict) -> list:
    counters: dict[int, int] = {}
    flowables = []
    for item in block.items:
        if item.ordered:
            counters[item.level] = counters.get(item.level, 0) + 1
            bullet = f"{counters[item.level]}."
        else:
            counters.pop(item.level, None)
            bullet = _BULLETS[item.level % len(_BULLETS)]
        for deeper in [key for key in counters if key > item.level]:
            counters.pop(deeper)
        indent = 16 + 16 * item.level
        style = styles["body"].clone(
            f"DossierListL{item.level}{'O' if item.ordered else 'B'}",
            leftIndent=indent,
            bulletIndent=indent - 14,
            spaceAfter=2.5,
            # Left unset, reportlab draws bullets in Helvetica against a Times page.
            # Ordered markers stay in the body serif so the numbers match the text
            # they label; unordered markers come from the TrueType face registered
            # above, which is the only way the glyph survives text extraction (see
            # _BULLET_FACE). Its dot is the heavier of the two, so it sets smaller.
            bulletFontName=_SERIF if item.ordered else _BULLET_FACE,
            bulletFontSize=styles["body"].fontSize * (1.0 if item.ordered else 0.9),
        )
        paragraph = _para(_markup(item.text), style)
        paragraph.bulletText = bullet
        flowables.append(paragraph)
    if flowables:
        from reportlab.platypus import Spacer

        flowables.append(Spacer(1, 5))
    return flowables


def _hard_wrap_code(text: str, font_size: float, available: float) -> str:
    """Pre-wrap to the frame width: Preformatted never wraps, and JSON lines overflow."""
    from reportlab.pdfbase.pdfmetrics import stringWidth

    advance = stringWidth("0", _MONO, font_size) or font_size * 0.6
    columns = max(int(available / advance), 24)
    wrapped: list[str] = []
    for line in (text or " ").split("\n"):
        stripped = line.lstrip(" ")
        indent = line[: len(line) - len(stripped)]
        if not stripped:
            wrapped.append("")
            continue
        wrapper = textwrap.TextWrapper(
            width=columns,
            initial_indent=indent,
            subsequent_indent=indent + "    ",
            break_long_words=True,
            break_on_hyphens=False,
            replace_whitespace=False,
            drop_whitespace=True,
        )
        wrapped.extend(wrapper.wrap(stripped) or [indent])
    return "\n".join(wrapped)


def _code_flowable(text: str, styles: dict, available: float, appendix: bool = False):
    from reportlab.platypus import XPreformatted

    style = styles["appendix_code" if appendix else "code"]
    inner = available - style.leftIndent - style.rightIndent - 2 * style.borderPadding
    body = _hard_wrap_code(text, style.fontSize, inner)
    return _para(_literal(body), style, klass=XPreformatted)


def _code_or_diagram(block: Code, styles: dict, available: float) -> list:
    """A fenced diagram is drawn; anything else fenced is set as the code it is.

    The caption ``number_figures_and_tables`` writes under the fence already calls
    this a figure, and for four exports it was a figure of eleven lines of Mermaid.
    A source the parser does not recognise still prints verbatim, which is the only
    honest fallback: a diagram half-drawn is worse than one not drawn.
    """
    from reportlab.platypus import Spacer

    if block.language.strip().lower() == "mermaid":
        drawing = flowchart_drawing(block.text, available)
        if drawing is not None:
            drawing.hAlign = "CENTER"
            return [Spacer(1, 6), drawing, Spacer(1, 4)]
    return [_code_flowable(block.text, styles, available)]


def _quote_flowable(block: Quote, styles: dict, available: float):
    from reportlab.lib import colors
    from reportlab.platypus import Table as PlatypusTable
    from reportlab.platypus import TableStyle

    table = PlatypusTable(
        [[_para(_markup(block.text), styles["quote"])]],
        colWidths=[available],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("LINEBEFORE", (0, 0), (0, -1), 2, colors.HexColor("#9aa0a6")),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


# Roughly a third of the A4 text column at 10.5/14. A closing section under this is
# moved whole rather than left to spill; above it, moving the section would cost more
# white on the page it left than it saves on the page it lands on.
_TAIL_MAX_LINES = 22
_CHARS_PER_LINE = 95


def _estimated_lines(blocks: list) -> int:
    """A rough line count for a run of prose, used only to decide whether it can move."""
    total = 0
    for block in blocks:
        if isinstance(block, Heading):
            total += 2
        elif isinstance(block, Para):
            total += max(1, -(-len(plain_text(block.text)) // _CHARS_PER_LINE)) + 1
        elif isinstance(block, ListBlock):
            total += sum(
                max(1, -(-len(plain_text(item.text)) // _CHARS_PER_LINE))
                for item in block.items
            )
        else:
            total += _TAIL_MAX_LINES + 1  # tables and code: never move the group
    return total


def _chapter_tails(blocks: list) -> dict[int, int]:
    """Where each chapter's short closing section starts and ends.

    A chapter ends on a forced page break, so whatever is left over when the previous
    page fills becomes the whole of the chapter's last page. On the live report that
    put two bullets and nine inches of white on page 18 of 80, three references on page
    20 and three criteria on page 5 -- pages a reader flipping through reads as a
    rendering failure. Keeping the closing section together moves its heading down with
    it, so the short page is a heading and the list under it: still short, but a
    deliberate-looking unit rather than an orphaned fragment.

    Only prose sections qualify. A section holding a table or a code box is left where
    it falls, because those cannot be re-flowed and a group that will not fit anywhere
    is a group reportlab has to split regardless.
    """
    starts = [
        index
        for index, block in enumerate(blocks)
        if isinstance(block, Heading) and block.level == 1
    ]
    tails: dict[int, int] = {}
    if not starts:
        return tails
    for start, end in zip(starts, [*starts[1:], len(blocks)], strict=True):
        headings = [
            index
            for index in range(start + 1, end)
            if isinstance(blocks[index], Heading) and blocks[index].level >= 2
        ]
        # The first entry is skipped: a chapter with one section has no closing
        # section distinct from the chapter, and moving it would move the chapter.
        if len(headings) < 2:
            continue
        head = headings[-1]
        if _estimated_lines(blocks[head:end]) <= _TAIL_MAX_LINES:
            tails[head] = end
    return tails


_CAPTION = re.compile(r"^(?:Table|Figure) \d+\.")


def _captions(block) -> bool:
    """Whether this paragraph is the numbered caption for the exhibit beside it."""
    return isinstance(block, Para) and bool(
        _CAPTION.match(plain_text(block.text).strip())
    )


def _introduces(block, following) -> bool:
    """Whether this paragraph is the sentence that announces the block under it.

    "Evidence gaps:" was the last line of page 51 and the four gaps it announces
    were on page 52; "Discovery recorded two qualifications on it:" ended page 7 the
    same way. A colon is a promise about the next line, so it cannot be the last
    thing on a leaf.
    """
    if not isinstance(block, Para) or not plain_text(block.text).rstrip().endswith(":"):
        return False
    return isinstance(following, ListBlock | Table)


_EXHIBIT_INDEX_TITLE = FIGURE_INDEX_HEADING.lstrip("# ").strip()


def _exhibit_index_lists(blocks: list) -> set[int]:
    """Where the back-matter index of exhibits is written out as a plain list."""
    return {
        index + 1
        for index, block in enumerate(blocks[:-1])
        if isinstance(block, Heading)
        and plain_text(block.text).strip() == _EXHIBIT_INDEX_TITLE
        and isinstance(blocks[index + 1], ListBlock)
    }


def _exhibit_wording(block: ListBlock) -> dict[str, str]:
    """Each exhibit's index entry, keyed by its label and kept in the order written."""
    wording: dict[str, str] = {}
    for item in block.items:
        text = " ".join(plain_text(item.text).split())
        label = _CAPTION.match(f"{text.partition(' — ')[0]}.")
        if label:
            wording[label.group(0).rstrip(".")] = text
    return wording


def _exhibit_index(block: ListBlock, styles: dict):
    """The index of exhibits, set with the page each exhibit is actually on.

    The Markdown index locates its entries by anchor, which is what a Markdown reader
    follows and what a PDF has no notion of. Flattened for export, all nineteen entries
    printed as a list of names over two pages with no page number against any of them
    and nothing to click -- an index of figures and tables that locates neither.

    The entries keep the Markdown's own wording and order so the two indexes are the
    same index; only the page numbers, which only the typeset copy has, are added.
    """
    from reportlab.platypus.tableofcontents import TableOfContents

    class _Index(TableOfContents):
        def __init__(self, wording: dict[str, str]):
            super().__init__(notifyKind="ExhibitEntry")
            self._wording = wording
            self._found: dict[str, tuple[int, str | None]] = {}
            self.levelStyles = [styles["toc"][1]]
            self.dotsMinLevel = 0

        def beforeBuild(self):
            super().beforeBuild()
            self._found = {}

        def addEntry(self, level, text, pageNum, key=None):
            # Captions are notified in the order they are typeset, which interleaves
            # the figures with the tables; the index is written figures first.
            self._found[text] = (pageNum, key)
            self._entries = [
                (0, _literal(wording), *self._found[label])
                for label, wording in self._wording.items()
                if label in self._found
            ]

    return _Index(_exhibit_wording(block))


def _keep_a_heading_with_what_it_heads(story: list) -> list:
    """Bind a heading past the space under it to the thing that space is above.

    reportlab binds a heading to the one flowable that follows it, and for every
    diagram in the report that flowable is the six-point spacer the drawing is padded
    with -- so the bind was satisfied by the padding and four "Proposed Workflow"
    headings printed as the last line of a page with their flowchart overleaf.
    """
    from reportlab.platypus import Spacer

    for index, flowable in enumerate(story):
        if not flowable.getKeepWithNext():
            continue
        for follower in story[index + 1 :]:
            if not isinstance(follower, Spacer):
                break
            follower.__dict__["keepWithNext"] = 1
    return story


def _story_from_blocks(
    blocks: list,
    styles: dict,
    available: float,
    cover_title: str = "",
    shared: dict[tuple, list[float]] | None = None,
) -> list:
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, KeepTogether, PageBreak, Spacer

    indexes = _exhibit_index_lists(blocks)
    # reportlab only registers an indexing flowable it finds at the top of the story, so
    # an index folded into a KeepTogether group is never asked for its entries and prints
    # "Placeholder for table of contents" against page 0. A short index -- a report with
    # two exhibits -- is exactly the closing section that qualifies as a chapter tail.
    tails = {
        head: end
        for head, end in _chapter_tails(blocks).items()
        if not any(head <= entry < end for entry in indexes)
    }
    story: list = []
    index = -1
    skip_until = 0
    for block in blocks:
        index += 1
        if index < skip_until:
            continue
        following = blocks[index + 1] if index + 1 < len(blocks) else None
        if index in indexes:
            story.append(_exhibit_index(block, styles))
            continue
        if index in tails:
            skip_until = tails[index]
            story.append(
                KeepTogether(
                    _story_from_blocks(
                        blocks[index:skip_until], styles, available, shared=shared
                    )
                )
            )
            continue
        if isinstance(block, Heading):
            if block.level == 1:
                # The document's own first H1 is the goal, which the title page has
                # already set at 26pt one leaf earlier. Printing it again opened the
                # body on a chapter heading with nothing under it, and the next H1
                # broke the page immediately -- so page four of every export was a
                # duplicated title and four inches of white.
                if index == 0 and plain_text(block.text).strip() == cover_title.strip():
                    continue
                # A chapter starts a page, but only if something is on the current one.
                if story and not isinstance(story[-1], PageBreak):
                    story.append(PageBreak())
            story.append(_para(_markup(block.text), styles[f"h{block.level}"]))
        elif isinstance(block, Para):
            paragraph = _para(_markup(block.text), styles["body"])
            if _introduces(block, following) or _captions(block):
                paragraph.keepWithNext = 1
            story.append(paragraph)
        elif isinstance(block, ListBlock):
            story.extend(_list_flowables(block, styles))
        elif isinstance(block, Table):
            story.extend(
                [_table_flowable(block, styles, available, shared), Spacer(1, 9)]
            )
        elif isinstance(block, Code):
            story.extend(_code_or_diagram(block, styles, available))
        elif isinstance(block, Quote):
            story.append(_quote_flowable(block, styles, available))
        elif isinstance(block, Rule):
            story.append(
                HRFlowable(
                    width="100%",
                    thickness=0.6,
                    color=colors.HexColor("#dadce0"),
                    spaceBefore=8,
                    spaceAfter=10,
                )
            )
        elif isinstance(block, Details):
            story.append(_para(_markup(block.summary), styles["summary"]))
            for nested in block.blocks:
                if isinstance(nested, Code):
                    story.append(
                        _code_flowable(nested.text, styles, available, appendix=True)
                    )
                else:
                    story.extend(
                        _story_from_blocks([nested], styles, available, shared=shared)
                    )
    return _keep_a_heading_with_what_it_heads(story)


def _make_doc_template(buffer, title: str, header: str, totals: dict):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    width, height = A4
    header_font = _CJK if has_cjk(header) else _SERIF

    class _DossierDocTemplate(SimpleDocTemplate):
        """Emits TOC entries and PDF bookmarks so multiBuild resolves real page numbers."""

        def handle_documentBegin(self):
            super().handle_documentBegin()
            self._entry_index = 0
            self._outline_depth = -1

        def _endBuild(self):
            # Without this the catalog's /PageMode stays /UseNone: the outline is
            # built and written and no reader shows it unless the person opening
            # the file knows to go looking for the bookmarks pane. A forty-page
            # report with a hundred-odd destinations opens as an undifferentiated
            # scroll.
            self.canv.showOutline()
            super()._endBuild()

        def afterPage(self):
            """Running head and folio, drawn once per page as that page is closed.

            This used to be deferred to a canvas that snapshotted every page's state
            and replayed it at save time -- the widely copied reportlab recipe for
            "Page n of m". The recipe's showPage() calls _startPage() directly
            instead of Canvas.showPage(), so the document's own page counter never
            advances during the build, and bookmarkPage() below binds every
            destination to page one. In a report with a hundred-odd contents entries
            and cross-references, all of them jumped to the title page. The total
            comes from the previous multiBuild pass instead (see _PageTotal), which
            costs one extra pass and keeps the page counter honest.
            """
            page = self.page
            if page <= 1:  # The title page carries no furniture.
                totals["counted"] = max(totals.get("counted", 0), page)
                return
            canvas = self.canv
            canvas.saveState()
            canvas.setFont(header_font, 8)
            canvas.setFillColor(colors.HexColor("#80868b"))
            canvas.drawCentredString(width / 2.0, height - _MARGIN + 14, header)
            canvas.setStrokeColor(colors.HexColor("#dadce0"))
            canvas.setLineWidth(0.4)
            canvas.line(
                _MARGIN, height - _MARGIN + 8, width - _MARGIN, height - _MARGIN + 8
            )
            canvas.setFont(_SERIF, 8.5)
            total = totals.get("total") or page
            canvas.drawCentredString(
                width / 2.0, _MARGIN - 22, f"Page {page} of {total}"
            )
            canvas.restoreState()
            totals["counted"] = max(totals.get("counted", 0), page)

        def afterFlowable(self, flowable):
            if not isinstance(flowable, Paragraph):
                return
            level = _TOC_LEVELS.get(flowable.style.name.removesuffix("CJK"))
            if level is None:
                # A caption is body type, so it falls through the contents machinery.
                # The back-matter index needs the page it landed on all the same.
                label = _CAPTION.match(flowable.getPlainText().strip())
                if label:
                    key = f"exhibit-{label.group(0).rstrip('.').lower()}"
                    self.canv.bookmarkPage(key)
                    self.notify(
                        "ExhibitEntry", (0, label.group(0).rstrip("."), self.page, key)
                    )
                return
            text = flowable.getPlainText()
            key = f"toc-{self._entry_index}"
            self._entry_index += 1
            self.canv.bookmarkPage(key)
            # A document that opens on a subheading -- because its own H1 duplicated
            # the title page and was dropped, or because it simply starts at H2 --
            # asks the outline to jump from nothing to depth one, which reportlab
            # refuses outright rather than flattening. The bookmark tree only needs
            # to nest consistently, so a level that skips its parent is pulled up to
            # sit beside it. The Contents page keeps the level the heading declared.
            self._outline_depth = min(level, self._outline_depth + 1)
            self.canv.addOutlineEntry(
                text[:120], key, self._outline_depth, closed=(self._outline_depth > 0)
            )
            # TableOfContents splices the label into paragraph markup verbatim, so it
            # must arrive escaped and with the CID face pinned for CJK headings.
            label = _literal(text)
            self.notify("TOCEntry", (level, label, self.page, key))

    return _DossierDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=_MARGIN,
        rightMargin=_MARGIN,
        topMargin=_MARGIN + 8,
        bottomMargin=_MARGIN,
        title=title,
        author="AI co-scientist",
        subject="Scientific research planning dossier",
    )


def _page_total_flowable(totals: dict):
    """A zero-size indexing flowable that makes multiBuild settle the page total.

    The folio has to print "of M" before the build that determines M has finished.
    multiBuild already re-runs the story until every indexing flowable declares
    itself satisfied, so the page count is registered as one more thing to
    converge: each pass records how many pages it produced, and the pass after it
    prints that number. A document whose total shifts because the folio widened
    simply takes one more pass, exactly as the table of contents does.
    """
    from reportlab.platypus.doctemplate import IndexingFlowable

    class _PageTotal(IndexingFlowable):
        width = 0
        height = 0

        def beforeBuild(self) -> None:
            totals["counted"] = 0

        def afterBuild(self) -> None:
            totals["previous"] = totals.get("total")
            totals["total"] = totals.get("counted", 0)

        def isSatisfied(self) -> int:
            return int(totals.get("total") == totals.get("previous"))

        def wrap(self, *args):
            return (0, 0)

        def draw(self) -> None:
            pass

    return _PageTotal()


def _title_page(title: str, question: str, stamp: str, styles: dict) -> list:
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, PageBreak, Spacer

    story = [
        Spacer(1, 150),
        _para(_literal(title), styles["title"]),
        Spacer(1, 14),
        HRFlowable(width="45%", thickness=1, color=colors.HexColor("#9aa0a6")),
        Spacer(1, 20),
    ]
    if question:
        story.extend(
            [
                _para(_literal(question), styles["subtitle"]),
                Spacer(1, 26),
            ]
        )
    story.extend(
        [
            _para(
                _literal(
                    stamp
                    or "Prepared by AI co-scientist on "
                    f"{date.today().isoformat()}. For research purposes only."
                ),
                styles["title_meta"],
            ),
            Spacer(1, 22),
            _para(_markup(_DEFAULT_NOTICE), styles["notice"]),
            PageBreak(),
        ]
    )
    return story


def render_pdf(content: str) -> bytes:
    """Render the dossier as a typeset PDF: title page, live TOC, tables, page numbers."""
    try:
        from reportlab.platypus import PageBreak
        from reportlab.platypus.tableofcontents import TableOfContents
    except ImportError as exc:
        raise RuntimeError("PDF export requires the reportlab dependency.") from exc

    _register_pdf_fonts()
    styles = _pdf_styles()
    # Both exporters build a native contents list and number the exhibits
    # themselves, so the Markdown one is dropped rather than set twice; the
    # numbering pass is idempotent and only fires on markup compiled elsewhere.
    blocks = parse_blocks(
        flatten_fragment_links(
            strip_table_of_contents(
                _without_cover_notice(number_figures_and_tables(content))
            )
        )
    )
    title, question, stamp = _dossier_meta(blocks)

    buffer = BytesIO()
    header = title if len(title) <= 96 else f"{title[:93].rstrip()}…"
    totals: dict[str, int | None] = {}
    document = _make_doc_template(buffer, title, header, totals)

    toc = TableOfContents()
    toc.levelStyles = styles["toc"]
    toc.dotsMinLevel = 0

    story = _title_page(title, question, stamp, styles)
    measure = document.width - 2 * _FRAME_PADDING
    story.extend(
        [
            _para("Contents", styles["toc_title"]),
            toc,
            PageBreak(),
            *_story_from_blocks(
                blocks,
                styles,
                measure,
                title,
                _shared_column_widths(blocks, measure),
            ),
            _page_total_flowable(totals),
        ]
    )

    document.multiBuild(story)
    return buffer.getvalue()


def _plain_markdown(value: str) -> str:
    """Keep dossier text readable in word processors without raw Markdown marks."""
    return plain_text(value)


def _docx_hyperlink(paragraph, href: str):
    """A real ``w:hyperlink`` so the exported link is clickable, not merely underlined.

    python-docx has no API for this, so the relationship and the element are built by
    hand. Without it the DOCX carried the link text underlined in blue and nothing
    else: every one of the report's references looked like a link and none of them
    was one, which is worse than plain text, because the reader tries clicking.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    relationship = paragraph.part.relate_to(
        href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    element = OxmlElement("w:hyperlink")
    element.set(qn("r:id"), relationship)
    paragraph._p.append(element)
    return element


def _docx_add_runs(
    paragraph, text: str, force_bold: bool = False, size: float | None = None
) -> None:
    """Turn inline Markdown into real Word runs instead of flattening the marks away."""
    from docx.shared import Pt, RGBColor

    for span in parse_inline(text):
        anchor = _docx_hyperlink(paragraph, span.href) if span.href else None
        for index, part in enumerate(span.text.split("\n")):
            if index:
                paragraph.add_run().add_break()
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = span.bold or force_bold or None
            run.italic = span.italic or None
            if size is not None:
                run.font.size = Pt(size)
            if span.code:
                run.font.name = "Courier New"
                run.font.size = Pt(min(9.0, size) if size is not None else 9.0)
            if anchor is not None:
                run.underline = True
                run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
                # Reparent the run under the hyperlink: a w:hyperlink with no w:r in
                # it is an empty anchor, and the text would sit outside the link.
                anchor.append(run._r)


def _docx_shade(cell, fill: str) -> None:
    """Solid cell fill, which python-docx exposes no API for."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _docx_repeat_header(row) -> None:
    """Mark a table's first row as a header row so Word repeats it across pages."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(marker)


def _docx_add_table(
    document,
    block: Table,
    available: float,
    shared: dict[tuple, list[float]] | None = None,
) -> None:
    """The same table the PDF sets: repeating header, shaded, weighted columns.

    The PDF gets a repeated header row, a grey header band and columns weighted by
    content. The DOCX got none of the three: the report's longest tables ran over a
    page boundary with their header left behind on the previous one, and every column
    was set to the same width, so a rank number and a paragraph of rationale were
    allotted the same room.

    Two things have to agree for that to hold. The cells are set at the same size the
    widths were measured at: they used to inherit Normal's 11pt against widths computed
    for 8.5pt type, so "Discriminating", "Round" and "Result" overflowed their columns
    and Word broke them mid-word -- the exact defect ``_column_widths`` was written to
    prevent, fixed in the PDF only. And the widths are written into ``w:tblGrid`` as
    well as into each ``w:tcW``: Word resolves in favour of the cell widths, but
    LibreOffice and Google Docs lay out from the grid, and python-docx leaves the grid
    at equal columns, so in those two readers every table fell back to equal columns.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt

    table = document.add_table(rows=1, cols=len(block.header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    points = _column_widths(block, available, shared)
    widths = [Pt(width) for width in points]
    _docx_set_grid(table, points)
    header_row = table.rows[0]
    _docx_repeat_header(header_row)
    for index, cell_text in enumerate(block.header):
        cell = header_row.cells[index]
        _docx_add_runs(cell.paragraphs[0], cell_text, True, size=_CELL_FONT_SIZE)
        _docx_shade(cell, "E8EAED")
        cell.width = widths[index]
    for row in block.rows:
        cells = table.add_row().cells
        for index, cell_text in enumerate(row):
            _docx_add_runs(cells[index].paragraphs[0], cell_text, size=_CELL_FONT_SIZE)
            cells[index].width = widths[index]
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:  # keep the grid last-written, after add_row() cloning
        _docx_set_grid(table, points)


def _docx_set_grid(table, points: list[float]) -> None:
    """Write the measured column widths into ``w:tblGrid`` and ``w:tblW``.

    python-docx emits one ``w:gridCol`` per column at an equal share of the page and
    an ``auto`` table width, both of which contradict the per-cell widths beside them.
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    twips = [str(round(width * 20)) for width in points]
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    for column in list(grid):
        grid.remove(column)
    for width in twips:
        element = OxmlElement("w:gridCol")
        element.set(qn("w:w"), width)
        grid.append(element)
    # Mutated in place rather than replaced: ``w:tblPr`` is a sequence, and an appended
    # ``w:tblW`` lands after ``w:tblLayout``, which is out of schema order.
    total = table._tbl.tblPr.find(qn("w:tblW"))
    if total is not None:
        total.set(qn("w:w"), str(sum(int(width) for width in twips)))
        total.set(qn("w:type"), "dxa")


def _docx_add_code(document, text: str, size: float) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph(style="No Spacing")
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)


def _docx_field(paragraph, instruction: str) -> None:
    """A Word field code, which Word evaluates on open and python-docx cannot write."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    for kind, text in (("begin", None), (None, instruction), ("end", None)):
        if kind is None:
            element = OxmlElement("w:instrText")
            element.set(qn("xml:space"), "preserve")
            element.text = text
        else:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        paragraph.add_run()._r.append(element)


_DOCX_TOC_DEPTH = 3


def _docx_toc_bookmarks(blocks: list) -> dict[int, str]:
    """A bookmark name for every heading the contents field covers, in reading order.

    Keyed on object identity so ``_docx_blocks`` can look each heading up as it walks
    the same tree, and the anchor the contents entry points at is the anchor the
    heading carries.
    """
    names: dict[int, str] = {}

    def walk(items: list) -> None:
        for item in items:
            if isinstance(item, Heading) and item.level <= _DOCX_TOC_DEPTH:
                names[id(item)] = f"_Toc_dossier{len(names)}"
            elif isinstance(item, Details):
                walk(item.blocks)

    walk(blocks)
    return names


def _docx_toc_entries(
    blocks: list, bookmarks: dict[int, str]
) -> list[tuple[int, str, str]]:
    """``(level, text, bookmark)`` for the contents list, in reading order."""
    entries: list[tuple[int, str, str]] = []

    def walk(items: list) -> None:
        for item in items:
            if isinstance(item, Heading) and id(item) in bookmarks:
                entries.append(
                    (item.level, plain_text(item.text).strip(), bookmarks[id(item)])
                )
            elif isinstance(item, Details):
                walk(item.blocks)

    walk(blocks)
    return entries


def _docx_bookmark(paragraph, name: str) -> None:
    """Wrap a heading in a ``w:bookmarkStart``/``w:bookmarkEnd`` pair for the TOC."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    identifier = str(abs(hash(name)) % 1_000_000)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), identifier)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), identifier)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _docx_anchor(paragraph, name: str):
    """An internal ``w:hyperlink``, which jumps to a bookmark rather than to a URL."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    element = OxmlElement("w:hyperlink")
    element.set(qn("w:anchor"), name)
    paragraph._p.append(element)
    return element


def _docx_update_fields(document) -> None:
    """Ask Word to evaluate every field on open, so the contents list builds itself."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    flag = OxmlElement("w:updateFields")
    flag.set(qn("w:val"), "true")
    settings.append(flag)


def _docx_cover(
    document,
    title: str,
    question: str,
    stamp: str,
    entries: list[tuple[int, str, str]] | None = None,
) -> None:
    """The same title page the PDF gets, followed by a Word-native contents field.

    The DOCX had neither. It opened straight onto the first heading, so the two
    exports of one report did not agree on what the document was -- and with no
    contents field a hundred-page dossier had no way in but the scroll bar.

    The field is written with a cached result rather than empty. Word evaluates a TOC
    field and would have built the list itself; Google Docs and LibreOffice do not
    evaluate fields on import at all, and with nothing cached to fall back on an
    eighty-page report opened in Docs -- the reader this export exists for -- had no
    contents list, permanently, under a note telling it to press a key that does
    nothing there. The cached entries are real internal hyperlinks, so the list works
    before any field is updated and Word replaces it with a paginated one on open.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Pt, RGBColor

    def centred(text: str, size: float, *, bold=False, italic=False, grey=False):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold or None
        run.italic = italic or None
        if grey:
            run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
        return paragraph

    document.add_paragraph().paragraph_format.space_after = Pt(120)
    centred(title, 26, bold=True)
    if question:
        centred(question, 13)
    centred(stamp, 10, grey=True).paragraph_format.space_before = Pt(24)
    centred(_DEFAULT_NOTICE, 9.5, italic=True, grey=True)

    # Not a heading: as Heading 2 under a TOC field scoped to levels 1-3, the word
    # "Contents" was the first entry in the contents list it captioned.
    contents = document.add_paragraph()
    contents.paragraph_format.page_break_before = True
    contents.paragraph_format.space_after = Pt(10)
    caption = contents.add_run("Contents")
    caption.bold = True
    caption.font.size = Pt(_DOCX_HEADING_SIZES[2])
    _docx_contents_field(document, entries or [])
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _docx_contents_field(document, entries: list[tuple[int, str, str]]) -> None:
    """A ``TOC`` field whose cached result is a working list of internal links."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
    from docx.shared import Pt, RGBColor

    opener = document.add_paragraph()
    for kind, text in (
        ("begin", None),
        (None, rf'TOC \o "1-{_DOCX_TOC_DEPTH}" \h \z \u'),
    ):
        if kind is None:
            element = OxmlElement("w:instrText")
            element.set(qn("xml:space"), "preserve")
            element.text = text
        else:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        opener.add_run()._r.append(element)
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    opener.add_run()._r.append(separator)

    for level, text, bookmark in entries:
        line = document.add_paragraph()
        line.paragraph_format.left_indent = Pt(14 * (level - 1))
        line.paragraph_format.space_after = Pt(2)
        anchor = _docx_anchor(line, bookmark)
        run = line.add_run(text)
        run.font.size = Pt(10.5)
        run.bold = level == 1 or None
        run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
        anchor.append(run._r)

    closer = document.add_paragraph()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    closer.add_run()._r.append(end)


_DOCX_QUOTE_STYLE = "Dossier Quote"


def _docx_quote_style(document):
    """The PDF's quote: grey italic behind a grey left rule, indented on one side.

    Word's built-in "Intense Quote" is a decorative pull-quote -- bold italic in accent
    blue, indented both sides, with a bottom border. The blocks set in this style are
    the verbatim fatal-flaw text and the adjudicator's justification, reprinted word for
    word so a reader can judge a safety override for themselves. Setting those as a
    pull-quote reads as emphasis rather than as citation, and it made the two exports of
    the same governance record look like different documents.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
    from docx.shared import Pt, RGBColor

    if _DOCX_QUOTE_STYLE in document.styles:
        return document.styles[_DOCX_QUOTE_STYLE]
    style = document.styles.add_style(_DOCX_QUOTE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles["Normal"]
    style.font.italic = True
    style.font.color.rgb = RGBColor(0x3C, 0x40, 0x43)
    _docx_drop_theme_fonts(style)
    # The border goes in before the indent and the spacing: ``w:pPr`` is a sequence in
    # which ``w:pBdr`` precedes both, and python-docx inserts those two after whatever
    # is already there.
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")  # eighths of a point, so 1.5pt
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "9AA0A6")
    borders.append(left)
    style.element.get_or_add_pPr().append(borders)
    style.paragraph_format.left_indent = Pt(18)
    style.paragraph_format.space_after = Pt(6)
    return style


def _docx_drop_theme_fonts(style) -> None:
    """Clear the theme font references so an explicit typeface actually applies.

    ``Font.name`` writes ``w:rFonts/@w:ascii``, and Word's template also carries
    ``@w:asciiTheme`` on the same element. ECMA-376 gives the theme attribute
    precedence, and the theme's major font is Calibri, so every one of the report's
    280-odd headings came out sans-serif over a serif body however the name was set.
    """
    from docx.oxml.ns import qn

    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        fonts.attrib.pop(qn(attribute), None)


def _docx_match_complex_size(style, size: float) -> None:
    """Set ``w:szCs`` alongside ``w:sz``, which python-docx leaves at the template's."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    properties = style.element.get_or_add_rPr()
    complex_size = properties.find(qn("w:szCs"))
    if complex_size is None:
        complex_size = OxmlElement("w:szCs")
        properties.append(complex_size)
    complex_size.set(qn("w:val"), str(round(size * 2)))


def _docx_add_page_number_footer(document) -> None:
    """Insert a PAGE field so Word/Docs paginate the exported dossier itself."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    _docx_field(paragraph, "PAGE")
    paragraph.add_run(" of ")
    _docx_field(paragraph, "NUMPAGES")


def _docx_own_numbering(document, style) -> int | None:
    """A numbering instance private to one ordered list, so it starts again at one.

    Every "List Number" paragraph python-docx creates points at the numbering
    instance the style carries, and Word counts an instance across the whole
    document. The report holds dozens of ordered lists, so they came out as a single
    running sequence: the reference list ended at 24 and the next numbered list in
    the document opened at 25. A new ``w:num`` against the same abstract definition
    keeps the format and resets the count.
    """
    from docx.oxml.ns import qn

    properties = style.element.find(qn("w:pPr"))
    source = properties.find(qn("w:numPr")) if properties is not None else None
    declared = source.find(qn("w:numId")) if source is not None else None
    if declared is None:
        return None
    numbering = document.part.numbering_part.element
    try:
        template = numbering.num_having_numId(int(declared.get(qn("w:val"))))
    except KeyError:
        return None
    fresh = numbering.add_num(template.abstractNumId.val)
    fresh.add_lvlOverride(ilvl=0).add_startOverride(1)
    return fresh.numId


def _docx_blocks(
    document,
    blocks: list,
    available: float,
    code_size: float = 8.0,
    *,
    top_level: bool = False,
    bookmarks: dict[int, str] | None = None,
    shared: dict[tuple, list[float]] | None = None,
) -> None:
    for index, block in enumerate(blocks):
        if isinstance(block, Heading):
            heading = document.add_heading("", level=block.level)
            _docx_add_runs(heading, block.text)
            anchor = (bookmarks or {}).get(id(block))
            if anchor:
                _docx_bookmark(heading, anchor)
            # A chapter starts a page in the PDF. In the DOCX it did not, so the two
            # exports of one report disagreed about where the chapters were: Provenance
            # opened four lines under the last idea in Word and on its own page in the
            # PDF. The first body chapter is exempt because the contents page has just
            # broken for it.
            if top_level and block.level == 1 and index:
                heading.paragraph_format.page_break_before = True
        elif isinstance(block, Para):
            _docx_add_runs(document.add_paragraph(), block.text)
        elif isinstance(block, ListBlock):
            own: dict[str, int | None] = {}
            for item in block.items:
                style = "List Number" if item.ordered else "List Bullet"
                if item.level:
                    style = f"{style} {min(item.level + 1, 3)}"
                paragraph = document.add_paragraph(style=style)
                if item.ordered:
                    # One instance per style per list: the nested levels of a single
                    # list each count separately, and both restart with the list.
                    if style not in own:
                        own[style] = _docx_own_numbering(document, paragraph.style)
                    if own[style] is not None:
                        numbering = paragraph._p.get_or_add_pPr().get_or_add_numPr()
                        numbering.get_or_add_numId().val = own[style]
                _docx_add_runs(paragraph, item.text)
        elif isinstance(block, Table):
            _docx_add_table(document, block, available, shared)
            document.add_paragraph()
        elif isinstance(block, Code):
            steps = (
                flowchart_steps(block.text)
                if block.language.strip().lower() == "mermaid"
                else None
            )
            if steps:
                # Word holds an image or nothing, and nothing here can rasterise the
                # drawing the PDF gets, so the arrows are read out in order. The
                # caption below still calls it a figure, which it is -- the same
                # graph, in the one form this format can carry.
                for step in steps:
                    _docx_add_runs(document.add_paragraph(style="List Bullet"), step)
            else:
                _docx_add_code(document, block.text, code_size)
        elif isinstance(block, Quote):
            paragraph = document.add_paragraph(style=_docx_quote_style(document))
            _docx_add_runs(paragraph, block.text)
        elif isinstance(block, Rule):
            _docx_rule(document)
        elif isinstance(block, Details):
            _docx_add_runs(document.add_paragraph(), block.summary, True)
            _docx_blocks(
                document,
                block.blocks,
                available,
                code_size=7.0,
                bookmarks=bookmarks,
                shared=shared,
            )


def _docx_rule(document) -> None:
    """A horizontal rule as a paragraph border, which is what the PDF draws.

    It used to be thirty literal em dashes, which is a row of text: it wraps at the
    margin, it is selectable, and it comes out a different length in every reader.
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    paragraph = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D3D7")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def render_docx(content: str) -> bytes:
    """Render an editable DOCX with real Word headings, tables, lists and runs."""
    try:
        from docx import Document
        from docx.shared import Inches, Mm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX export requires the python-docx dependency.") from exc

    document = Document()
    section = document.sections[0]
    # python-docx's template is US Letter. The PDF is A4, and the two exports of one
    # report paginated differently and handed _column_widths a different text column.
    section.page_width = Mm(_PAGE_WIDTH_MM)
    section.page_height = Mm(_PAGE_HEIGHT_MM)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Both exporters build a native contents list and number the exhibits
    # themselves, so the Markdown one is dropped rather than set twice; the
    # numbering pass is idempotent and only fires on markup compiled elsewhere.
    blocks = parse_blocks(
        flatten_fragment_links(
            strip_table_of_contents(
                _without_cover_notice(number_figures_and_tables(content))
            )
        )
    )
    title, question, stamp = _dossier_meta(blocks)
    document.core_properties.title = title
    document.core_properties.subject = "Scientific research planning dossier"
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    _docx_drop_theme_fonts(normal)
    # Word's default heading theme is a blue sans face; keep the serif hierarchy.
    for level, size in _DOCX_HEADING_SIZES.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        style.font.bold = True
        # Heading 4 and Heading 6 are italic in Word's template, and Heading 4 is the
        # report's most-used rank -- every idea carries eight of them. Left alone the
        # DOCX set in bold italic what the PDF sets in bold roman.
        style.font.italic = False
        _docx_drop_theme_fonts(style)
        _docx_match_complex_size(style, size)

    _docx_add_page_number_footer(document)
    # Same reason as the PDF: the document's own first H1 is the goal, which the
    # cover page has just set. Printed again it is a chapter heading with no chapter.
    # Dropped before the contents list is built, not after, so the list does not open
    # with an entry for a heading the document no longer holds.
    if blocks and isinstance(blocks[0], Heading) and blocks[0].level == 1:
        if plain_text(blocks[0].text).strip() == title.strip():
            blocks = blocks[1:]
    bookmarks = _docx_toc_bookmarks(blocks)
    _docx_cover(
        document,
        title,
        question,
        stamp
        or f"Prepared by AI co-scientist on {date.today().isoformat()}. "
        "For research purposes only.",
        _docx_toc_entries(blocks, bookmarks),
    )
    available = (
        section.page_width - section.left_margin - section.right_margin
    ) / 12700.0  # EMU per point
    _docx_blocks(
        document,
        blocks,
        available,
        top_level=True,
        bookmarks=bookmarks,
        shared=_shared_column_widths(blocks, available),
    )
    _docx_update_fields(document)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_dossier(path: str | Path, content: str) -> None:
    """Write Markdown, PDF, or Google Docs-compatible DOCX."""
    destination = Path(path)
    suffix = destination.suffix.lower()
    if suffix == ".pdf":
        destination.write_bytes(render_pdf(content))
    elif suffix == ".docx":
        destination.write_bytes(render_docx(content))
    else:
        destination.write_text(content, encoding="utf-8")
