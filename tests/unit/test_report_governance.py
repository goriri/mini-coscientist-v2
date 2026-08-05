"""The report has to tell the truth about a human answering a governance block.

A hypothesis that a safety officer removed must not simply be missing from the
ranking, and a hypothesis that someone allowed to stand while carrying a fatal flaw
must not read like any other idea. Both are silent-omission defects, so every test
here is written against the invariant rather than against the wording: the fixtures
are built with the real ``governance`` helpers, so a change to how withdrawal or
override is recorded breaks these tests instead of quietly bypassing them.
"""

from __future__ import annotations

import re
from io import BytesIO

import pytest

from coscientist.dossier import compile_dossier, render_docx, render_pdf
from coscientist.governance import (
    governance_blockers,
    record_adjudication,
    withdraw_candidate,
)
from coscientist.models import STAGES, Artifact, CandidateReview, ReviewSet
from coscientist.narrative import load_record

FLAW = (
    "Annealing the assembled electrode at 400 C decomposes the PVDF binder and vents "
    "hydrogen fluoride into the laboratory, which is an acute inhalation hazard to "
    "the person running the cell build."
)
SECOND_FLAW = (
    "The protocol as written cycles cells above their rated voltage window, which "
    "risks lithium plating and thermal runaway on an unattended cycler."
)
THIRD_FLAW = (
    "The dry-room step vents solvent vapour into a shared corridor, which no local "
    "extraction covers and no risk assessment names."
)
ADJUDICATOR = "J. Reviewer (battery safety officer)"
SECOND_ADJUDICATOR = "P. Okonkwo (institutional biosafety chair)"
JUSTIFICATION = (
    "Confirmed with the binder supplier: PVDF decomposes above 350 C, so the anneal "
    "cannot be run on an assembled electrode. The confounder is worth testing, but "
    "only on binder-free electrodes, which is a different hypothesis."
)
SECOND_JUSTIFICATION = (
    "Accepted for a single pilot cell only, inside the vented test chamber, with the "
    "cycler attended for the full run. The upside is large enough to justify the "
    "exposure at n=1 and at no larger scale."
)


def _block(session, candidate_id: str, flaw: str):
    """Record a real fatal governance finding and hand back its blocker."""
    review = CandidateReview(
        candidate_id=candidate_id,
        criterion="safety_governance",
        reviewer="ethics_safety_governance",
        recommendation="reject",
        findings=["The hazard is intrinsic to the step rather than to its execution."],
        objections=["No engineering control was proposed for the released gas."],
        fatal_flaws=[flaw],
        confidence=0.9,
    )
    session.artifacts.append(
        Artifact(
            stage="reflect",
            agent="ethics_safety_governance",
            artifact_type="specialist_output",
            content="",
            schema_name="ReviewSet",
            payload=ReviewSet(reviews=[review]).model_dump(),
        )
    )
    return next(
        item for item in governance_blockers(session) if item.review_id == review.id
    )


def _withdraw(session, candidate_id: str, *, flaw: str = FLAW) -> str:
    blocker = _block(session, candidate_id, flaw)
    withdraw_candidate(session, candidate_id)
    record_adjudication(
        session,
        blocker,
        resolution="withdraw",
        adjudicator=ADJUDICATOR,
        justification=JUSTIFICATION,
    )
    return candidate_id


def _override(session, candidate_id: str, *, flaw: str = SECOND_FLAW) -> str:
    blocker = _block(session, candidate_id, flaw)
    record_adjudication(
        session,
        blocker,
        resolution="override",
        adjudicator=SECOND_ADJUDICATOR,
        justification=SECOND_JUSTIFICATION,
    )
    return candidate_id


def _flat(text: str) -> str:
    return " ".join(text.split())


def _population_ids(session) -> list[str]:
    """The candidate ids the report will rank, read the way the renderer reads them."""
    return [candidate.id for candidate in load_record(session).candidates]


def _idea_slots(report: str) -> list[str]:
    """The numbered subsection headings that list the run's ideas, in document order."""
    return [
        line[len("##### ") :].strip()
        for line in report.splitlines()
        if re.match(r"^##### 4\.\d+ ", line)
    ]


def _section(report: str, heading: str) -> str:
    """One heading's text, up to the next heading at the same or a shallower level."""
    level = len(heading) - len(heading.lstrip("#"))
    start = report.index(heading)
    tail = report[start + len(heading) :]
    following = [
        match.start()
        for match in re.finditer(r"^#{1,6} ", tail, flags=re.MULTILINE)
        if len(tail[match.start() :].split(" ")[0]) <= level
    ]
    return tail[: following[0]] if following else tail


# One idea's chapter runs from its title to the next idea's, through the Reviews and
# Tournament headings the reference layout nests one level inside it. The chapters sit
# under "# Top ideas", so a chapter title is a level two and the next level two ends it
# -- either the next idea or, for the last one, the provenance appendix's first section.
_NEXT_IDEA = re.compile(r"^## ", re.MULTILINE)


def _idea_chapter(report: str, title: str) -> str:
    heading = f"\n## {title}\n"
    tail = report[report.index(heading) + len(heading) :]
    match = _NEXT_IDEA.search(tail)
    return tail[: match.start()] if match else tail


def _pdf_text(report: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(report)))
    return _flat("\n".join(page.extract_text() or "" for page in reader.pages))


def _docx_text(report: str) -> str:
    from docx import Document

    document = Document(BytesIO(render_docx(report)))
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts += [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return _flat("\n".join(parts))


# --- the common case: nothing to say, so nothing is said ---------------------


def test_a_run_with_no_adjudication_renders_no_governance_material(rich_session):
    report = compile_dossier(rich_session)

    assert "Governance adjudications" not in report
    assert "adjudicat" not in report.lower()
    assert "withdrawn" not in report.lower()


def test_the_governance_heading_never_appears_without_something_under_it(rich_session):
    """An empty heading is the failure mode a conditional section invites."""
    assert "Governance adjudications" not in compile_dossier(rich_session)

    _withdraw(rich_session, _population_ids(rich_session)[-1])
    block = _section(compile_dossier(rich_session), "## Governance adjudications")
    body = [line for line in block.splitlines() if line.strip()]

    assert body and not body[0].startswith("#")
    assert len(block.split()) > 50


# --- withdrawal: visibly removed, never silently absent ----------------------


def test_a_withdrawn_hypothesis_keeps_a_numbered_slot(rich_session):
    before = len(_population_ids(rich_session))
    victim = _withdraw(rich_session, _population_ids(rich_session)[-1])
    record = load_record(rich_session)

    report = compile_dossier(rich_session)
    slots = _idea_slots(report)

    assert victim not in [candidate.id for candidate in record.candidates]
    # The reader must be able to count what was generated, not what survived.
    assert len(slots) == before
    assert any(record.title_for(victim) in slot for slot in slots)


def test_a_withdrawn_hypothesis_is_named_as_withdrawn_not_merely_missing(rich_session):
    victim = _withdraw(rich_session, _population_ids(rich_session)[-1])
    title = load_record(rich_session).title_for(victim)

    report = compile_dossier(rich_session)
    slot = next(
        _section(report, f"##### {line}")
        for line in _idea_slots(report)
        if title in line
    )

    assert "withdraw" in slot.lower()
    assert ADJUDICATOR in slot


def test_a_withdrawn_hypothesis_gets_no_rank_and_no_deep_dive(rich_session):
    victim = _withdraw(rich_session, _population_ids(rich_session)[-1])
    record = load_record(rich_session)
    title = record.title_for(victim)

    report = compile_dossier(rich_session)

    assert f"\n## {title}\n" not in report
    ranks = re.findall(
        r"^Rank: (\d+)(?:, tied on Elo with .+)?$", report, flags=re.MULTILINE
    )
    assert [int(item) for item in ranks] == list(range(1, len(record.candidates) + 1))


def test_the_withdrawn_hypothesis_text_survives_its_removal(rich_session):
    victim = _withdraw(rich_session, _population_ids(rich_session)[0])
    record = load_record(rich_session)

    assert record.superseded_populations == 1
    assert victim in record.titles
    note = next(item for item in record.adjudications if item.candidate_id == victim)
    assert note.claim, "the withdrawn claim must be recoverable from superseded history"
    assert _flat(note.claim).rstrip(".") in _flat(compile_dossier(rich_session))


def test_a_single_rewrite_of_the_population_is_counted_in_words(rich_session):
    """One withdrawal read as "rewritten 1 time", which is a template showing through.

    Every other small count in the report is spelled, and a run that withdrew a
    hypothesis once has rewritten the population once -- the figure and the
    singular-as-plural noun together made the provenance note read as generated
    rather than written.
    """
    _withdraw(rich_session, _population_ids(rich_session)[0])
    report = _flat(compile_dossier(rich_session))

    assert "The candidate population was rewritten once after" in report
    assert "rewritten 1 time" not in report


def test_a_withdrawal_does_not_move_the_generate_stage_after_the_reviews(rich_session):
    """The stage table reports the workflow, not the order the session was written in.

    Withdrawal replaces the population and appends the replacement, so the table
    headed "what each stage produced" listed generate below all five reflect rows --
    which says the hypotheses were written after the reviews of them.
    """
    _withdraw(rich_session, _population_ids(rich_session)[0])
    stages = [note.stage for note in load_record(rich_session).provenance]

    assert stages.count("generate") == 1
    assert stages.index("generate") < stages.index("reflect")
    assert stages == sorted(stages, key=STAGES.index)


# --- every adjudication is reported, verbatim --------------------------------


@pytest.mark.parametrize(
    ("apply", "flaw", "adjudicator", "justification"),
    [
        (_withdraw, FLAW, ADJUDICATOR, JUSTIFICATION),
        (_override, SECOND_FLAW, SECOND_ADJUDICATOR, SECOND_JUSTIFICATION),
    ],
)
def test_an_adjudication_reprints_the_flaw_and_the_reason_word_for_word(
    rich_session, apply, flaw, adjudicator, justification
):
    apply(rich_session, _population_ids(rich_session)[1])

    report = _flat(compile_dossier(rich_session))

    assert _flat(flaw) in report
    assert _flat(justification) in report
    assert adjudicator in report


@pytest.mark.parametrize(
    ("apply", "justification"),
    [(_withdraw, JUSTIFICATION), (_override, SECOND_JUSTIFICATION)],
)
def test_the_reason_for_a_decision_is_printed_once_and_pointed_at_elsewhere(
    rich_session, apply, justification
):
    """The flaw travels with the idea; the reasoning that answered it does not.

    An overridden idea printed the flaw and the justification verbatim in the idea
    listing, again in its own Critical Flaws subsection and again in the governance
    block. Three copies of one decision make a reader compare them to establish that
    they are one decision.
    """
    apply(rich_session, _population_ids(rich_session)[1])
    report = _flat(compile_dossier(rich_session))

    assert report.count(_flat(justification)) == 1
    assert report.count("under Governance adjudications") >= 1


def test_every_adjudication_gets_its_own_reported_decision(rich_session):
    ids = _population_ids(rich_session)
    _override(rich_session, ids[0])
    _withdraw(rich_session, ids[-1])

    report = compile_dossier(rich_session)
    block = _section(report, "## Governance adjudications")
    decisions = [line for line in block.splitlines() if re.match(r"^### \d+\. ", line)]

    assert len(decisions) == len(load_record(rich_session).adjudications) == 2
    assert sum("Withdrawn" in line for line in decisions) == 1
    assert sum("Override" in line for line in decisions) == 1
    for adjudicator in (ADJUDICATOR, SECOND_ADJUDICATOR):
        assert adjudicator in block
    for flaw in (FLAW, SECOND_FLAW):
        assert _flat(flaw) in _flat(block)


def test_the_report_does_not_pass_off_a_typed_name_as_a_verified_one(rich_session):
    """``--adjudicator`` is free text and nothing checks it, so a report that says a
    fatal safety flaw was accepted by "a named person" is inviting the reader to read
    an unverified string as the record of who is accountable for the decision."""
    _override(rich_session, _population_ids(rich_session)[0])

    report = compile_dossier(rich_session)
    block = _flat(_section(report, "## Governance adjudications"))

    assert SECOND_ADJUDICATOR in block, "the name is still on the record"
    assert "does not authenticate" in block
    assert "as entered by whoever ran the adjudication" in block
    # And the claim is not re-made in the per-idea warning, which is printed
    # wherever the overridden idea appears.
    assert "a named person" not in _flat(report)


def test_the_resolution_of_each_adjudication_is_stated(rich_session):
    ids = _population_ids(rich_session)
    _override(rich_session, ids[0])
    _withdraw(rich_session, ids[-1])

    block = _flat(
        _section(compile_dossier(rich_session), "## Governance adjudications")
    )

    assert "Resolution: withdrawal" in block
    assert "Resolution: override" in block


def test_an_adjudication_names_its_adjudicator_where_the_name_does_work(rich_session):
    """The block ran "Resolution: <name> withdrew ...", "Adjudicated by: <name>",
    "Justification given by <name>" in four consecutive lines. Two of those attach the
    name to something -- the decision and the quotation -- and the label between them
    was a third copy of one string with nothing new attached to it."""
    _withdraw(rich_session, _population_ids(rich_session)[-1])

    block = _flat(
        _section(compile_dossier(rich_session), "## Governance adjudications")
    )

    assert block.count(ADJUDICATOR) == 2
    assert "Adjudicated by" not in block


def test_a_withdrawal_says_once_that_the_idea_never_reached_the_tournament(
    rich_session,
):
    """Ranklessness followed from one decision and was reported as though it were
    three findings: the resolution sentence, the paragraph under it in the governance
    block, and the paragraph under the notice in the ideas section all said it."""
    victim = _withdraw(rich_session, _population_ids(rich_session)[-1])
    title = load_record(rich_session).title_for(victim)

    report = compile_dossier(rich_session)
    block = _flat(_section(report, "## Governance adjudications"))
    slot = next(
        _flat(_section(report, f"##### {line}"))
        for line in _idea_slots(report)
        if title in line
    )

    for text in (block, slot):
        assert text.count("no rank and no Elo") == 1
        assert "absent from the ranking" not in text
        assert "It has no rank, no Elo" not in text
    # What does not follow from the resolution sentence is still said, once each.
    assert "absent from the per-idea sections" in block
    assert "no section of its own further down" in slot


# --- override: at least as loud as a withdrawal ------------------------------


def test_an_overridden_idea_cannot_be_read_without_meeting_the_flaw(rich_session):
    victim = _override(rich_session, _population_ids(rich_session)[0])
    title = load_record(rich_session).title_for(victim)

    report = compile_dossier(rich_session)
    dive = _idea_chapter(report, title)
    preamble = dive.split("### Idea Proposal")[0]

    assert _flat(SECOND_FLAW) in _flat(preamble), (
        "the accepted flaw must precede the idea itself, not follow it"
    )
    assert SECOND_ADJUDICATOR in preamble


def test_an_overridden_idea_carries_the_flaw_into_its_critical_flaws_summary(
    rich_session,
):
    victim = _override(rich_session, _population_ids(rich_session)[0])
    title = load_record(rich_session).title_for(victim)

    report = compile_dossier(rich_session)
    dive = _idea_chapter(report, title)
    critical = _section(dive, "##### 2. Critical Flaws")

    assert _flat(SECOND_FLAW) in _flat(critical)
    assert SECOND_ADJUDICATOR in critical


def test_an_overridden_idea_is_flagged_where_the_ideas_are_listed(rich_session):
    victim = _override(rich_session, _population_ids(rich_session)[0])
    record = load_record(rich_session)
    title = record.title_for(victim)
    claim = next(item.claim for item in record.candidates if item.id == victim)

    report = compile_dossier(rich_session)
    slot = next(
        _section(report, f"##### {line}")
        for line in _idea_slots(report)
        if title in line
    )

    assert _flat(SECOND_FLAW) in _flat(slot)
    # The claim must never be the first thing a reader of this idea meets.
    assert slot.index(SECOND_ADJUDICATOR) < slot.index(_flat(claim).rstrip("."))


def test_an_overridden_idea_still_competes(rich_session):
    """An override keeps the hypothesis live, so it must keep its rank and its dive."""
    victim = _override(rich_session, _population_ids(rich_session)[0])
    record = load_record(rich_session)

    report = compile_dossier(rich_session)

    assert victim in [candidate.id for candidate in record.candidates]
    assert f"\n## {record.title_for(victim)}\n" in report


# --- both together, and the unanswered case ----------------------------------


def test_withdrawal_and_override_are_reported_side_by_side(rich_session):
    ids = _population_ids(rich_session)
    overridden = _override(rich_session, ids[0])
    withdrawn = _withdraw(rich_session, ids[-1])
    record = load_record(rich_session)

    report = compile_dossier(rich_session)

    assert record.override_for(overridden) is not None
    assert record.override_for(withdrawn) is None
    assert [item.candidate_id for item in record.withdrawals] == [withdrawn]
    assert f"\n## {record.title_for(overridden)}\n" in report
    assert f"\n## {record.title_for(withdrawn)}\n" not in report
    for text in (FLAW, SECOND_FLAW, JUSTIFICATION, SECOND_JUSTIFICATION):
        assert _flat(text) in _flat(report)


def test_the_adjudication_lead_counts_answers_and_admits_an_unanswered_block(
    rich_session,
):
    """Two false statements opened this section. The count was the number of
    adjudications printed as a number of hypotheses carrying a fatal flaw, so two
    blocks answered against one hypothesis read as two hypotheses; and "each block was
    answered by hand" was asserted whether or not a block was still open."""
    ids = _population_ids(rich_session)
    _override(rich_session, ids[0], flaw=FLAW)
    _override(rich_session, ids[0], flaw=SECOND_FLAW)
    _block(rich_session, ids[1], THIRD_FLAW)
    record = load_record(rich_session)

    block = _flat(
        _section(compile_dossier(rich_session), "## Governance adjudications")
    )

    assert len(record.adjudications) == 2
    assert len({item.candidate_id for item in record.adjudications}) == 1
    assert "Two governance adjudications are recorded for this run" in block
    assert "hypotheses" not in block, "answers were counted as hypotheses"
    assert "Each block was answered by hand" not in block
    assert (
        "One further fatal finding has been recorded and not answered by anyone"
        in block
    )
    assert "listed under Unanswered governance blocks below" in block


def test_the_adjudication_lead_says_so_when_nothing_was_left_unanswered(rich_session):
    """The other side of the same branch: with no open blocker the section has to say
    that, rather than leave the reader to infer it from a heading that is not there."""
    _override(rich_session, _population_ids(rich_session)[0])

    block = _flat(
        _section(compile_dossier(rich_session), "## Governance adjudications")
    )

    assert "One governance adjudication is recorded for this run" in block
    assert "No other fatal finding was left unanswered in this run." in block
    assert "Unanswered governance blocks" not in block


def test_an_unanswered_governance_block_is_reported_as_blocking(rich_session):
    victim = _population_ids(rich_session)[2]
    _block(rich_session, victim, FLAW)
    record = load_record(rich_session)

    report = compile_dossier(rich_session)

    assert [item.candidate_id for item in record.open_governance_blocks] == [victim]
    assert not record.adjudications
    assert _flat(FLAW) in _flat(report)
    assert "blocked" in report.lower()
    # Nobody decided anything, so nothing may be filed under a decision heading.
    assert "## Governance adjudications" not in report
    assert "## Unanswered governance blocks" in report


def test_an_answered_block_is_no_longer_reported_as_open(rich_session):
    _withdraw(rich_session, _population_ids(rich_session)[-1])

    assert not load_record(rich_session).open_governance_blocks


# --- the same in every export format -----------------------------------------


def test_the_adjudication_reaches_pdf_and_docx(rich_session):
    ids = _population_ids(rich_session)
    _override(rich_session, ids[0])
    _withdraw(rich_session, ids[-1])
    report = compile_dossier(rich_session)

    for rendered in (_pdf_text(report), _docx_text(report)):
        assert "Governance adjudications" in rendered
        for adjudicator in (ADJUDICATOR, SECOND_ADJUDICATOR):
            assert adjudicator in rendered
        for text in (FLAW, SECOND_FLAW, JUSTIFICATION, SECOND_JUSTIFICATION):
            assert _flat(text) in rendered


def test_a_clean_run_carries_no_governance_heading_into_pdf_or_docx(rich_session):
    report = compile_dossier(rich_session)

    for rendered in (_pdf_text(report), _docx_text(report)):
        assert "Governance adjudications" not in rendered
