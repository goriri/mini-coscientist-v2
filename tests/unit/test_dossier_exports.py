import re
from io import BytesIO
from zipfile import ZipFile

from coscientist.dossier import render_docx, render_pdf
from coscientist.markdown_render import _unique_anchor

DOSSIER = """# Co-Scientist Research Dossier

**Question:** Does the intervention improve the measured outcome?

## Executive synthesis

- Candidate 1 remains testable.
- Independent replication is required.

## Decision audit

| Candidate | Decision |
| --- | --- |
| Candidate 1 | Conditional advance |
"""

RICH_DOSSIER = """# Co-Scientist Research Dossier

**Question:** 鉴定治疗 PD-1 耐药的非小细胞肺癌的 scRNA-seq 驱动新型协同靶点

## Section one

Inline **bold**, *italic*, `code_span`, [link](https://example.com/?a=1&b=2), 5 > 3.

> Quoted guidance.

---

1. First step
2. Second step
   - nested detail

| Rank | Candidate | Elo |
| ---: | --- | ---: |
| 1 | `cand_a` | 1500.0 |

```json
{"schema_version": "1.0"}
```

# Appendix

<details><summary>Validated typed payload</summary>

```json
{"payload": true}
```

</details>
"""


def test_pdf_export_is_a_valid_pdf():
    exported = render_pdf(DOSSIER)

    assert exported.startswith(b"%PDF-")
    assert len(exported) > 1000


def test_pdf_export_has_title_page_toc_and_page_numbers():
    from pypdf import PdfReader

    exported = render_pdf(RICH_DOSSIER)
    reader = PdfReader(BytesIO(exported))

    assert len(exported) > 6000
    assert len(reader.pages) >= 3
    title_page = reader.pages[0].extract_text() or ""
    assert "Prepared by AI co-scientist on" in title_page
    assert "research purposes only" in title_page
    # The title page carries no furniture; numbering starts on the contents page.
    assert "Page" not in title_page
    assert "Page 2 of " in (reader.pages[1].extract_text() or "")


def test_pdf_toc_entries_carry_resolved_page_numbers():
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(RICH_DOSSIER)))
    contents = reader.pages[1].extract_text() or ""

    assert "Contents" in contents
    for heading in ("Section one", "Appendix"):
        assert heading in contents

    # Outline entries prove the bookmark keys resolved during multiBuild. The
    # document's own opening H1 is not among them: it repeats the title page.
    # The outline nests, so it is read depth-first rather than as a flat list.
    def titles(entries) -> list[str]:
        found = []
        for entry in entries:
            if isinstance(entry, list):
                found.extend(titles(entry))
            else:
                found.append(str(entry["/Title"]))
        return found

    assert titles(reader.outline) == [
        "Section one",
        "Appendix",
        "Index of Figures and Tables",
    ]


LONG_DOSSIER = "# Long dossier\n\n**Question:** Does it paginate?\n\n" + "".join(
    f"# Chapter {chapter}\n\n"
    + "".join(
        f"## Section {chapter}.{section}\n\n" + "Filler sentence. " * 90 + "\n\n"
        for section in range(1, 4)
    )
    for chapter in range(1, 5)
)


def test_every_bookmark_lands_on_the_page_its_heading_is_on():
    """Every internal destination in the report resolved to the title page.

    The canvas that printed "Page n of m" was the widely copied reportlab recipe,
    whose showPage() calls _startPage() rather than Canvas.showPage(). The document's
    page counter therefore never advanced during the build, and bookmarkPage() bound
    all hundred-odd contents entries and cross-references to page one.
    """
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(LONG_DOSSIER)))

    def flatten(items):
        for item in items:
            if isinstance(item, list):
                yield from flatten(item)
            else:
                yield item

    entries = list(flatten(reader.outline))
    pages = [reader.get_destination_page_number(entry) + 1 for entry in entries]

    assert len(entries) == 16, "four chapters and their three sections each"
    assert min(pages) > 1, "no destination may point at the title page"
    assert len(set(pages)) > 1, "the destinations all resolved to one page"
    assert pages == sorted(pages), "bookmarks are out of document order"
    # And the reader is shown the outline rather than left to find the pane.
    assert reader.trailer["/Root"].get("/PageMode") == "/UseOutlines"


def test_the_folio_counts_every_page_of_the_finished_document():
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(LONG_DOSSIER)))
    total = len(reader.pages)

    assert total > 3
    for number, page in enumerate(reader.pages[1:], start=2):
        assert f"Page {number} of {total}" in (page.extract_text() or "")


def test_list_bullets_survive_being_copied_out_of_the_pdf():
    """reportlab encodes U+2022 as byte 0x7f, which viewers draw as a bullet and
    every text extraction reads as a control character."""
    from pypdf import PdfReader

    reader = PdfReader(
        BytesIO(render_pdf("# Doc\n\n## List\n\n- First item\n- Second item\n"))
    )
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert "\x7f" not in text
    assert text.count("•") == 2


def test_one_header_row_is_set_the_same_way_in_every_table_that_shares_it():
    """Column widths were weighted by body content alone, so the same two headers
    wrapped onto two lines in one table and fitted on one in the next."""
    from reportlab.pdfbase.pdfmetrics import stringWidth

    from coscientist.dossier import _column_widths
    from coscientist.markdown_render import Table

    header = ["Judge", "Rationale"]
    narrow = Table(header=header, rows=[["Model", "Short."]], aligns=["left", "left"])
    wide = Table(
        header=header,
        rows=[["Model", "A considerably longer rationale " * 6]],
        aligns=["left", "left"],
    )

    for block in (narrow, wide):
        for label, width in zip(header, _column_widths(block, 400.0), strict=True):
            assert width >= stringWidth(label, "Times-Bold", 8.5), (
                f"{label!r} is set narrower than the word itself"
            )


def test_pdf_export_renders_markup_rather_than_markdown_source():
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(RICH_DOSSIER)))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    for literal in ("**bold**", "| ---", "|---", "<details>", "<summary>", "```"):
        assert literal not in text
    assert "bold" in text
    assert "Quoted guidance." in text
    assert "鉴定治疗" in text


def test_pdf_export_uses_serif_body_type_not_eight_point():
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(RICH_DOSSIER)))
    fonts = {
        str(font.get("/BaseFont"))
        for page in reader.pages
        for font in page.get("/Resources", {}).get("/Font", {}).values()
    }

    assert any("Times" in name for name in fonts)
    assert any("STSong" in name for name in fonts)


def test_docx_export_is_google_docs_compatible():
    exported = render_docx(DOSSIER)

    assert exported.startswith(b"PK")
    with ZipFile(BytesIO(exported)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "Co-Scientist Research Dossier" in document_xml
    assert "Candidate 1 remains testable" in document_xml


def test_docx_export_uses_native_word_structures():
    from docx import Document

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    styles = [paragraph.style.name for paragraph in document.paragraphs]

    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "List Number" in styles
    assert "List Bullet 2" in styles
    assert [table.style.name for table in document.tables] == ["Table Grid"]

    header_cells = document.tables[0].rows[0].cells
    assert [cell.text for cell in header_cells] == ["Rank", "Candidate", "Elo"]
    assert all(run.bold for cell in header_cells for run in cell.paragraphs[0].runs)

    runs = [
        run
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text in {"bold", "italic", "code_span"}
    ]
    assert {run.text for run in runs} == {"bold", "italic", "code_span"}
    assert next(run for run in runs if run.text == "bold").bold
    assert next(run for run in runs if run.text == "italic").italic
    assert next(run for run in runs if run.text == "code_span").font.name == (
        "Courier New"
    )


def test_docx_tables_repeat_their_header_and_weight_their_columns():
    """The PDF repeats the header across pages, bands it grey and weights the
    columns. The DOCX did none of the three, so a table that broke over a page
    boundary carried on with no header, and a rank number was allotted the same
    width as a paragraph of rationale."""
    from docx import Document
    from docx.oxml.ns import qn

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    header_row = document.tables[0].rows[0]
    properties = header_row._tr.find(qn("w:trPr"))

    assert properties is not None and properties.find(qn("w:tblHeader")) is not None
    shading = header_row.cells[0]._tc.tcPr.find(qn("w:shd"))
    assert shading is not None and shading.get(qn("w:fill")) == "E8EAED"
    widths = [cell.width.pt for cell in header_row.cells]
    assert len(set(widths)) > 1, "every column was set to the same width"
    assert widths[1] > widths[0], "the candidate column is the one carrying prose"


def test_a_chapter_starts_a_page_in_both_exports():
    """The PDF breaks the page before every H1 and the DOCX did not, so the two
    exports of one report disagreed about where its chapters began."""
    from docx import Document

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    chapters = [p for p in document.paragraphs if p.style.name == "Heading 1"]

    assert [p.text for p in chapters] == ["Appendix"]
    assert all(p.paragraph_format.page_break_before for p in chapters)


def test_the_contents_caption_is_not_an_entry_in_its_own_contents():
    """Set as Heading 2 above a TOC field scoped to levels 1-3, the word "Contents"
    was the first line of the list it captioned."""
    from docx import Document

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    caption = next(p for p in document.paragraphs if p.text == "Contents")

    assert caption.style.name == "Normal"
    assert caption.runs[0].bold
    assert caption.paragraph_format.page_break_before


def test_docx_export_has_a_page_number_field():
    with ZipFile(BytesIO(render_docx(DOSSIER))) as archive:
        footer = archive.read("word/footer1.xml").decode("utf-8")

    assert "PAGE" in footer


def test_docx_export_opens_on_a_cover_page_that_does_not_repeat_as_a_heading():
    """The document's own opening H1 is the cover title. Emitted as Heading 1 as
    well, it appeared twice on the first two pages and again in the contents."""
    from docx import Document

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]

    assert headings == ["Appendix"]
    cover = next(
        p for p in document.paragraphs if p.text.startswith("Co-Scientist Research")
    )
    assert cover.style.name == "Normal"
    assert cover.runs[0].bold
    assert document.element.xml.count('w:type="page"') == 1, (
        "the cover runs into the contents, or the body starts on the cover"
    )


def test_docx_contents_is_a_field_word_can_refresh():
    """A contents list written out as text is stale the first time anyone edits."""
    with ZipFile(BytesIO(render_docx(RICH_DOSSIER))) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml
    assert document_xml.count("fldChar") >= 2


def test_docx_contents_works_before_any_field_is_updated():
    """Google Docs and LibreOffice do not evaluate a TOC field on import.

    An empty field leaves them with no contents list at all, permanently. The field is
    written with a cached result -- one internal hyperlink per heading -- so the list
    is usable in every reader, and Word replaces it with a paginated one on open.
    """
    with ZipFile(BytesIO(render_docx(RICH_DOSSIER))) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")

    assert 'w:fldCharType="separate"' in document_xml, "the field has no cached result"
    assert 'w:updateFields w:val="true"' in settings_xml
    anchors = re.findall(r'w:hyperlink w:anchor="([^"]+)"', document_xml)
    bookmarks = set(re.findall(r'w:bookmarkStart[^>]*w:name="([^"]+)"', document_xml))
    assert anchors, "the cached contents list is empty"
    assert set(anchors) <= bookmarks, "a contents entry points at no heading"


def test_docx_export_carries_links_as_relationships_rather_than_as_text():
    from docx import Document

    document = Document(BytesIO(render_docx(RICH_DOSSIER)))
    external = [
        rel.target_ref for rel in document.part.rels.values() if rel.is_external
    ]

    assert "https://example.com/?a=1&b=2" in external
    assert "](http" not in "\n".join(p.text for p in document.paragraphs)


# A researchgate locator 254 characters long, which is wider than the text column of
# either export and so has to be broken across three lines to be set at all.
WRAPPING_URL = (
    "https://www.researchgate.net/publication/317195885_Tristrimethylsilyl_Phosphite"
    "_TMSPi_and_Triethyl_Phosphite_TEPi_as_Electrolyte_Additives_for_Lithium_Ion_"
    "Batteries_Mechanistic_Insights_into_Differences_during_LiNi_05_Mn_03_Co_02_O_2_"
    "-Graphite_Full_Cells"
)

WRAPPING_LINK_DOSSIER = f"""# Co-Scientist Research Dossier

**Question:** Does a reference too wide for the column survive being set?

## References

1. Untitled source on researchgate.net. [{WRAPPING_URL}]({WRAPPING_URL})
"""


def test_a_reference_too_wide_for_the_column_links_to_the_whole_url():
    """A reference URL that wraps must still target the URL, not a piece of it.

    An audit of the rendered PDF found three separate link annotations over the one
    researchgate reference and read them as a linkifier matching sub-spans of the
    locator, each pointing somewhere dead. They are the three lines the one link is
    set on, and all three carry the whole target -- which is what this pins, because
    a linkifier that really did split on the underscores would leave targets ending
    mid-word and no other test would notice.
    """
    from docx import Document
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(render_pdf(WRAPPING_LINK_DOSSIER)))
    targets = [
        str(annotation.get_object()["/A"]["/URI"])
        for page in reader.pages
        for annotation in (page.get("/Annots") or [])
        if "/URI" in (annotation.get_object().get("/A") or {})
    ]

    assert targets, "the reference was set without a link at all"
    assert set(targets) == {WRAPPING_URL}

    document = Document(BytesIO(render_docx(WRAPPING_LINK_DOSSIER)))
    external = [
        rel.target_ref for rel in document.part.rels.values() if rel.is_external
    ]

    assert external == [WRAPPING_URL]


def test_each_ordered_list_starts_again_at_one():
    """One numbering instance is counted across the whole document, so the second
    ordered list opened where the first left off -- references at 5 through 8."""
    from docx import Document
    from docx.oxml.ns import qn

    document = Document(
        BytesIO(
            render_docx(
                "# Report\n\n## First\n\n1. One\n2. Two\n\n## Second\n\n1. One again\n"
            )
        )
    )
    numbered = [p for p in document.paragraphs if p.style.name == "List Number"]
    assert len(numbered) == 3

    ids = [
        p._p.pPr.find(qn("w:numPr")).find(qn("w:numId")).get(qn("w:val"))
        for p in numbered
    ]
    assert ids[0] == ids[1] != ids[2], "the two lists share one running sequence"
    numbering = document.part.numbering_part.element
    for num_id in {ids[0], ids[2]}:
        instance = numbering.num_having_numId(int(num_id))
        override = instance.find(qn("w:lvlOverride"))
        start = override.find(qn("w:startOverride")) if override is not None else None
        assert start is not None and start.get(qn("w:val")) == "1"


def test_docx_headings_are_set_in_the_face_the_exporter_asks_for():
    """``Font.name`` writes ``w:ascii``; Word's template also carries ``w:asciiTheme``
    on the same element, and the theme attribute wins. Every heading in the report
    came out in the theme's Calibri over a Times body."""
    with ZipFile(BytesIO(render_docx(RICH_DOSSIER))) as archive:
        styles_xml = archive.read("word/styles.xml").decode("utf-8")

    for level in range(1, 7):
        block = re.search(
            rf'<w:style [^>]*w:styleId="Heading{level}".*?</w:style>', styles_xml, re.S
        )
        assert block, f"Heading {level} is not defined"
        fonts = re.search(r"<w:rFonts[^/]*/>", block.group()).group()
        assert "Theme" not in fonts, f"Heading {level} still resolves through the theme"
        assert "Times New Roman" in fonts
        # Heading 4 and Heading 6 are italic in the template; the PDF sets them roman.
        assert '<w:i w:val="0"/>' in block.group()


def test_the_two_exports_agree_on_the_paper_they_are_set_on():
    """The DOCX inherited python-docx's US Letter while the PDF is A4, so one report
    paginated two ways and the table arithmetic was handed a measure it did not have."""
    with ZipFile(BytesIO(render_docx(RICH_DOSSIER))) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert '<w:pgSz w:w="11906" w:h="16838"/>' in document_xml


def test_docx_table_cells_are_set_at_the_size_their_columns_were_measured_at():
    """Widths are computed at 8.5pt and the cells inherited Normal's 11pt, so headers
    overflowed their columns and Word broke them mid-word -- the defect the width
    arithmetic exists to prevent, fixed in the PDF only. The grid is written too:
    LibreOffice and Docs lay out from ``w:tblGrid``, which python-docx leaves equal."""
    from docx.oxml.ns import qn

    with ZipFile(BytesIO(render_docx(RICH_DOSSIER))) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    grid = re.search(r"<w:tblGrid>(.*?)</w:tblGrid>", document_xml, re.S)
    columns = re.findall(r'w:w="(\d+)"', grid.group(1))
    assert len(set(columns)) > 1, "the grid is still at equal column widths"
    assert '<w:tblW w:type="auto"' not in document_xml
    assert f'<w:sz w:val="{round(8.5 * 2)}"/>' in document_xml

    from docx import Document

    table = Document(BytesIO(render_docx(RICH_DOSSIER))).tables[0]
    widths = [cell.width for cell in table.rows[0].cells]
    total = sum(int(width) for width in columns)
    emu = sum(w for w in widths)
    assert abs(total * 12700 / 20 - emu) < 12700, "grid and cell widths disagree"
    assert table._tbl.tblPr.find(qn("w:tblW")).get(qn("w:type")) == "dxa"


def test_the_same_table_is_set_at_the_same_widths_wherever_it_appears():
    """Weighted per table, "Judge" was set anywhere from 56 to 85pt across eight
    tournament grids that are the same grid. A moved column rule reads as data."""
    from coscientist.dossier import _column_widths, _shared_column_widths
    from coscientist.markdown_render import Table

    header = ["Judge", "Rationale"]
    narrow = Table(header=header, rows=[["Model", "Short."]], aligns=["left", "left"])
    wide = Table(
        header=header,
        rows=[["Model", "A considerably longer rationale " * 6]],
        aligns=["left", "left"],
    )
    other = Table(header=["Only"], rows=[["one"]], aligns=["left"])

    shared = _shared_column_widths([narrow, wide, other], 400.0)
    assert _column_widths(narrow, 400.0, shared) == _column_widths(wide, 400.0, shared)
    assert ("Only",) not in shared, "a table with no twin is still measured on its own"


def test_a_chapter_does_not_end_on_two_orphaned_bullets():
    """A chapter ends on a forced page break, so whatever spills becomes the whole of
    its last page: the live report closed one chapter with two bullets and nine inches
    of white. The closing section moves down whole instead."""
    from pypdf import PdfReader

    dossier = (
        "# Report\n\n**Question:** Does it break well?\n\n"
        "# Chapter one\n\n## Long section\n\n"
        + "Filler sentence. "
        * 260
        + "\n\n## Closing section\n\n- First point.\n- Second point.\n\n"
        "# Chapter two\n\nShort.\n"
    )
    reader = PdfReader(BytesIO(render_pdf(dossier)))
    pages = [page.extract_text() or "" for page in reader.pages]
    closing = next(page for page in pages if "First point." in page)

    assert "Closing section" in closing, "the heading was left on the previous page"
    assert "Second point." in closing, "the list was split across the break"


def test_pdf_export_renders_tables_as_reportlab_table_objects():
    table_dossier = """# Title

| Col A | Col B |
| --- | --- |
| Val 1 | Val 2 |
"""
    exported = render_pdf(table_dossier)
    assert exported.startswith(b"%PDF-")
    assert len(exported) > 1000


def test_compile_dossier_toc_alignment_and_evidence_discovery():
    import re

    from coscientist.dossier import compile_dossier
    from coscientist.models import ApprovalProfile
    from coscientist.orchestration import CoScientistWorkflow

    flow = CoScientistWorkflow(
        "Design a 2026-State-of-the-Art Synthesis Strategy for a 45-mer Hydrophobic Therapeutic Peptide",
        approval_profile=ApprovalProfile.AUTO,
        workflow_version=1,
    )
    flow.accept_literature_only()
    flow.run_auto()
    from coscientist.models import Artifact, DeepResearchRun, DiscoveryManifest

    manifest = DiscoveryManifest(
        question=flow.session.question,
        runs=[DeepResearchRun(pass_number=1, status="completed")],
    )
    flow.session.artifacts.append(
        Artifact(
            stage="evidence",
            agent="deep_research_discovery",
            artifact_type="specialist_output",
            content="Evidence discovered.",
            schema_name="DiscoveryManifest",
            payload=manifest.model_dump(mode="json"),
        )
    )
    report_md = compile_dossier(flow.session)

    # The contents list, not every bracketed link in the document: the index of
    # exhibits at the back is written in the same notation and is not part of it.
    contents = report_md.split("## Table of Contents", 1)[1].split("\n## ", 1)[0]
    entries = re.findall(r"^\s*-\s*\[([^\]]+)\]\((#[^)]*)\)$", contents, re.M)
    assert entries, "the dossier carries no contents list"

    anchors: dict[str, int] = {}
    headings: list[tuple[int, str, str]] = []
    for match in re.finditer(r"^(#{1,2}) (.+)$", report_md, re.M):
        text = match.group(2).strip()
        headings.append(
            (len(match.group(1)), text, "#" + _unique_anchor(text, anchors))
        )

    # Every entry points at a heading that exists. A contents list whose links
    # resolve to nothing is worse than none: it looks navigable and is not.
    by_anchor = {anchor: text for _, text, anchor in headings}
    for title, anchor in entries:
        assert anchor in by_anchor, f"contents entry {title!r} points at nothing"
        assert by_anchor[anchor] == title

    # And every chapter and section is in it, bar the two that would be noise:
    # the document's own title, which the entry would sit directly under, and
    # the contents list itself.
    listed = {anchor for _, anchor in entries}
    for index, (level, text, anchor) in enumerate(headings):
        if index == 0 or text == "Table of Contents":
            continue
        assert anchor in listed, f"heading {text!r} (h{level}) is not in the contents"

    assert "## Literature discovery" in report_md
    discovery_section = report_md.split("## Literature discovery", 1)[1].split("\n## ")[
        0
    ]
    assert "Deep Research ran one pass" in discovery_section
    assert "source lead" in discovery_section

    main_section = report_md.split("## Complete artifact appendix")[0]
    assert not re.search(r"\bcandidate_[a-z0-9_-]+\b", main_section, re.IGNORECASE)


def test_render_pdf_clickable_toc_mermaid_and_index():
    md = """# Co-Scientist Research Dossier

## Table of Contents

- [Executive synthesis (Executive Summary)](#executive-synthesis-executive-summary)
- [Evidence Discovery](#evidence-discovery)

## Executive synthesis (Executive Summary)

Summary content here.
| Col A | Col B |
| --- | --- |
| Val 1 | Val 2 |

## Evidence Discovery

```mermaid
graph TD
    A[Baseline Strategy] --> B(Evaluation Stage)
```
"""
    from pypdf import PdfReader

    pdf_bytes = render_pdf(md)
    assert pdf_bytes.startswith(b"%PDF-")
    assert len(pdf_bytes) > 2500

    assert b"/Link" in pdf_bytes or b"/Annot" in pdf_bytes
    assert b"/Dest" in pdf_bytes or b"/Dests" in pdf_bytes or b"/Names" in pdf_bytes

    # Read the pages rather than the file. ReportLab deflates its content
    # streams, so a substring search over the bytes finds only what happens to
    # live outside them -- it would pass on a document whose body was blank.
    reader = PdfReader(BytesIO(pdf_bytes))
    rendered = "\n".join((page.extract_text() or "") for page in reader.pages)
    assert "Index of Figures and Tables" in rendered
    assert "Figure 1." in rendered
    assert "Table 1." in rendered
    assert "Page" in rendered


def test_every_export_tells_the_reader_what_the_document_is(rich_session):
    """The notice was cover matter in the PDF and the DOCX and nowhere at all in
    the Markdown -- the copy most likely to be pasted into a message. A reader
    holding the .md had to reach an appendix two thousand seven hundred lines down
    to learn that nothing in it is a finding."""
    from pypdf import PdfReader

    from coscientist.dossier import _DEFAULT_NOTICE, compile_dossier
    from coscientist.markdown_render import CONTENTS_HEADING

    markdown = compile_dossier(rich_session)

    assert markdown.count(_DEFAULT_NOTICE) == 1
    # Above the contents list, because forty links between the title and that
    # sentence is the same thing as not putting it on the cover.
    assert markdown.index(_DEFAULT_NOTICE) < markdown.index(CONTENTS_HEADING)

    # And set once in the two formats that already had it, not twice: they build
    # their own title page out of the same compiled Markdown.
    pdf = "\n".join(
        (page.extract_text() or "")
        for page in PdfReader(BytesIO(render_pdf(markdown))).pages
    )
    assert pdf.count("Draft for internal review") == 1

    with ZipFile(BytesIO(render_docx(markdown))) as archive:
        document = archive.read("word/document.xml").decode("utf-8")
    assert document.count("Draft for internal review") == 1


def test_the_pdf_sets_the_characters_the_markdown_wrote():
    """Every one of these came off a live ninety-two page PDF as a filled box or as
    nothing at all. The star is the worst of the set: it marks the shortlisted ideas
    in the executive table, it sat inside the pictograph drop range, and so the
    caption saying "a star marks an idea the tournament shortlisted" printed above a
    column with nothing in it on every page it appeared."""
    from coscientist.dossier import _renderable

    assert _renderable("Candidate 1 ★") == "Candidate 1 *"
    assert _renderable("signiﬁcant") == "significant"
    assert _renderable("1\u20105 nm") == "1-5 nm"
    assert _renderable("ΔG ≤ 0") == "DeltaG <= 0"

    # Decoration goes, and the gap it leaves goes with it.
    assert _renderable("Prepared \U0001f9ec by the panel") == "Prepared by the panel"

    # Anything rarer is named rather than dropped, because a character silently
    # missing from a formula is a different formula.
    assert _renderable("E = ∮ B") == "E = (contour integral) B"

    # ACS sets the dash in a paper's name as a box-drawing rule, so a title reaches
    # the reference list with one in. Named by the rule above rather than folded, it
    # printed the character's own Unicode name mid-title on a live page.
    assert (
        _renderable("Degradation Effects in Cells─Learning from Potential Profiles")
        == "Degradation Effects in Cells—Learning from Potential Profiles"
    )
    assert "box drawings" not in _renderable("Cells━Learning ⸺ and ⸻ and ―")

    # The CID face draws these and _shift_scripts turns those into markup: neither
    # is the body face's problem, and folding either would corrupt the page.
    passed = "锂酸锂 Li₂CO₃"
    assert _renderable(passed) == passed


def test_the_shortlist_star_reaches_the_rendered_page():
    from pypdf import PdfReader

    dossier = (
        "# Title\n\n## Executive Candidate Summary\n\n"
        "| Rank | Candidate |\n| --- | --- |\n"
        "| 1 | A 2 nm Coating ★ |\n\n"
        "A star marks an idea the tournament shortlisted.\n"
    )
    text = "\n".join(
        (page.extract_text() or "")
        for page in PdfReader(BytesIO(render_pdf(dossier))).pages
    )

    assert "A 2 nm Coating *" in text


DIAGRAM = """graph TD
    A[High Voltage] --> B{Cathode Surface}
    B -->|Uncoated| C[Oxygen Evolution]
    C --> D[TM Migration]
    B -->|Coated| G[Al-O-TM Bonds]
    G --> I[Suppressed Migration]
"""


def _story(markdown: str) -> list:
    from coscientist.dossier import _pdf_styles, _register_pdf_fonts, _story_from_blocks
    from coscientist.markdown_render import parse_blocks

    _register_pdf_fonts()
    return _story_from_blocks(parse_blocks(markdown), _pdf_styles(), 460.0)


def test_a_heading_is_bound_past_its_padding_to_the_figure_it_heads():
    """Four "Proposed Workflow" headings printed at the foot of a page with the
    flowchart they head overleaf. The heading style keeps with the next flowable and
    always had -- but for a diagram that next flowable is the six-point spacer the
    drawing is padded with, so the bind was satisfied by the padding and reportlab
    was free to break underneath it.
    """
    from reportlab.graphics.shapes import Drawing
    from reportlab.platypus import Spacer

    story = _story("### Proposed Workflow\n\n```mermaid\n" + DIAGRAM + "```\n")
    kept = [flowable.getKeepWithNext() for flowable in story]

    assert isinstance(story[1], Spacer) and isinstance(story[2], Drawing), (
        "the padding this test is about is no longer where the fix reaches"
    )
    assert kept[:3] == [1, 1, 0], "the heading is bound no further than its own padding"


def test_a_lead_in_is_bound_to_the_list_it_announces():
    """ "Evidence gaps:" was the last line of page 51 and the four gaps it announces
    were on page 52. A colon is a promise about the next line."""
    story = _story("**Evidence gaps:**\n\n- The vacancy energy is unmeasured.\n")

    assert story[0].getKeepWithNext(), "the lead-in can still be left on its own"


def test_a_caption_is_bound_to_the_exhibit_it_names():
    story = _story(
        "**Table 4.** Match summary.\n\n| Judge | Verdict |\n| --- | --- |\n| A | Held |\n"
    )

    assert story[0].getKeepWithNext(), "the caption can still be left on its own"


def test_a_paragraph_that_announces_nothing_is_not_bound_to_what_follows():
    """Binding every paragraph would push text down the page for no reason: the fix
    is for the ones that promise something, not for prose that happens to precede a
    table."""
    story = _story("The tournament ran three rounds.\n\n| A |\n| --- |\n| 1 |\n")

    assert not story[0].getKeepWithNext()


def test_the_index_of_exhibits_gives_a_page_and_a_link_for_every_entry():
    """The Markdown index locates its entries by anchor, which is what a Markdown
    reader follows and what a PDF has no notion of. Flattened for export, all
    twenty-nine entries of a live report printed over two pages with no page number
    against any of them and nothing to click."""
    from pypdf import PdfReader

    dossier = (
        "# Report\n\n## Findings\n\n| Rank | Idea |\n| --- | --- |\n| 1 | Coating |\n\n"
        + "Filler sentence about the tournament. " * 400
        + "\n\n## Later findings\n\n| Rank | Idea |\n| --- | --- |\n| 2 | Doping |\n"
    )
    reader = PdfReader(BytesIO(render_pdf(dossier)))
    pages = {
        page.indirect_reference.idnum: number
        for number, page in enumerate(reader.pages, start=1)
    }
    index = reader.pages[-1]
    text = index.extract_text() or ""

    entries = dict(re.findall(r"\.\s(\d+)\n(Table \d+) —", text))
    assert {"Table 1", "Table 2"} == set(entries.values())
    printed = {label: int(page) for page, label in entries.items()}
    assert printed["Table 1"] != printed["Table 2"], (
        "both exhibits were listed against the same page"
    )
    for label, page in printed.items():
        assert label in (reader.pages[page - 1].extract_text() or ""), (
            f"{label} is not on the page the index sends the reader to"
        )

    # And the entries are links, not text that happens to name a page.
    targets = {
        pages[annotation.get_object()["/Dest"][0].idnum]
        for annotation in (index.get("/Annots") or [])
    }
    assert targets == set(printed.values())


def test_a_run_that_waited_overnight_is_not_measured_out_in_minutes():
    """Provenance printed "5463 minutes after it started" over two timestamps four
    days apart. The number is right and nobody converts it, and the two dates it
    sits between have already said the same thing in a unit a reader holds."""
    from coscientist.dossier import _elapsed

    assert _elapsed("2026-08-14T14:28:47+00:00", "2026-08-18T09:31:44+00:00") == (
        ", 3 days and 19 hours after it started"
    )
    assert _elapsed("2026-08-14T14:28:47+00:00", "2026-08-14T20:15:47+00:00") == (
        ", 5 hours and 47 minutes after it started"
    )
    # The remainder is dropped where there is none, rather than printed as a zero.
    assert _elapsed("2026-08-14T14:28:47+00:00", "2026-08-15T14:28:47+00:00") == (
        ", 1 day after it started"
    )
    # And the two units that already read well are left as they were.
    assert _elapsed("2026-08-14T14:28:47+00:00", "2026-08-14T14:58:47+00:00") == (
        ", 30 minutes after it started"
    )
    assert _elapsed("2026-08-14T14:28:47+00:00", "2026-08-14T14:29:20+00:00") == (
        ", 33 seconds after it started"
    )
